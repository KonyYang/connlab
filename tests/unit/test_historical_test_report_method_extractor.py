from __future__ import annotations

from pathlib import Path

from docx import Document

from backend.infrastructure.office.historical_test_report_method_extractor import (
    HistoricalTestReportMethodExtractor,
)


def test_extractor_reads_methods_requirement_table_with_section5_heading(tmp_path: Path) -> None:
    source = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("5. TEST METHODS/REQUIREMENTS")
    table = document.add_table(rows=3, cols=4)
    table.cell(0, 0).text = "Test Item"
    table.cell(0, 1).text = "Method"
    table.cell(0, 2).text = "Condition"
    table.cell(0, 3).text = "Requirement"
    table.cell(1, 0).text = "Contact Resistance (Low Level)"
    table.cell(1, 1).text = "EIA-364-23D"
    table.cell(1, 2).text = "20mV max, 100mA max"
    table.cell(1, 3).text = "Initial <= 0.25 mΩ; ΔR <= 0.17 mΩ"
    table.cell(2, 0).text = "Temperature rise"
    table.cell(2, 1).text = "EIA-364-70"
    table.cell(2, 2).text = "75 A"
    table.cell(2, 3).text = "<= 30 ℃"
    document.save(source)

    result = HistoricalTestReportMethodExtractor().extract(source)

    assert result.source_table_index == 1
    assert len(result.rows) == 2
    assert result.rows[0].test_item == "Contact Resistance (Low Level)"


def test_extractor_skips_same_shape_table_without_heading_evidence(tmp_path: Path) -> None:
    source = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("7. APPENDIX")
    table = document.add_table(rows=2, cols=4)
    table.cell(0, 0).text = "Test Item"
    table.cell(0, 1).text = "Method"
    table.cell(0, 2).text = "Condition"
    table.cell(0, 3).text = "Requirement"
    table.cell(1, 0).text = "Dummy"
    table.cell(1, 1).text = "EIA-000"
    table.cell(1, 2).text = "N/A"
    table.cell(1, 3).text = "N/A"
    document.save(source)

    result = HistoricalTestReportMethodExtractor().extract(source)

    assert result.source_table_index is None
    assert result.rows == ()


def test_extractor_handles_bilingual_header_variant(tmp_path: Path) -> None:
    source = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("5. TEST METHODS/REQUIREMENTS")
    table = document.add_table(rows=2, cols=4)
    table.cell(0, 0).text = "Test Item 测试项目"
    table.cell(0, 1).text = "Test Method/Standard"
    table.cell(0, 2).text = "Test Condition 条件"
    table.cell(0, 3).text = "Judgement Criteria"
    table.cell(1, 0).text = "Mating/Un-mating Force"
    table.cell(1, 1).text = "EIA-364-13"
    table.cell(1, 2).text = "25.4 mm/min"
    table.cell(1, 3).text = "Mating <= 20N; Un-mating >= 6N"
    document.save(source)

    result = HistoricalTestReportMethodExtractor().extract(source)

    assert result.source_table_index == 1
    assert len(result.rows) == 1
    assert result.rows[0].method == "EIA-364-13"


def test_extractor_chooses_target_table_when_non_target_table_precedes(tmp_path: Path) -> None:
    source = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("3. SCOPE")
    non_target = document.add_table(rows=2, cols=4)
    non_target.cell(0, 0).text = "Test Item"
    non_target.cell(0, 1).text = "Method"
    non_target.cell(0, 2).text = "Condition"
    non_target.cell(0, 3).text = "Requirement"
    non_target.cell(1, 0).text = "Not target"
    non_target.cell(1, 1).text = "X"
    non_target.cell(1, 2).text = "Y"
    non_target.cell(1, 3).text = "Z"
    document.add_paragraph("5. TEST METHODS/REQUIREMENTS")
    target = document.add_table(rows=2, cols=4)
    target.cell(0, 0).text = "Test Item"
    target.cell(0, 1).text = "Method"
    target.cell(0, 2).text = "Condition"
    target.cell(0, 3).text = "Requirement"
    target.cell(1, 0).text = "LLCR"
    target.cell(1, 1).text = "EIA-364-23D"
    target.cell(1, 2).text = "20mV max, 100mA max"
    target.cell(1, 3).text = "Initial <= 0.25 mΩ; ΔR <= 0.17 mΩ"
    document.save(source)

    result = HistoricalTestReportMethodExtractor().extract(source)

    assert result.source_table_index == 2
    assert len(result.rows) == 1
    assert result.rows[0].test_item == "LLCR"
