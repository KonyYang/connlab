from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

import backend.infrastructure.office.word_document_gateway as word_gateway_module
from backend.infrastructure.office.office_lifecycle import OfficeAutomationUnavailable
from backend.infrastructure.office.word_document_gateway import WordDocumentGateway


def test_application_form_plain_table_writes_same_row_and_next_row_fields(
    tmp_path: Path,
) -> None:
    target = tmp_path / "application.docx"
    _write_application_form_plain_docx(target)

    result = WordDocumentGateway().write_application_form_fields(
        target,
        {
            "project_type": "New Product Development",
            "lab": "Dongguan Lab",
            "project_leader": "BI Leader",
            "received_date": "20 Jun 2026",
            "estimated_completion_date": "30 Jun 2026",
            "sample_condition": "Acceptable",
        },
    )

    assert {field.field_key for field in result.changed_fields} == {
        "project_type",
        "lab",
        "project_leader",
        "received_date",
        "estimated_completion_date",
        "sample_condition",
    }
    values = _read_application_form_plain_docx(target)
    assert values["Project Type"] == "New Product Development"
    assert values["Lab Performing the Tests:"] == "Dongguan Lab"
    assert values["Lab Personnel Assigned:"] == "BI Leader"
    assert values["Date Lab Received Samples:"] == "20 Jun 2026"
    assert values["Estimated Completion Date:"] == "30 Jun 2026"
    assert values["Condition of Samples when Received:"] == "Acceptable"


def test_application_form_plain_table_blocks_missing_critical_field(
    tmp_path: Path,
) -> None:
    target = tmp_path / "application.docx"
    _write_application_form_plain_docx(target, include_lab=False)

    with pytest.raises(ValueError, match="lab"):
        WordDocumentGateway().write_application_form_fields(
            target,
            {
                "lab": "Dongguan Lab",
            },
        )


def test_application_form_requires_com_does_not_fallback_to_false_success(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    target = tmp_path / "application.docx"
    _write_application_form_plain_docx(target)

    monkeypatch.setattr(
        word_gateway_module,
        "application_form_requires_com",
        lambda path: True,
    )

    def _raise_unavailable(path: Path, fields: dict[str, str]):
        raise OfficeAutomationUnavailable("Word COM automation requires pywin32.")

    monkeypatch.setattr(
        word_gateway_module,
        "write_application_form_fields_with_com",
        _raise_unavailable,
    )

    with pytest.raises(OfficeAutomationUnavailable):
        WordDocumentGateway().write_application_form_fields(
            target,
            {"requested_by": "Ming-Peng.Cao"},
        )


def _write_application_form_plain_docx(
    path: Path,
    *,
    include_lab: bool = True,
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    next_row = document.add_table(rows=2, cols=1)
    next_row.cell(0, 0).text = "Project Type"
    same_row = document.add_table(rows=3, cols=4)
    if include_lab:
        same_row.cell(0, 0).text = "Lab Performing the Tests:"
    same_row.cell(0, 2).text = "Lab Personnel Assigned:"
    same_row.cell(1, 0).text = "Date Lab Received Samples:"
    same_row.cell(1, 2).text = "Estimated Completion Date:"
    same_row.cell(2, 0).text = "Condition of Samples when Received:"
    document.save(path)


def _read_application_form_plain_docx(path: Path) -> dict[str, str]:
    document = Document(path)
    next_row = document.tables[0]
    same_row = document.tables[1]
    return {
        next_row.cell(0, 0).text.strip(): next_row.cell(1, 0).text.strip(),
        same_row.cell(0, 0).text.strip(): same_row.cell(0, 1).text.strip(),
        same_row.cell(0, 2).text.strip(): same_row.cell(0, 3).text.strip(),
        same_row.cell(1, 0).text.strip(): same_row.cell(1, 1).text.strip(),
        same_row.cell(1, 2).text.strip(): same_row.cell(1, 3).text.strip(),
        same_row.cell(2, 0).text.strip(): same_row.cell(2, 1).text.strip(),
    }
