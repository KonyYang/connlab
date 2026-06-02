from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from backend.application.confirmed_matrix_test_record_document_generation_service import (
    TestRecordHeaderMetadata,
)
from backend.application.test_record_fee_dataset_preview_service import (
    TestRecordGroupDataset as RecordGroup,
    TestRecordStepDataset as RecordStep,
)
from backend.infrastructure.office.test_record_document_gateway import TestRecordDocumentGateway


def test_test_record_gateway_generates_docx(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    Document().save(template)
    output = tmp_path / "generated.docx"

    step = RecordStep(
        sequence=1,
        test_item="LLCR",
        condition_summary="After conditioning",
        method_summary="Measure",
        reference_standard="EIA-364-23",
        judgement_criteria="20 mOhm max",
        duration_hint="1 day(s)",
        source_section="5.4",
        source_table_index=21,
        source_row_index=5,
        warnings=(),
    )
    group = RecordGroup(
        group_key="group_1",
        group_label="Group 1",
        source_table_index=21,
        steps=(step,),
        warnings=(),
    )

    result = TestRecordDocumentGateway().generate(
        template_path=template,
        output_path=output,
        source_document_name="spec.docx",
        groups=(group,),
        warnings=("missing duration",),
    )

    assert output.exists()
    assert result.status == "generated"


def test_test_record_gateway_rejects_non_docx_template(tmp_path: Path) -> None:
    template = tmp_path / "template.txt"
    template.write_text("x", encoding="utf-8")

    with pytest.raises(ValueError, match="Only .docx"):
        TestRecordDocumentGateway().generate(
            template_path=template,
            output_path=tmp_path / "out.docx",
            source_document_name="spec.docx",
            groups=(),
            warnings=(),
        )


def test_gateway_generates_confirmed_matrix_test_record_docx(tmp_path: Path) -> None:
    template = _build_confirmed_matrix_template(tmp_path / "template.docx")
    output = tmp_path / "confirmed-record.docx"
    group = _ConfirmedGroup()

    result = TestRecordDocumentGateway().generate_from_confirmed_matrix(
        template_path=template,
        output_path=output,
        project_id="P1",
        project_no="DL-001",
        product_description="Connector",
        applicable_specification="GS-12-1507",
        confirmed_matrix_id="cmv-1",
        groups=(group,),
        header_metadata=TestRecordHeaderMetadata(),
    )

    assert result == output
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert (
        "Group Number 组别编号: 1 ;   "
        "Sample Quantity & Number 样品数量及编号: 5 sets (Group1-1#~5#)"
    ) in text
    assert len(document.tables) == 2
    step_table = document.tables[0]
    assert step_table.rows[1].cells[0].text == "1"
    assert step_table.rows[1].cells[1].text == "Visual"
    assert step_table.rows[1].cells[2].text == "EIA-364-18"
    assert step_table.rows[1].cells[3].text == "10x"
    assert step_table.rows[1].cells[8].text == ""
    assert step_table.rows[1].cells[4].text == ""
    assert step_table.rows[1].cells[7].text == ""
    equipment_table = document.tables[1]
    assert equipment_table.rows[1].cells[0].text == "EQUIPMENT USED 使用的设备:"


def test_gateway_repeats_template_group_blocks_for_multiple_groups(tmp_path: Path) -> None:
    template = _build_confirmed_matrix_template(tmp_path / "template.docx")
    output = tmp_path / "confirmed-record.docx"
    group1 = _ConfirmedGroup()
    group2 = _ConfirmedGroup(group_key="g2", group_label="Group 2", sample_quantity_expression="7")

    TestRecordDocumentGateway().generate_from_confirmed_matrix(
        template_path=template,
        output_path=output,
        project_id="P1",
        project_no="DL-001",
        product_description="Connector",
        applicable_specification="GS-12-1507",
        confirmed_matrix_id="cmv-1",
        groups=(group1, group2),
        header_metadata=TestRecordHeaderMetadata(),
    )

    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert (
        "Group Number 组别编号: 1 ;   "
        "Sample Quantity & Number 样品数量及编号: 5 sets (Group1-1#~5#)"
    ) in text
    assert (
        "Group Number 组别编号: 2 ;   "
        "Sample Quantity & Number 样品数量及编号: 7 sets (Group2-1#~7#)"
    ) in text
    assert text.count("EQUIPMENT USED 使用的设备:") == 2
    assert document.element.xml.count('w:type="page"') >= 1
    assert len(document.tables) == 4
    assert document.tables[1].rows[1].cells[0].text == "EQUIPMENT USED 使用的设备:"
    assert document.tables[3].rows[1].cells[0].text == "EQUIPMENT USED 使用的设备:"


def test_gateway_orders_group_steps_by_step_token(tmp_path: Path) -> None:
    template = _build_confirmed_matrix_template(tmp_path / "template.docx")
    output = tmp_path / "confirmed-record.docx"
    group = _ConfirmedGroup(
        steps=(
            _ConfirmedStep(sequence=9, raw_token="9", test_item="Examination"),
            _ConfirmedStep(sequence=2, raw_token="2", test_item="LLCR"),
            _ConfirmedStep(sequence=4, raw_token="4(b)", test_item="High temp"),
        )
    )

    TestRecordDocumentGateway().generate_from_confirmed_matrix(
        template_path=template,
        output_path=output,
        project_id="P1",
        project_no="DL-001",
        product_description="Connector",
        applicable_specification="GS-12-1507",
        confirmed_matrix_id="cmv-1",
        groups=(group,),
        header_metadata=TestRecordHeaderMetadata(),
    )

    step_table = Document(output).tables[0]
    assert [row.cells[0].text for row in step_table.rows[1:]] == ["2", "4(b)", "9"]
    assert [row.cells[1].text for row in step_table.rows[1:]] == [
        "LLCR",
        "High temp",
        "Examination",
    ]


def test_gateway_writes_step_level_requirement_values_into_remark_column(tmp_path: Path) -> None:
    template = _build_confirmed_matrix_template(tmp_path / "template.docx")
    output = tmp_path / "confirmed-record.docx"
    group = _ConfirmedGroup(
        steps=(
            _ConfirmedStep(sequence=2, raw_token="2", test_item="LLCR"),
            _ConfirmedStep(sequence=5, raw_token="5", test_item="LLCR"),
        )
    )
    group.steps[0].requirement = "≤ 0.25 mΩ"
    group.steps[1].requirement = "ΔR ≤ 0.17 mΩ"

    TestRecordDocumentGateway().generate_from_confirmed_matrix(
        template_path=template,
        output_path=output,
        project_id="P1",
        project_no="DL-001",
        product_description="Connector",
        applicable_specification="GS-12-1507",
        confirmed_matrix_id="cmv-1",
        groups=(group,),
        header_metadata=TestRecordHeaderMetadata(),
    )

    step_table = Document(output).tables[0]
    assert step_table.rows[1].cells[8].text == "≤ 0.25 mΩ"
    assert step_table.rows[2].cells[8].text == "ΔR ≤ 0.17 mΩ"


def test_gateway_remark_only_keeps_numeric_requirement_content(tmp_path: Path) -> None:
    template = _build_confirmed_matrix_template(tmp_path / "template.docx")
    output = tmp_path / "confirmed-record.docx"
    group = _ConfirmedGroup(
        steps=(
            _ConfirmedStep(sequence=1, raw_token="1", test_item="Visual"),
            _ConfirmedStep(sequence=2, raw_token="2", test_item="LLCR"),
        )
    )
    group.steps[0].requirement = "No detrimental condition"
    group.steps[1].requirement = "Initial ≤ 25mΩ"

    TestRecordDocumentGateway().generate_from_confirmed_matrix(
        template_path=template,
        output_path=output,
        project_id="P1",
        project_no="DL-001",
        product_description="Connector",
        applicable_specification="GS-12-1507",
        confirmed_matrix_id="cmv-1",
        groups=(group,),
        header_metadata=TestRecordHeaderMetadata(),
    )

    step_table = Document(output).tables[0]
    assert step_table.rows[1].cells[8].text == ""
    assert step_table.rows[2].cells[8].text == "Initial ≤ 25mΩ"


def test_gateway_remark_drops_plain_numeric_reference_without_threshold(tmp_path: Path) -> None:
    template = _build_confirmed_matrix_template(tmp_path / "template.docx")
    output = tmp_path / "confirmed-record.docx"
    group = _ConfirmedGroup(
        steps=(
            _ConfirmedStep(sequence=1, raw_token="1", test_item="Reference"),
        )
    )
    group.steps[0].requirement = "See section 5.4 for details"

    TestRecordDocumentGateway().generate_from_confirmed_matrix(
        template_path=template,
        output_path=output,
        project_id="P1",
        project_no="DL-001",
        product_description="Connector",
        applicable_specification="GS-12-1507",
        confirmed_matrix_id="cmv-1",
        groups=(group,),
        header_metadata=TestRecordHeaderMetadata(),
    )

    step_table = Document(output).tables[0]
    assert step_table.rows[1].cells[8].text == ""


class _ConfirmedStep:
    def __init__(
        self,
        *,
        sequence: int = 1,
        raw_token: str = "1",
        test_item: str = "Visual",
    ) -> None:
        self.sequence = sequence
        self.raw_token = raw_token
        self.test_item = test_item
        self.section = "6.1"
        self.method = "EIA-364-18"
        self.condition = "10x"
        self.requirement = "No damage"


class _ConfirmedGroup:
    def __init__(
        self,
        *,
        group_key: str = "g1",
        group_label: str = "Group 1",
        sample_quantity_expression: str = "5",
        steps: tuple[_ConfirmedStep, ...] | None = None,
    ) -> None:
        self.group_key = group_key
        self.group_label = group_label
        self.sample_quantity_expression = sample_quantity_expression
        self.steps = steps or (_ConfirmedStep(),)
        self.step_count = len(self.steps)


def _build_confirmed_matrix_template(path: Path) -> Path:
    document = Document()
    paragraph = document.add_paragraph("Group Number 组别编号: ")
    paragraph.add_run("#")
    paragraph.add_run(" ;   Sample Quantity & Number 样品数量及编号: ")
    paragraph.add_run("#")
    step_table = document.add_table(rows=2, cols=9)
    step_table.rows[0].cells[0].text = "Step"
    step_table.rows[1].cells[0].text = "placeholder"
    document.add_paragraph("EQUIPMENT USED 使用的设备:")
    equipment_table = document.add_table(rows=2, cols=7)
    equipment_table.rows[0].cells[0].text = "Equipment ID No."
    equipment_table.rows[1].cells[0].text = "EQUIPMENT USED 使用的设备:"
    document.save(path)
    return path


def test_gateway_fills_test_record_header_metadata(tmp_path: Path) -> None:
    template = _build_template_with_header(tmp_path / "template.docx")
    output = tmp_path / "record.docx"

    TestRecordDocumentGateway().generate_from_confirmed_matrix(
        template_path=template,
        output_path=output,
        project_id="P1",
        project_no="DL-001",
        product_description="legacy",
        applicable_specification="legacy",
        confirmed_matrix_id="cmv-1",
        groups=(_ConfirmedGroup(),),
        header_metadata=TestRecordHeaderMetadata(
            lab_test_request_number="DL-2026-05-003",
            product_description="Coolpower HDF 3.40mm pin",
            applicable_specification="GS-12-1507",
        ),
    )

    document = Document(output)
    header_tables = document.sections[0].header.tables
    assert "Lab Test Request Number" in header_tables[0].cell(0, 2).text
    assert "DL-2026-05-003" in header_tables[0].cell(0, 2).text
    assert header_tables[1].cell(0, 1).text == "Coolpower HDF 3.40mm pin"
    assert header_tables[1].cell(0, 3).text == "GS-12-1507"
    assert header_tables[1].cell(0, 5).text == ""


def _build_template_with_header(path: Path) -> Path:
    section_width = 7 * 914400
    document = Document()
    header = document.sections[0].header
    table0 = header.add_table(rows=1, cols=3, width=section_width)
    table0.cell(0, 2).text = "Lab Test Request Number:\n实验室测试项目编号："
    table1 = header.add_table(rows=1, cols=6, width=section_width)
    table1.cell(0, 0).text = "Product Description\n产品描述"
    table1.cell(0, 2).text = "Applicable Specification\n适用的规范"
    table1.cell(0, 4).text = "Estimated Completion Date\n预计完成日期"
    document.add_paragraph("Group Number 组别编号: ")
    step_table = document.add_table(rows=1, cols=9)
    step_table.rows[0].cells[0].text = "Step"
    document.add_paragraph("EQUIPMENT USED 使用的设备:")
    equipment_table = document.add_table(rows=1, cols=7)
    equipment_table.rows[0].cells[0].text = "Equipment ID No."
    document.save(path)
    return path
