from __future__ import annotations

from copy import deepcopy
from dataclasses import replace
from datetime import date
from decimal import Decimal
from hashlib import sha256
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pytest

from backend.application.confirmed_matrix_test_record_preview_service import (
    ConfirmedMatrixTestRecordPreviewGroup,
    ConfirmedMatrixTestRecordPreviewStep,
)
from backend.application.test_report_draft_service import TestReportDraftData
from backend.application.project_basic_information_service import (
    ProjectBasicInformationSampleRow,
)
from backend.domain.result_dataset_models import (
    LlcrDatasetPayload,
    LlcrMeasurement,
    LlcrResultEntry,
    ResultDatasetRevision,
    ResultDatasetSourceIdentity,
)
from backend.infrastructure.office.test_report_document_gateway import (
    TestReportDocumentGateway,
)


def test_generates_e3707_draft_without_mutating_approved_template(tmp_path: Path) -> None:
    template = _build_template(tmp_path / "E-3707_H.docx")
    template_hash = sha256(template.read_bytes()).hexdigest()
    output = tmp_path / "draft.docx"
    output.touch()

    written = TestReportDocumentGateway().generate(
        template_path=template,
        output_path=output,
        report=_report(),
    )

    assert written == output
    assert sha256(template.read_bytes()).hexdigest() == template_hash
    document = Document(output)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    assert any(
        "qualification testing conducted on Coolpower HDF" in text
        and "GS-12-2113 Rev.7" in text
        for text in paragraphs
    )
    assert any("DRAFT" in text and "testing and review" in text for text in paragraphs)
    assert any("May 20, 2026" in text for text in paragraphs)
    assert all("[PRODUCT NAME]" not in text for text in paragraphs)

    headers = [
        text_node.text or ""
        for section in document.sections
        for header in (section.header, section.first_page_header)
        for text_node in header._element.xpath(".//w:t")
    ]
    assert any("DL-2026-05-011" in text for text in headers)
    assert all("WW-XXXX-YY-ZZZ" not in text for text in headers)
    assert any("MP Cao" in text for text in headers)
    assert any("Even Yang" in text for text in headers)
    assert any("Coolpower HDF qualification testing" in text for text in headers)
    assert any("05/Dec/2025 to 03/Mar/2026" in text for text in headers)
    assert all("TBD" not in text for text in headers)

    sample = document.tables[0]
    assert len(sample.rows) == 3
    assert [cell.text for cell in sample.rows[1].cells] == [
        "Pin",
        "PN-PIN",
        "202510",
        "C1100",
        "Ag",
        "No",
        "NA",
    ]
    assert [cell.text for cell in sample.rows[2].cells] == [
        "Socket",
        "PN-SOCKET",
        "202511",
        "C19010",
        "Au",
        "Yes",
        "LCP",
    ]

    description = document.tables[1]
    assert [cell.text for cell in description.rows[1].cells] == [
        "Test Items",
        "1",
        "2",
    ]
    assert _cell_fill(description.cell(1, 1)) == "B2B2B2"
    assert _cell_fill(description.cell(1, 2)) == "B2B2B2"
    assert [cell.text for cell in description.rows[-1].cells] == [
        "Samples Size(sets)",
        "5",
        "3",
    ]
    assert all(_cell_fill(cell) == "8DB3E2" for cell in description.rows[-1].cells)
    assert description.cell(2, 2)._tc.tcPr.find(qn("w:tcBorders")) is not None

    all_tables = list(document.tables)
    for section in document.sections:
        for header in (section.header, section.first_page_header):
            all_tables.extend(header.tables)
        for footer in (section.footer, section.first_page_footer):
            all_tables.extend(footer.tables)
    populated_table_runs = [
        run
        for table in all_tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text
    ]
    assert populated_table_runs
    assert all(run.font.name == "Arial" for run in populated_table_runs)

    test_description_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "4. TEST DESCRIPTION"
    )
    assert test_description_heading.paragraph_format.keep_with_next is True

    methods = document.tables[2]
    method_rows = [[cell.text for cell in row.cells] for row in methods.rows]
    assert [
        "Visual Examination",
        "EIA-364-18",
        "10x",
        "No detrimental condition",
    ] in method_rows
    assert ["LLCR", "EIA-364-23", "20mV max", "Initial ≤0.25mΩ; ΔR ≤0.17mΩ"] in method_rows

    result_tables = [
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells]
        == ["Step", "Test", "Requirement", "Step Description", "Result", "Comment"]
    ]
    assert len(result_tables) == 2
    assert [cell.text for cell in result_tables[0].rows[1].cells] == [
        "1",
        "Visual Examination",
        "No detrimental condition",
        "Visual Examination",
        "No detriment",
        "Pass",
    ]
    assert [cell.text for cell in result_tables[0].rows[2].cells][-2:] == [
        "Initial ≤_mΩ",
        "Pass",
    ]
    assert [cell.text for cell in result_tables[1].rows[2].cells][-2:] == [
        "ΔR ≤_mΩ",
        "Pass",
    ]
    assert "Group 1 Test Results" in paragraphs
    assert "Group 2 Test Results" in paragraphs

    revision = document.tables[-1]
    assert [cell.text for cell in revision.rows[1].cells] == [
        "A",
        "All",
        "Initial draft - not released",
        "28/Aug/2026",
    ]
    equipment_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "7. EQUIPMENTS"
    )
    assert equipment_heading._p.getprevious().xpath('.//w:br[@w:type="page"]')


def test_synchronizes_only_managed_llcr_result_cells_into_a_new_draft(tmp_path: Path) -> None:
    template = _build_template(tmp_path / "E-3707_H.docx")
    source = tmp_path / "revision-1.docx"
    gateway = TestReportDocumentGateway()
    gateway.generate(template_path=template, output_path=source, report=_report())
    document = Document(source)
    for paragraph in document.paragraphs:
        if "This report summarizes" in paragraph.text:
            paragraph.text = "Operator-maintained purpose text."
            break
    first_results = _result_tables(document)[0]
    first_results.cell(1, 4).text = "Operator visual result"
    first_results.cell(1, 5).text = "Reviewed"
    document.save(source)
    source_hash = sha256(source.read_bytes()).hexdigest()
    output = tmp_path / "revision-2.docx"

    written = gateway.synchronize_llcr_results(
        source_path=source,
        output_path=output,
        dataset=_llcr_dataset(),
    )

    assert written == output
    assert sha256(source.read_bytes()).hexdigest() == source_hash
    synchronized = Document(output)
    assert any(
        paragraph.text == "Operator-maintained purpose text."
        for paragraph in synchronized.paragraphs
    )
    result_table = _result_tables(synchronized)[0]
    assert result_table.cell(1, 4).text == "Operator visual result"
    assert result_table.cell(1, 5).text == "Reviewed"
    assert result_table.cell(2, 4).text == "Initial ≤0.198mΩ"
    assert result_table.cell(2, 5).text == "Pass"


def test_llcr_sync_fails_without_partial_output_when_target_is_ambiguous(tmp_path: Path) -> None:
    template = _build_template(tmp_path / "E-3707_H.docx")
    source = tmp_path / "revision-1.docx"
    gateway = TestReportDocumentGateway()
    gateway.generate(template_path=template, output_path=source, report=_report())
    document = Document(source)
    first_results = _result_tables(document)[0]
    first_results._tbl.append(deepcopy(first_results.rows[2]._tr))
    document.save(source)
    output = tmp_path / "revision-2.docx"

    with pytest.raises(ValueError, match="uniquely locate"):
        gateway.synchronize_llcr_results(
            source_path=source,
            output_path=output,
            dataset=_llcr_dataset(),
        )

    assert not output.exists()


def test_generates_from_approved_numbered_heading_without_separator_space(
    tmp_path: Path,
) -> None:
    template = _build_template(
        tmp_path / "E-3707_H.docx",
        compact_numbered_headings=True,
    )
    output = tmp_path / "draft.docx"

    TestReportDocumentGateway().generate(
        template_path=template,
        output_path=output,
        report=_report(),
    )

    document = Document(output)
    assert "1.PURPOSE" in [paragraph.text for paragraph in document.paragraphs]
    test_description_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "4.TEST DESCRIPTION"
    )
    assert test_description_heading.paragraph_format.keep_with_next is True


def test_generates_when_first_page_header_placeholders_are_split_across_runs(
    tmp_path: Path,
) -> None:
    template = _build_template(
        tmp_path / "E-3707_H.docx",
        split_first_page_header_placeholders=True,
    )
    output = tmp_path / "draft.docx"

    TestReportDocumentGateway().generate(
        template_path=template,
        output_path=output,
        report=_report(),
    )

    document = Document(output)
    first_page_text = "".join(
        text_node.text or ""
        for text_node in document.sections[0].first_page_header._element.xpath(
            ".//w:t"
        )
    )
    assert "DL-2026-05-011" in first_page_text
    assert "WW-XXXX-YY-ZZZ" not in first_page_text
    assert "28/Aug/2026" in first_page_text
    assert "05/Dec/2025 to 03/Mar/2026" in first_page_text
    assert "TBD" not in first_page_text


def test_populates_default_approved_by_name_from_template(tmp_path: Path) -> None:
    template = _build_template(tmp_path / "E-3707_H.docx")
    template_document = Document(template)
    approved_by_cell = template_document.sections[0].first_page_header.tables[0].cell(
        1,
        3,
    )
    approved_by_cell.text = "Name"
    template_document.save(template)
    output = tmp_path / "draft.docx"

    TestReportDocumentGateway().generate(
        template_path=template,
        output_path=output,
        report=_report(),
    )

    generated = Document(output)
    assert (
        generated.sections[0].first_page_header.tables[0].cell(1, 3).text
        == "Gentle Zeng"
    )


def test_populates_legacy_purpose_and_received_date_placeholders(
    tmp_path: Path,
) -> None:
    template = _build_template(
        tmp_path / "E-3707_H.docx",
        legacy_body_placeholders=True,
    )
    output = tmp_path / "draft.docx"

    TestReportDocumentGateway().generate(
        template_path=template,
        output_path=output,
        report=_report(),
    )

    paragraphs = [paragraph.text for paragraph in Document(output).paragraphs]
    assert any(
        text
        == (
            "This report summarizes the qualification testing conducted on Coolpower HDF "
            "to assess the conformance to AFCI product specification GS-12-2113 Rev.7."
        )
        for text in paragraphs
    )
    assert any("May 20, 2026" in text for text in paragraphs)
    assert all("XXXXXXXXXXXXXXXX" not in text for text in paragraphs)
    assert all("DDMMMYYYY" not in text for text in paragraphs)


def test_generates_result_headings_from_numeric_confirmed_group_labels(
    tmp_path: Path,
) -> None:
    template = _build_template(tmp_path / "E-3707_H.docx")
    output = tmp_path / "draft.docx"
    report = _report()
    numeric_groups = tuple(
        replace(group, group_label=str(index))
        for index, group in enumerate(report.groups, start=1)
    )

    TestReportDocumentGateway().generate(
        template_path=template,
        output_path=output,
        report=replace(report, groups=numeric_groups),
    )

    paragraphs = [paragraph.text for paragraph in Document(output).paragraphs]
    assert "Group 1 Test Results" in paragraphs
    assert "Group 2 Test Results" in paragraphs


def test_result_groups_flow_without_forced_page_breaks(tmp_path: Path) -> None:
    template = _build_template(tmp_path / "E-3707_H.docx")
    output = tmp_path / "draft.docx"

    TestReportDocumentGateway().generate(
        template_path=template,
        output_path=output,
        report=_report(),
    )

    document = Document(output)
    group_headings = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Group ")
        and paragraph.text.endswith(" Test Results")
    ]

    assert [paragraph.text for paragraph in group_headings] == [
        "Group 1 Test Results",
        "Group 2 Test Results",
    ]
    assert all(
        paragraph.paragraph_format.keep_with_next is True
        for paragraph in group_headings
    )
    assert all(
        not paragraph._p.getprevious().xpath('.//w:br[@w:type="page"]')
        for paragraph in group_headings[1:]
    )


def test_large_matrix_description_uses_consistent_fixed_geometry(
    tmp_path: Path,
) -> None:
    template = _build_template(tmp_path / "E-3707_H.docx")
    template_document = Document(template)
    _set_fixture_table_geometry(template_document.tables[1], (3024, 7857))
    template_document.save(template)
    output = tmp_path / "draft.docx"
    report = _report()
    groups = tuple(
        replace(
            report.groups[0],
            group_key=f"g{index}",
            group_label=f"Group {index}",
        )
        for index in range(1, 13)
    )

    TestReportDocumentGateway().generate(
        template_path=template,
        output_path=output,
        report=replace(report, groups=groups),
    )

    description = Document(output).tables[1]
    grid_widths = [
        int(column.get(qn("w:w")))
        for column in description._tbl.tblGrid.gridCol_lst
    ]
    table_width = description._tbl.tblPr.find(qn("w:tblW"))
    table_layout = description._tbl.tblPr.find(qn("w:tblLayout"))

    assert [cell.text for cell in description.rows[1].cells][-3:] == [
        "10",
        "11",
        "12",
    ]
    assert len(grid_widths) == 13
    assert table_width is not None
    assert table_width.get(qn("w:type")) == "dxa"
    assert int(table_width.get(qn("w:w"))) == sum(grid_widths) == 10881
    assert table_layout is not None
    assert table_layout.get(qn("w:type")) == "fixed"
    assert 2500 <= grid_widths[0] <= 3024
    assert min(grid_widths[1:]) >= 690
    assert max(grid_widths[1:]) - min(grid_widths[1:]) <= 1

    for row in description.rows:
        grid_index = 0
        for cell_xml in row._tr.tc_lst:
            span_xml = cell_xml.tcPr.find(qn("w:gridSpan"))
            span = 1 if span_xml is None else int(span_xml.get(qn("w:val")))
            cell_width = int(cell_xml.tcPr.tcW.get(qn("w:w")))
            assert cell_width == sum(grid_widths[grid_index : grid_index + span])
            grid_index += span
        assert grid_index == len(grid_widths)


def test_rejects_template_contract_drift_without_replacing_reserved_output(
    tmp_path: Path,
) -> None:
    template = _build_template(tmp_path / "E-3707_H.docx")
    document = Document(template)
    document.tables[2].cell(0, 0).text = "Changed heading"
    document.save(template)
    output = tmp_path / "draft.docx"
    output.touch()

    with pytest.raises(ValueError, match="Test Methods/Requirements table"):
        TestReportDocumentGateway().generate(
            template_path=template,
            output_path=output,
            report=_report(),
        )

    assert output.read_bytes() == b""


def test_generates_llcr_stage_descriptions_and_stage_specific_result_shells(
    tmp_path: Path,
) -> None:
    template = _build_template(tmp_path / "E-3707_H.docx")
    output = tmp_path / "draft.docx"
    group = ConfirmedMatrixTestRecordPreviewGroup(
        group_key="g1",
        group_label="Group 1",
        sample_quantity_expression="5",
        step_count=5,
        steps=(
            _step(1, "Contact Resistance (Low Level)", "≤0.25mΩ"),
            _step(2, "Pe-Durability", "No damage"),
            _step(3, "Contact Resistance (Low Level)", "ΔR ≤0.17mΩ"),
            _step(4, "Thermal Shock", "No damage"),
            _step(5, "Contact Resistance (Low Level)", "ΔR ≤0.17mΩ"),
        ),
    )

    TestReportDocumentGateway().generate(
        template_path=template,
        output_path=output,
        report=replace(_report(), groups=(group,)),
    )

    document = Document(output)
    results = next(
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells]
        == ["Step", "Test", "Requirement", "Step Description", "Result", "Comment"]
    )
    assert [row.cells[3].text for row in results.rows[1:]] == [
        "LLCR",
        "Pe-Durability",
        "After Pe-Durability",
        "Thermal Shock",
        "Final ΔR",
    ]
    assert results.rows[1].cells[2].text == "Initial ≤0.25mΩ"
    assert [row.cells[4].text for row in results.rows[1:]] == [
        "Initial ≤_mΩ",
        "No damage",
        "ΔR ≤_mΩ",
        "No damage",
        "ΔR ≤_mΩ",
    ]
    assert all(row.cells[5].text == "Pass" for row in results.rows[1:])


def test_rejects_first_page_header_contract_drift(tmp_path: Path) -> None:
    template = _build_template(tmp_path / "E-3707_H.docx")
    document = Document(template)
    first_page = document.sections[0].first_page_header
    first_page.tables[0].cell(0, 0).text = "REPORT NUMBER"
    document.save(template)
    output = tmp_path / "draft.docx"
    output.touch()

    with pytest.raises(ValueError, match="first-page header placeholder"):
        TestReportDocumentGateway().generate(
            template_path=template,
            output_path=output,
            report=_report(),
        )

    assert output.read_bytes() == b""


def test_refuses_to_replace_nonempty_output(tmp_path: Path) -> None:
    template = _build_template(tmp_path / "E-3707_H.docx")
    output = tmp_path / "manual.docx"
    output.write_bytes(b"manual")

    with pytest.raises(FileExistsError, match="already exists"):
        TestReportDocumentGateway().generate(
            template_path=template,
            output_path=output,
            report=_report(),
        )

    assert output.read_bytes() == b"manual"


def _build_template(
    path: Path,
    *,
    compact_numbered_headings: bool = False,
    split_first_page_header_placeholders: bool = False,
    legacy_body_placeholders: bool = False,
) -> Path:
    document = Document()
    header = document.sections[0].header
    header.add_table(rows=1, cols=2, width=7 * 914400).cell(0, 0).text = "XX-YY-ZZZ"
    document.sections[0].different_first_page_header_footer = True
    first_page = document.sections[0].first_page_header.add_table(
        rows=2,
        cols=4,
        width=7 * 914400,
    )
    for cell, value in zip(
        first_page.rows[0].cells,
        ["WW-XXXX-YY-ZZZ", "DDMMMYYYY", "DDMMMYYYY-DDMMMYYYY", "Name(s)"],
        strict=True,
    ):
        cell.text = value
    for cell, value in zip(
        first_page.rows[1].cells,
        ["NAME", "PRODUCT NAME/TEST DESCRIPTION ", "Name", "(s)"],
        strict=True,
    ):
        cell.text = value
    if split_first_page_header_placeholders:
        _set_cell_runs(first_page.rows[0].cells[0], "WW-XX", "XX-YY-ZZZ")
        _set_cell_runs(first_page.rows[0].cells[1], "DD", "MMM", "YYYY")
        _set_cell_runs(
            first_page.rows[0].cells[2],
            "DD",
            "MMM",
            "YYYY-DD",
            "MMM",
            "YYYY",
        )

    heading = (
        (lambda value: value.replace(". ", "."))
        if compact_numbered_headings
        else (lambda value: value)
    )

    document.add_paragraph(heading("1. PURPOSE"))
    if legacy_body_placeholders:
        document.add_paragraph(
            "This report summarizes theXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX assess "
            "the conformance to product specification GS-XX-XXXX(Rev.)"
        )
    else:
        document.add_paragraph(
            "This report summarizes the [TEST DESCRIPTION] conducted on [PRODUCT NAME] "
            "to assess the conformance to AFCI product specification "
            "[GS-XX-XXXX (Rev.X, DATE)]."
        )
    document.add_paragraph(heading("2. CONCLUSIONS"))
    document.add_paragraph(
        "The [PRODUCT NAME] successfully completed the tested items and met the specified "
        "requirements of AFCI product specification [GS-XX-XXXX (Rev.X, DATE)]."
    )
    document.add_paragraph(heading("3. SAMPLE DESCRIPTION"))
    received_date_placeholder = (
        "DDMMMYYYY" if legacy_body_placeholders else "[RECEIVED SAMPLES DATE]"
    )
    document.add_paragraph(
        f"Samples were received at the laboratory on {received_date_placeholder}. Prior to "
        "testing, the samples were examined at low magnification and judged to be acceptable "
        "for testing."
    )
    sample = document.add_table(rows=2, cols=7)
    for cell, value in zip(
        sample.rows[0].cells,
        [
            "Description",
            "Part #",
            "Lot#",
            "Base Mat’l",
            "Contact Plating",
            "Contact Lubricant",
            "Housing Material",
        ],
        strict=True,
    ):
        cell.text = value
    document.add_paragraph("Figure 1:")
    document.add_paragraph(heading("4. TEST DESCRIPTION"))
    description = document.add_table(rows=3, cols=2)
    description.cell(0, 0).text = "Test Items"
    description.cell(0, 1).text = "Test Sequence"
    for row in description.rows:
        _add_direct_cell_borders(row.cells[1])
    document.add_paragraph(heading("5. TEST METHODS/REQUIREMENTS"))
    methods = document.add_table(rows=2, cols=4)
    for cell, value in zip(
        methods.rows[0].cells,
        ["Item", "Test Method", "Condition", "Requirement"],
        strict=True,
    ):
        cell.text = value
    document.add_paragraph(heading("6. TEST RESULTS"))
    document.add_paragraph(
        "Unless otherwise specified, assessment of conformity to requirements is based on "
        "simple acceptance."
    )
    document.add_paragraph("Group # Test Results")
    results = document.add_table(rows=3, cols=6)
    for cell, value in zip(
        results.rows[0].cells,
        ["Step", "Test", "Requirement", "Step Description", "Result", "Comment"],
        strict=True,
    ):
        cell.text = value
    document.add_paragraph(heading("7. EQUIPMENTS"))
    equipment = document.add_table(rows=2, cols=5)
    for cell, value in zip(
        equipment.rows[0].cells,
        ["Item", "Manufacturer", "ID Number", "Last Cal.", "Cal. Due"],
        strict=True,
    ):
        cell.text = value
    document.add_paragraph(heading("8. REVISION RECORD"))
    revision = document.add_table(rows=4, cols=4)
    for cell, value in zip(
        revision.rows[0].cells,
        ["Revision Level", "Affected Pages", "Description", "Revision Date"],
        strict=True,
    ):
        cell.text = value
    revision.rows[1].cells[0].text = "A"
    revision.rows[1].cells[1].text = "All"
    revision.rows[1].cells[2].text = "Original Release"
    document.add_paragraph("*** End of Report ***")

    second = document.add_section(WD_SECTION.NEW_PAGE)
    second.header.is_linked_to_previous = False
    second.header.add_table(rows=1, cols=2, width=7 * 914400).cell(0, 0).text = (
        "Report No. WW-XXXX-YY-ZZZ"
    )
    document.save(path)
    return path


def _report() -> TestReportDraftData:
    group1 = ConfirmedMatrixTestRecordPreviewGroup(
        group_key="g1",
        group_label="Group 1",
        sample_quantity_expression="5",
        step_count=2,
        steps=(
            ConfirmedMatrixTestRecordPreviewStep(
                sequence=1,
                raw_token="1",
                test_item="Visual Examination",
                section="6.1",
                method="EIA-364-18",
                condition="10x",
                requirement="No detrimental condition",
            ),
            ConfirmedMatrixTestRecordPreviewStep(
                sequence=2,
                raw_token="2",
                test_item="LLCR",
                section="6.2",
                method="EIA-364-23",
                condition="20mV max",
                requirement="Initial ≤0.25mΩ",
            ),
        ),
    )
    group2 = ConfirmedMatrixTestRecordPreviewGroup(
        group_key="g2",
        group_label="Group 2",
        sample_quantity_expression="3",
        step_count=2,
        steps=(
            ConfirmedMatrixTestRecordPreviewStep(
                sequence=1,
                raw_token="1",
                test_item="Visual Examination",
                section="6.1",
                method="EIA-364-18",
                condition="10x",
                requirement="No detrimental condition",
            ),
            ConfirmedMatrixTestRecordPreviewStep(
                sequence=4,
                raw_token="4",
                test_item="LLCR",
                section="6.2",
                method="EIA-364-23",
                condition="20mV max",
                requirement="ΔR ≤0.17mΩ",
            ),
        ),
    )
    return TestReportDraftData(
        project_id="P1",
        report_number="DL-2026-05-011",
        product_name="Coolpower HDF",
        test_description="qualification testing",
        applicable_specification="GS-12-2113 Rev.7",
        received_samples_date="2026-05-20",
        start_test_date="2025-12-05",
        finish_test_date="2026-03-03",
        description_part_number="10179696-0001LF",
        requestor="MP Cao",
        project_leader="Even Yang",
        confirmed_matrix_id="cmv-1",
        groups=(group1, group2),
        generated_on=date(2026, 8, 28),
        sample_rows=(
            ProjectBasicInformationSampleRow(
                product_name="Pin",
                part_number="PN-PIN",
                lot_or_traceability="202510",
                material="C1100",
                plating="Ag",
                lubricant="No",
                housing_material="NA",
                row_index=0,
            ),
            ProjectBasicInformationSampleRow(
                product_name="Socket",
                part_number="PN-SOCKET",
                lot_or_traceability="202511",
                material="C19010",
                plating="Au",
                lubricant="Yes",
                housing_material="LCP",
                row_index=1,
            ),
        ),
    )


def _step(
    sequence: int,
    test_item: str,
    requirement: str,
) -> ConfirmedMatrixTestRecordPreviewStep:
    return ConfirmedMatrixTestRecordPreviewStep(
        sequence=sequence,
        raw_token=str(sequence),
        test_item=test_item,
        section="6.1",
        method="EIA-364",
        condition="Condition",
        requirement=requirement,
    )


def _llcr_dataset() -> ResultDatasetRevision:
    measurement = LlcrMeasurement(
        1, "SIG1", Decimal("0.198"), "mΩ", "SIG", "K10",
        Decimal("0.248"), "mΩ", "D10",
    )
    entry = LlcrResultEntry(
        result_id="group-1:row-llcr:2",
        confirmed_group_id="group-1",
        group_label="1",
        confirmed_row_id="row-llcr",
        matrix_step_sequence=2,
        matrix_step_token="2",
        stage="initial",
        stage_label="Initial LLCR",
        requirement="Initial ≤0.25mΩ",
        requirement_comparator="<=",
        requirement_limit=Decimal("0.25"),
        requirement_unit="mΩ",
        measurements=(measurement,),
        summary_min=Decimal("0.198"),
        summary_max=Decimal("0.198"),
        summary_average=Decimal("0.198"),
        provisional_outcome="pass",
        confirmed_outcome="pass",
        source_range="SIG!K10:K10",
    )
    return ResultDatasetRevision(
        dataset_id="dataset-1",
        dataset_type="llcr",
        revision=1,
        project_id="P1",
        confirmed_matrix_id="cmv-1",
        confirmed_matrix_revision=1,
        source=ResultDatasetSourceIdentity("LLCR.xlsx", "a" * 64, 100),
        imported_at="2026-08-29T08:00:00Z",
        imported_by="Even Yang",
        confirmed_at="2026-08-29T08:01:00Z",
        confirmed_by="Even Yang",
        parser_profile_version="connlab-llcr-macro-v1",
        validation_status="confirmed",
        payload=LlcrDatasetPayload((entry,)),
    )


def _result_tables(document):
    headers = ["Step", "Test", "Requirement", "Step Description", "Result", "Comment"]
    return [
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells] == headers
    ]


def _add_direct_cell_borders(cell) -> None:
    borders = OxmlElement("w:tcBorders")
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        borders.append(edge)
    cell._tc.get_or_add_tcPr().append(borders)


def _cell_fill(cell) -> str | None:
    shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return shading.get(qn("w:fill")) if shading is not None else None


def _set_cell_runs(cell, *values: str) -> None:
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    for value in values:
        paragraph.add_run(value)


def _set_fixture_table_geometry(table, widths: tuple[int, ...]) -> None:
    table.autofit = False
    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    assert table_width is not None
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(widths)))
    for column, width in zip(
        table._tbl.tblGrid.gridCol_lst,
        widths,
        strict=True,
    ):
        column.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell._tc.tcPr.tcW.set(qn("w:type"), "dxa")
            cell._tc.tcPr.tcW.set(qn("w:w"), str(width))
