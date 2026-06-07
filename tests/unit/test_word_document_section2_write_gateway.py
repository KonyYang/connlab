from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from backend.infrastructure.office import WordDocumentGateway


def test_word_gateway_writes_section2_fields(tmp_path: Path) -> None:
    path = _section2_docx(tmp_path / "request.docx")

    result = WordDocumentGateway().write_section2_fields(
        path,
        {
            "lab": "Connector Lab",
            "assigned_personnel": "White",
            "received_date": "2026-05-12",
            "estimated_completion_date": "2026-05-19",
            "sample_condition": "Good condition",
        },
    )

    values = _table_values(path)
    assert values["Lab"] == "Connector Lab"
    assert values["Assigned Personnel"] == "White"
    assert values["Received Date"] == "2026-05-12"
    assert values["Estimated Completion Date"] == "2026-05-19"
    assert values["Sample Condition"] == "Good condition"
    assert {item.field_key for item in result.changed_fields} == {
        "lab",
        "assigned_personnel",
        "received_date",
        "estimated_completion_date",
        "sample_condition",
    }
    assert result.unchanged_fields == ()


def test_word_gateway_matches_lab_performing_the_tests_label(tmp_path: Path) -> None:
    """Real Section 2 label text maps to the existing lab field."""
    path = _section2_docx(
        tmp_path / "request.docx",
        lab_label="Lab Performing the Tests:",
    )

    result = WordDocumentGateway().write_section2_fields(path, {"lab": "Dongguan"})

    values = _table_values(path)
    assert values["Lab Performing the Tests:"] == "Dongguan"
    assert result.changed_fields[0].field_key == "lab"


def test_word_gateway_rejects_missing_section2_location_without_save(
    tmp_path: Path,
) -> None:
    path = _section2_docx(tmp_path / "request.docx", include_sample_condition=False)

    with pytest.raises(ValueError, match="sample_condition"):
        WordDocumentGateway().write_section2_fields(
            path,
            {
                "lab": "Connector Lab",
                "sample_condition": "Good condition",
            },
        )

    values = _table_values(path)
    assert values["Lab"] == ""
    assert "Sample Condition" not in values


def test_word_gateway_rejects_non_docx(tmp_path: Path) -> None:
    path = tmp_path / "request.doc"
    path.write_text("legacy", encoding="utf-8")

    with pytest.raises(ValueError, match="Only .docx"):
        WordDocumentGateway().write_section2_fields(path, {"lab": "Connector Lab"})


def _section2_docx(
    path: Path,
    *,
    include_sample_condition: bool = True,
    lab_label: str = "Lab",
) -> Path:
    document = Document()
    labels = [
        lab_label,
        "Assigned Personnel",
        "Received Date",
        "Estimated Completion Date",
    ]
    if include_sample_condition:
        labels.append("Sample Condition")
    table = document.add_table(rows=len(labels), cols=2)
    for index, label in enumerate(labels):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = ""
    document.save(path)
    return path


def _table_values(path: Path) -> dict[str, str]:
    document = Document(path)
    table = document.tables[0]
    return {
        row.cells[0].text.strip(): row.cells[1].text.strip()
        for row in table.rows
    }
