from __future__ import annotations

from pathlib import Path
from datetime import date
import json

import pytest

from backend.application.confirmed_matrix_test_record_document_generation_service import (
    ConfirmedMatrixTestRecordDocumentGenerationError,
    ConfirmedMatrixTestRecordDocumentGenerationNotFoundError,
    ConfirmedMatrixTestRecordDocumentGenerationService,
    GenerateConfirmedMatrixTestRecordDocumentCommand,
)
from backend.application.project_basic_information_output import (
    ConfirmedBasicInformationSnapshot,
)
from backend.application.confirmed_matrix_test_record_preview_service import (
    ConfirmedMatrixTestRecordPreview,
    ConfirmedMatrixTestRecordPreviewGroup,
    ConfirmedMatrixTestRecordPreviewNotFoundError,
    ConfirmedMatrixTestRecordPreviewStep,
    ConfirmedMatrixTestRecordStepQuantity,
)
from backend.domain.enums import LtrStatus


def test_generation_service_writes_preview_groups_to_controlled_output(tmp_path: Path) -> None:
    writer = _Writer()
    service = ConfirmedMatrixTestRecordDocumentGenerationService(
        preview_service=_PreviewService(_preview()),
        project_store=_ProjectStore(),
        writer=writer,
        ltr_store=_LtrStore(),
        intake_case_store=_IntakeCaseStore(),
        intake_draft_store=_IntakeDraftStore(),
        application_form_store=_ApplicationFormStore(),
    )

    result = service.generate(
        GenerateConfirmedMatrixTestRecordDocumentCommand(
            project_id="P1",
            output_dir=tmp_path,
            template_path=_template(tmp_path),
        )
    )

    assert result.project_id == "P1"
    assert result.confirmed_matrix_id == "cmv-1"
    assert result.output_path.name == "DL-001 Test Record.docx"
    assert writer.calls[0]["product_description"] == "Coolpower HDF 3.40mm pin"
    header_metadata = writer.calls[0]["header_metadata"]
    assert header_metadata.lab_test_request_number == "DL-2026-05-003"
    assert header_metadata.product_description == "Coolpower HDF 3.40mm pin"
    assert header_metadata.applicable_specification == "GS-12-1507"
    assert writer.calls[0]["groups"][0].group_label == "Group 1"
    assert writer.calls[0]["groups"][0].sample_quantity_expression == "5"
    assert writer.calls[0]["groups"][0].steps[0].raw_token == "1"


def test_generation_service_requires_confirmed_basic_information_when_reader_is_configured(
    tmp_path: Path,
) -> None:
    writer = _Writer()
    service = ConfirmedMatrixTestRecordDocumentGenerationService(
        preview_service=_PreviewService(_preview()),
        project_store=_ProjectStore(),
        writer=writer,
        basic_information_reader=_BasicInformationReader(None),
    )

    with pytest.raises(
        ConfirmedMatrixTestRecordDocumentGenerationError,
        match="Confirm Basic Information before generating Test Record",
    ):
        service.generate(
            GenerateConfirmedMatrixTestRecordDocumentCommand(
                project_id="P1",
                output_dir=tmp_path,
                template_path=_template(tmp_path),
            )
        )

    assert writer.calls == []


def test_generation_service_uses_confirmed_basic_information_for_header_metadata(
    tmp_path: Path,
) -> None:
    writer = _Writer()
    service = ConfirmedMatrixTestRecordDocumentGenerationService(
        preview_service=_PreviewService(_preview()),
        project_store=_ProjectStore(),
        writer=writer,
        basic_information_reader=_BasicInformationReader(
            _basic_information(
                {
                    "dl_number": "DL-2026-05-011",
                    "product_description": "Confirmed Coolpower HDF",
                    "description_pn": "SHOULD-NOT-BE-USED",
                    "applicable_specifications": "GS-12-9999",
                },
                version=3,
            )
        ),
        ltr_store=_LtrStore(),
        intake_case_store=_IntakeCaseStore(),
        intake_draft_store=_IntakeDraftStore(),
    )

    result = service.generate(
        GenerateConfirmedMatrixTestRecordDocumentCommand(
            project_id="P1",
            output_dir=tmp_path,
            template_path=_template(tmp_path),
        )
    )

    header_metadata = writer.calls[0]["header_metadata"]
    assert header_metadata.lab_test_request_number == "DL-2026-05-011"
    assert header_metadata.product_description == "Confirmed Coolpower HDF"
    assert header_metadata.applicable_specification == "GS-12-9999"
    assert writer.calls[0]["product_description"] == "Confirmed Coolpower HDF"
    assert result.confirmed_basic_information_version == 3
    assert result.confirmed_basic_information_source_signature_hash == (
        _basic_information({}, version=3).source_signature_hash
    )


def test_generation_service_uses_unique_draft_name_when_target_exists(
    tmp_path: Path,
) -> None:
    writer = _Writer()
    (tmp_path / "DL-001 Test Record.docx").write_bytes(b"user edited")
    service = ConfirmedMatrixTestRecordDocumentGenerationService(
        preview_service=_PreviewService(_preview()),
        project_store=_ProjectStore(),
        writer=writer,
    )

    result = service.generate(
        GenerateConfirmedMatrixTestRecordDocumentCommand(
            project_id="P1",
            output_dir=tmp_path,
            template_path=_template(tmp_path),
        )
    )

    assert result.output_path.name == "DL-001 Test Record (2).docx"
    assert (tmp_path / "DL-001 Test Record.docx").read_bytes() == b"user edited"


def test_generation_service_uses_active_confirmed_preview_only(tmp_path: Path) -> None:
    service = ConfirmedMatrixTestRecordDocumentGenerationService(
        preview_service=_PreviewService(None),
        project_store=_ProjectStore(),
        writer=_Writer(),
    )

    with pytest.raises(ConfirmedMatrixTestRecordDocumentGenerationNotFoundError):
        service.generate(
            GenerateConfirmedMatrixTestRecordDocumentCommand(
                project_id="P1",
                output_dir=tmp_path,
                template_path=_template(tmp_path),
            )
        )


def test_generation_service_rejects_empty_preview(tmp_path: Path) -> None:
    empty = ConfirmedMatrixTestRecordPreview(
        project_id="P1",
        confirmed_matrix_id="cmv-empty",
        preview_status="empty",
        groups=(),
    )
    service = ConfirmedMatrixTestRecordDocumentGenerationService(
        preview_service=_PreviewService(empty),
        project_store=_ProjectStore(),
        writer=_Writer(),
    )

    with pytest.raises(ConfirmedMatrixTestRecordDocumentGenerationError, match="no previewable"):
        service.generate(
            GenerateConfirmedMatrixTestRecordDocumentCommand(
                project_id="P1",
                output_dir=tmp_path,
                template_path=_template(tmp_path),
            )
        )


def test_generation_service_converts_template_structure_error_to_business_error(
    tmp_path: Path,
) -> None:
    service = ConfirmedMatrixTestRecordDocumentGenerationService(
        preview_service=_PreviewService(_preview()),
        project_store=_ProjectStore(),
        writer=_FailingWriter(),
    )

    with pytest.raises(ConfirmedMatrixTestRecordDocumentGenerationError, match="step table"):
        service.generate(
            GenerateConfirmedMatrixTestRecordDocumentCommand(
                project_id="P1",
                output_dir=tmp_path,
                template_path=_template(tmp_path),
            )
        )


def test_generation_service_prefers_latest_project_folder_submitted_material(tmp_path: Path) -> None:
    writer = _Writer()
    folder_path = tmp_path / "project-folder"
    service = ConfirmedMatrixTestRecordDocumentGenerationService(
        preview_service=_PreviewService(_preview()),
        project_store=_ProjectStore(),
        writer=writer,
        folder_store=_FolderStore(folder_path),
    )

    result = service.generate(
        GenerateConfirmedMatrixTestRecordDocumentCommand(
            project_id="P1",
            output_dir=tmp_path / "fallback",
            template_path=_template(tmp_path),
        )
    )

    assert result.output_path == folder_path / "Submitted Material" / "DL-001 Test Record.docx"


def test_generation_service_keeps_header_metadata_blank_when_sources_missing(tmp_path: Path) -> None:
    writer = _Writer()
    service = ConfirmedMatrixTestRecordDocumentGenerationService(
        preview_service=_PreviewService(_preview()),
        project_store=_ProjectStore(),
        writer=writer,
        ltr_store=_LtrStore(records=()),
        intake_case_store=_IntakeCaseStore(case=None),
        intake_draft_store=_IntakeDraftStore(draft=None),
        application_form_store=_ApplicationFormStore(forms=()),
    )

    service.generate(
        GenerateConfirmedMatrixTestRecordDocumentCommand(
            project_id="P1",
            output_dir=tmp_path,
            template_path=_template(tmp_path),
        )
    )

    header_metadata = writer.calls[0]["header_metadata"]
    assert header_metadata.lab_test_request_number == ""
    assert header_metadata.applicable_specification == ""
    assert header_metadata.product_description == "Connector"


def test_generation_service_passes_step_mapped_requirements_to_writer(tmp_path: Path) -> None:
    writer = _Writer()
    preview = ConfirmedMatrixTestRecordPreview(
        project_id="P1",
        confirmed_matrix_id="cmv-1",
        preview_status="ready",
        groups=(
            ConfirmedMatrixTestRecordPreviewGroup(
                group_key="g1",
                group_label="Group 1",
                sample_quantity_expression="5",
                step_count=2,
                steps=(
                    ConfirmedMatrixTestRecordPreviewStep(
                        sequence=2,
                        raw_token="2",
                        test_item="LLCR",
                        section="6.2",
                        method="EIA-364-23",
                        condition="20mV max, 100mA max",
                        requirement="≤ 0.25 mΩ",
                    ),
                    ConfirmedMatrixTestRecordPreviewStep(
                        sequence=5,
                        raw_token="5",
                        test_item="LLCR",
                        section="6.2",
                        method="EIA-364-23",
                        condition="20mV max, 100mA max",
                        requirement="ΔR ≤ 0.17 mΩ",
                    ),
                ),
            ),
        ),
    )
    service = ConfirmedMatrixTestRecordDocumentGenerationService(
        preview_service=_PreviewService(preview),
        project_store=_ProjectStore(),
        writer=writer,
    )

    service.generate(
        GenerateConfirmedMatrixTestRecordDocumentCommand(
            project_id="P1",
            output_dir=tmp_path,
            template_path=_template(tmp_path),
        )
    )

    written_steps = writer.calls[0]["groups"][0].steps
    assert written_steps[0].requirement == "≤ 0.25 mΩ"
    assert written_steps[1].requirement == "ΔR ≤ 0.17 mΩ"


def test_generation_service_passes_step_quantity_projection_to_writer(tmp_path: Path) -> None:
    writer = _Writer()
    service = ConfirmedMatrixTestRecordDocumentGenerationService(
        preview_service=_PreviewService(_preview_with_quantity()),
        project_store=_ProjectStore(),
        writer=writer,
    )

    service.generate(
        GenerateConfirmedMatrixTestRecordDocumentCommand(
            project_id="P1",
            output_dir=tmp_path,
            template_path=_template(tmp_path),
        )
    )

    quantity = writer.calls[0]["groups"][0].steps[0].quantity
    assert quantity.status == "ready"
    assert quantity.total_readings == "6"


class _PreviewService:
    def __init__(self, preview: ConfirmedMatrixTestRecordPreview | None) -> None:
        self.preview = preview

    def build_preview(self, command):
        if self.preview is None:
            raise ConfirmedMatrixTestRecordPreviewNotFoundError("Active confirmed matrix not found.")
        return self.preview


class _Project:
    project_id = "P1"
    product_name = "Connector"
    project_no = "DL-001"


class _ProjectStore:
    def get(self, project_id: str):
        return _Project() if project_id == "P1" else None


class _BasicInformationReader:
    def __init__(self, snapshot: ConfirmedBasicInformationSnapshot | None) -> None:
        self.snapshot = snapshot

    def get_latest_confirmed(self, project_id: str):
        return self.snapshot if project_id == "P1" else None


class _Writer:
    def __init__(self) -> None:
        self.calls: list[dict[str, object]] = []

    def generate_from_confirmed_matrix(self, **kwargs):
        self.calls.append(kwargs)
        kwargs["output_path"].write_bytes(b"docx")
        return kwargs["output_path"]


class _FailingWriter:
    def generate_from_confirmed_matrix(self, **kwargs):
        raise ValueError("Template step table must contain at least 9 columns.")


class _FolderRecord:
    def __init__(self, folder_path: Path) -> None:
        self.folder_path = folder_path
        self.created_on = None


class _FolderStore:
    def __init__(self, folder_path: Path) -> None:
        self._folders = [_FolderRecord(folder_path)]

    def list_by_project(self, project_id: str):
        return self._folders if project_id == "P1" else []


def _preview() -> ConfirmedMatrixTestRecordPreview:
    return ConfirmedMatrixTestRecordPreview(
        project_id="P1",
        confirmed_matrix_id="cmv-1",
        preview_status="ready",
        groups=(
            ConfirmedMatrixTestRecordPreviewGroup(
                group_key="g1",
                group_label="Group 1",
                sample_quantity_expression="5",
                step_count=1,
                steps=(
                    ConfirmedMatrixTestRecordPreviewStep(
                        sequence=1,
                        raw_token="1",
                        test_item="Visual",
                        section="6.1",
                        method="EIA-364-18",
                        condition="10x",
                        requirement="No damage",
                    ),
                ),
            ),
        ),
    )


def _preview_with_quantity() -> ConfirmedMatrixTestRecordPreview:
    quantity = ConfirmedMatrixTestRecordStepQuantity(
        test_points_per_sample="3",
        readings_per_point="2",
        contact_points_per_sample="6",
        total_readings="6",
        status="ready",
        source="matrix_step_override",
        review_reason=None,
    )
    return ConfirmedMatrixTestRecordPreview(
        project_id="P1",
        confirmed_matrix_id="cmv-1",
        preview_status="ready",
        groups=(
            ConfirmedMatrixTestRecordPreviewGroup(
                group_key="g1",
                group_label="Group 1",
                sample_quantity_expression="5",
                step_count=1,
                steps=(
                    ConfirmedMatrixTestRecordPreviewStep(
                        sequence=1,
                        raw_token="1",
                        test_item="LLCR",
                        section="6.2",
                        method="EIA-364-23",
                        condition="20mV max, 100mA max",
                        requirement="≤ 0.25 mΩ",
                        quantity=quantity,
                    ),
                ),
            ),
        ),
    )


def _template(tmp_path: Path) -> Path:
    from docx import Document

    template = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Group Number: placeholder")
    step_table = document.add_table(rows=1, cols=9)
    step_table.rows[0].cells[0].text = "Step"
    equipment_table = document.add_table(rows=1, cols=7)
    equipment_table.rows[0].cells[0].text = "Equipment"
    document.save(template)
    return template


class _Ltr:
    def __init__(
        self,
        *,
        ltr_number: str = "DL-2026-05-003",
        status=LtrStatus.REGISTERED,
        notes: str | None = None,
        registered_on: date | None = None,
    ) -> None:
        self.ltr_number = ltr_number
        self.status = status
        self.notes = notes or json.dumps({"sample_description": "Coolpower HDF 3.40mm pin"})
        self.registered_on = registered_on or date(2026, 5, 30)


class _LtrStore:
    def __init__(self, records: tuple[object, ...] | None = None) -> None:
        self._records = records if records is not None else (_Ltr(),)

    def list_by_project(self, project_id: str):
        return list(self._records) if project_id == "P1" else []


class _IntakeCase:
    case_id = "case-1"


class _IntakeCaseStore:
    def __init__(self, case: object = "__default__") -> None:
        self._case = _IntakeCase() if case == "__default__" else case

    def get_by_confirmed_project(self, project_id: str):
        if project_id != "P1":
            return None
        return self._case


class _IntakeDraft:
    requested_testing_json = json.dumps(
        [
            {"test_to_be_performed": "Visual", "applicable_specification": "GS-12-1507"},
            {"test_to_be_performed": "LLCR", "applicable_specification": "GS-12-1507"},
        ]
    )


class _IntakeDraftStore:
    def __init__(self, draft: object = "__default__") -> None:
        self._draft = _IntakeDraft() if draft == "__default__" else draft

    def get_by_case(self, case_id: str):
        if case_id != "case-1":
            return None
        return self._draft


class _ApplicationForm:
    requested_testing = "Visual test by EIA-364-18; GS-00-000"


class _ApplicationFormStore:
    def __init__(self, forms: tuple[object, ...] | None = None) -> None:
        self._forms = forms if forms is not None else (_ApplicationForm(),)

    def list_by_project(self, project_id: str):
        return list(self._forms) if project_id == "P1" else []


def _basic_information(
    values: dict[str, str],
    *,
    version: int = 1,
) -> ConfirmedBasicInformationSnapshot:
    return ConfirmedBasicInformationSnapshot(
        project_id="P1",
        version=version,
        values=values,
        source_signature='{"source":"test"}',
        confirmed_at="2026-06-20T00:00:00+00:00",
        confirmed_by="tester",
    )
