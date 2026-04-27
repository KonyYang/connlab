from __future__ import annotations

from pathlib import Path

from docx import Document

from backend.infrastructure.office import (
    OfficeAutomationUnavailable,
    OfficeFacade,
    OfficeFileKind,
    OfficeLifecycleManager,
    WordDocumentGateway,
)


ROOT = Path(__file__).resolve().parents[2]


def test_office_facade_classifies_supported_file_types(tmp_path: Path) -> None:
    """OfficeFacade classifies common intake file types without business logic."""
    samples = {
        "request.docx": OfficeFileKind.DOCX,
        "legacy.doc": OfficeFileKind.DOC,
        "spec.pdf": OfficeFileKind.PDF,
        "mail.msg": OfficeFileKind.OUTLOOK_MSG,
        "image003.jpg": OfficeFileKind.IMAGE,
        "notes.bin": OfficeFileKind.UNKNOWN,
    }
    facade = OfficeFacade()

    for filename, expected_kind in samples.items():
        path = tmp_path / filename
        path.write_bytes(b"sample")

        classification = facade.classify_file(path)

        assert classification.original_name == filename
        assert classification.kind is expected_kind
        assert classification.size_bytes == 6
        assert classification.supported is (expected_kind is not OfficeFileKind.UNKNOWN)


def test_word_gateway_reads_docx_snapshot_with_header_footer_and_table(
    tmp_path: Path,
) -> None:
    """WordDocumentGateway extracts neutral Word content snapshots."""
    docx_path = tmp_path / "application.docx"
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Laboratory Testing Request"
    section.footer.paragraphs[0].text = "Form No. E-3718 Rev H"
    document.add_paragraph("Requested By: Jane Engineer")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Part Number / Revision"
    table.cell(0, 1).text = "Quantity"
    table.cell(1, 0).text = "ABC-123 Rev A"
    table.cell(1, 1).text = "12"
    document.save(docx_path)

    snapshot = WordDocumentGateway().read_word_document(docx_path)

    assert "Requested By: Jane Engineer" in snapshot.paragraphs
    assert snapshot.tables == [[
        ["Part Number / Revision", "Quantity"],
        ["ABC-123 Rev A", "12"],
    ]]
    assert "Laboratory Testing Request" in snapshot.headers
    assert "Form No. E-3718 Rev H" in snapshot.footers
    assert "ABC-123 Rev A | 12" in snapshot.raw_text


def test_office_facade_delegates_word_snapshot_read(tmp_path: Path) -> None:
    """OfficeFacade exposes Word document reading through the boundary."""
    docx_path = tmp_path / "minimal.docx"
    document = Document()
    document.add_paragraph("ConnLab application form")
    document.save(docx_path)

    snapshot = OfficeFacade().read_word_document(docx_path)

    assert snapshot.raw_text == "ConnLab application form"


def test_office_lifecycle_rejects_unimplemented_com_fallback() -> None:
    """COM automation remains centralized and unavailable until a scoped task."""
    manager = OfficeLifecycleManager()

    try:
        manager.require_com_fallback("Word")
    except OfficeAutomationUnavailable as exc:
        assert "Word COM fallback is not implemented" in str(exc)
    else:  # pragma: no cover - defensive assertion
        raise AssertionError("COM fallback should be unavailable")


def test_application_api_and_frontend_do_not_import_office_libraries_directly() -> None:
    """Application/API/frontend layers must not directly import Office libraries."""
    checked_roots = [
        ROOT / "backend" / "application",
        ROOT / "backend" / "api",
        ROOT / "frontend" / "src",
    ]
    violations: list[str] = []
    for checked_root in checked_roots:
        for path in checked_root.rglob("*"):
            if path.suffix not in {".py", ".ts", ".tsx"}:
                continue
            source = path.read_text(encoding="utf-8")
            if "win32com" in source or "from docx import" in source or "import docx" in source:
                violations.append(path.relative_to(ROOT).as_posix())

    assert violations == []


def test_only_office_infrastructure_imports_docx_in_backend_infrastructure() -> None:
    """python-docx usage is contained in the Office gateway for new infra code."""
    office_source = (
        ROOT / "backend" / "infrastructure" / "office" / "word_document_gateway.py"
    ).read_text(encoding="utf-8")

    assert "from docx import Document" in office_source
