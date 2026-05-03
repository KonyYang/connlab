from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from backend.application.intake_asset_preview_service import (
    IntakeAssetPreviewNotFoundError,
    IntakeAssetPreviewService,
)
from backend.domain import IntakeAsset, IntakeAssetRole


class AssetStore:
    """In-memory intake asset store for preview service tests."""

    def __init__(self, assets: list[IntakeAsset]) -> None:
        """Create the store with known assets."""
        self.assets = {asset.asset_id: asset for asset in assets}

    def get(self, asset_id: str) -> IntakeAsset | None:
        """Return one asset by id."""
        return self.assets.get(asset_id)


def test_docx_preview_returns_structured_application_form_sections(
    tmp_path: Path,
) -> None:
    """A registered DOCX asset returns path-free structured preview content."""
    docx_path = _create_application_docx(tmp_path / "application.docx")
    service = IntakeAssetPreviewService(
        AssetStore([_asset("asset-docx", docx_path, ".docx")])
    )

    preview = service.preview_asset("asset-docx")

    assert preview.kind == "docx_application_form"
    assert preview.metadata.original_name == "application.docx"
    assert not hasattr(preview.metadata, "stored_path")
    assert ("Requested By", "Alice Requestor") in [
        (field.label, field.value) for field in preview.fields
    ]
    assert any(table.title == "Test Sample Information" for table in preview.tables)
    sample_table = next(
        table for table in preview.tables if table.title == "Test Sample Information"
    )
    assert sample_table.rows[0][0] == "Connector A"
    assert any(table.title == "Requested Testing" for table in preview.tables)


def test_unsupported_asset_returns_metadata_preview(tmp_path: Path) -> None:
    """Non-rendered assets return metadata instead of failing."""
    pdf_path = tmp_path / "drawing.pdf"
    pdf_path.write_bytes(b"%PDF-1.4")
    service = IntakeAssetPreviewService(AssetStore([_asset("asset-pdf", pdf_path, ".pdf")]))

    preview = service.preview_asset("asset-pdf")

    assert preview.kind == "metadata_only"
    assert preview.metadata.original_name == "drawing.pdf"
    assert ("File type", "PDF") in [(field.label, field.value) for field in preview.fields]
    assert preview.message is not None
    assert "Detailed rendering is not implemented" in preview.message


def test_image_asset_returns_inline_preview(tmp_path: Path) -> None:
    """Small image attachments return a safe data URL preview."""
    image_path = tmp_path / "photo.png"
    image_path.write_bytes(b"\x89PNG\r\n\x1a\n")
    service = IntakeAssetPreviewService(
        AssetStore([_asset("asset-image", image_path, ".png", "image/png")])
    )

    preview = service.preview_asset("asset-image")

    assert preview.kind == "image"
    assert preview.title == "Image preview"
    assert preview.image_data_url is not None
    assert preview.image_data_url.startswith("data:image/png;base64,")


def test_non_application_docx_returns_metadata_only(tmp_path: Path) -> None:
    """Word attachments that are not LTR application forms are not rendered as forms."""
    docx_path = tmp_path / "notes.docx"
    document = Document()
    document.add_paragraph("General meeting notes")
    document.save(docx_path)
    service = IntakeAssetPreviewService(
        AssetStore([_asset("asset-notes", docx_path, ".docx")])
    )

    preview = service.preview_asset("asset-notes")

    assert preview.kind == "metadata_only"
    assert preview.title == "Word attachment"
    assert preview.message is not None
    assert "does not look like" in preview.message


def test_missing_asset_raises_not_found() -> None:
    """Missing asset ids are reported as not found."""
    service = IntakeAssetPreviewService(AssetStore([]))

    with pytest.raises(IntakeAssetPreviewNotFoundError):
        service.preview_asset("missing")


def _asset(
    asset_id: str,
    path: Path,
    extension: str,
    mime_type: str = "application/octet-stream",
) -> IntakeAsset:
    """Build one registered intake asset for tests."""
    return IntakeAsset(
        asset_id=asset_id,
        package_id="pkg-1",
        original_name=path.name,
        stored_path=path,
        extension=extension,
        mime_type=mime_type,
        size_bytes=path.stat().st_size,
        sha256=asset_id * 64,
        asset_role=IntakeAssetRole.APPLICATION_FORM_CANDIDATE,
    )


def _create_application_docx(path: Path) -> Path:
    """Create a small table-driven Laboratory Testing Request document."""
    document = Document()
    header = document.add_table(rows=4, cols=2)
    for row_index, (label, value) in enumerate(
        [
            ("Form No.", "E-3718"),
            ("Requested By", "Alice Requestor"),
            ("Email", "alice@example.com"),
            ("Description of Requested Testing", "Thermal cycling"),
        ]
    ):
        header.cell(row_index, 0).text = label
        header.cell(row_index, 1).text = value

    samples = document.add_table(rows=2, cols=5)
    for index, header_text in enumerate(
        ["Product Name", "Part Number", "Revision", "Contact Base Material", "Quantity"]
    ):
        samples.cell(0, index).text = header_text
    for index, value in enumerate(["Connector A", "PN-075", "A", "Ag", "12"]):
        samples.cell(1, index).text = value
    document.save(path)
    return path
