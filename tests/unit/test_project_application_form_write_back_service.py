from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document

from backend.application.project_basic_information_output import (
    ConfirmedBasicInformationSnapshot,
)
from backend.application.official_project_workspace_service import OfficialWorkspaceRecord
from backend.application.project_application_form_write_back_service import (
    ProjectApplicationFormWriteBackError,
    ProjectApplicationFormWriteBackService,
)
from backend.application.project_folder_required_forms_service import compute_sha256
from backend.application.project_output_record_service import ProjectOutputStatusItem
from backend.domain import (
    ApplicationForm,
    FileAsset,
    FileAssetType,
    Project,
    ProjectOutputKind,
    ProjectOutputSource,
    ProjectOutputStatus,
    ProjectStatus,
)


def test_application_form_write_back_updates_copied_submitted_material_docx(
    tmp_path: Path,
) -> None:
    official = tmp_path / "DL-001" / "DL-001 Connector Qualification test"
    submitted = official / "Submitted Material"
    target = submitted / "application.docx"
    _write_docx(target)
    output_store = _OutputStore()
    service = ProjectApplicationFormWriteBackService(
        project_store=_ProjectStore(),
        workspace_store=_WorkspaceStore(official),
        application_form_store=_ApplicationFormStore(),
        file_asset_store=_FileAssetStore(target),
        basic_information_reader=_BasicInformationReader(_basic_information()),
        output_record_service=output_store,
    )

    result = service.write_back("P1")

    assert result.status == "updated"
    assert result.target_path == target
    values = _read_table_values(target)
    assert values["LTR Number"] == "DL-001"
    assert values["Lab Performing the Tests"] == "Dongguan"
    assert values["Lab Personnel Assigned"] == "BI Leader"
    assert values["Date Lab Received Samples"] == "20 Jun 2026"
    assert values["Estimated Completion Date"] == "30 Jun 2026"
    assert values["Condition of Samples when Received"] == "Acceptable"
    assert output_store.commands
    assert str(target) == output_store.commands[-1].output_path
    assert output_store.commands[-1].source_context_signature == (
        "application-form:F1|"
        "basic:2@394f0d9772b800b7086b0d43d7a5bb748f33efafc474c39e9e25d4dc481712fe"
    )


def test_application_form_write_back_blocks_without_confirmed_basic_information(
    tmp_path: Path,
) -> None:
    official = tmp_path / "DL-001" / "DL-001 Connector Qualification test"
    target = official / "Submitted Material" / "application.docx"
    _write_docx(target)
    service = ProjectApplicationFormWriteBackService(
        project_store=_ProjectStore(),
        workspace_store=_WorkspaceStore(official),
        application_form_store=_ApplicationFormStore(),
        file_asset_store=_FileAssetStore(target),
        basic_information_reader=_BasicInformationReader(None),
        output_record_service=_OutputStore(),
        office=_RejectingOffice(),
    )

    try:
        service.write_back("P1")
    except ProjectApplicationFormWriteBackError as exc:
        assert "Basic Information" in str(exc)
    else:
        raise AssertionError("Expected Basic Information blocker.")


def test_application_form_write_back_blocks_user_changed_managed_target(
    tmp_path: Path,
) -> None:
    official = tmp_path / "DL-001" / "DL-001 Connector Qualification test"
    target = official / "Submitted Material" / "application.docx"
    _write_docx(target)
    original_sha = compute_sha256(target)
    target.write_text("operator changed", encoding="utf-8")
    output_store = _OutputStore(
        items=(
            ProjectOutputStatusItem(
                output_kind=ProjectOutputKind.SECTION2_WRITE_BACK,
                status=ProjectOutputStatus.CURRENT,
                output_path=str(target),
                source=ProjectOutputSource.SYSTEM_GENERATED,
                draft_id="D1",
                draft_version=1,
                reason="current",
                updated_at="2026-06-20T00:00:00+00:00",
                output_sha256=original_sha,
                output_size_bytes=1,
                source_context_signature="application-form:F1|basic:1@old",
            ),
        )
    )
    office = _RejectingOffice()
    service = ProjectApplicationFormWriteBackService(
        project_store=_ProjectStore(),
        workspace_store=_WorkspaceStore(official),
        application_form_store=_ApplicationFormStore(),
        file_asset_store=_FileAssetStore(target),
        basic_information_reader=_BasicInformationReader(_basic_information()),
        output_record_service=output_store,
        office=office,
    )

    try:
        service.write_back("P1")
    except ProjectApplicationFormWriteBackError as exc:
        assert "changed outside ConnLab" in str(exc)
    else:
        raise AssertionError("Expected managed fingerprint blocker.")
    assert office.calls == 0


def test_application_form_write_back_uses_basic_information_without_project_fallback(
    tmp_path: Path,
) -> None:
    official = tmp_path / "DL-001" / "DL-001 Connector Qualification test"
    target = official / "Submitted Material" / "application.docx"
    _write_docx(target)
    office = _CapturingOffice()
    service = ProjectApplicationFormWriteBackService(
        project_store=_ProjectStore(),
        workspace_store=_WorkspaceStore(official),
        application_form_store=_ApplicationFormStore(),
        file_asset_store=_FileAssetStore(target),
        basic_information_reader=_BasicInformationReader(
            _basic_information_with_conflicting_sources()
        ),
        output_record_service=_OutputStore(),
        office=office,
    )

    service.write_back("P1")

    assert "description_pn" not in office.fields
    assert "product_description" not in office.fields
    assert "test_item" not in office.fields
    assert "applicable_specifications" not in office.fields
    assert "requested_by" not in office.fields
    assert "requester" not in office.fields
    assert "location" not in office.fields
    assert "manufacturing_site" not in office.fields
    assert office.fields["project_leader"] == "BI Leader"
    assert "assigned_personnel" not in office.fields
    assert "Connector" not in office.fields.values()
    assert "Qualification Testing" not in office.fields.values()
    assert "MP Cao" not in office.fields.values()


def test_application_form_write_back_uses_parsed_assigned_personnel_only_as_fallback(
    tmp_path: Path,
) -> None:
    official = tmp_path / "DL-001" / "DL-001 Connector Qualification test"
    target = official / "Submitted Material" / "application.docx"
    _write_docx(target)
    office = _CapturingOffice()
    service = ProjectApplicationFormWriteBackService(
        project_store=_ProjectStore(),
        workspace_store=_WorkspaceStore(official),
        application_form_store=_ApplicationFormStore(assigned_personnel="Parsed Engineer"),
        file_asset_store=_FileAssetStore(target),
        basic_information_reader=_BasicInformationReader(
            _basic_information_with_blank_project_leader()
        ),
        output_record_service=_OutputStore(),
        office=office,
    )

    service.write_back("P1")

    assert office.fields["project_leader"] == "Parsed Engineer"
    assert "assigned_personnel" not in office.fields


def test_application_form_write_back_blocks_missing_required_value_before_office(
    tmp_path: Path,
) -> None:
    official = tmp_path / "DL-001" / "DL-001 Connector Qualification test"
    target = official / "Submitted Material" / "application.docx"
    _write_docx(target)
    office = _RejectingOffice()
    service = ProjectApplicationFormWriteBackService(
        project_store=_ProjectStore(),
        workspace_store=_WorkspaceStore(official),
        application_form_store=_ApplicationFormStore(),
        file_asset_store=_FileAssetStore(target),
        basic_information_reader=_BasicInformationReader(_basic_information_missing_lab()),
        output_record_service=_OutputStore(),
        office=office,
    )

    try:
        service.write_back("P1")
    except ProjectApplicationFormWriteBackError as exc:
        assert "Lab Performing the Tests" in str(exc)
    else:
        raise AssertionError("Expected required Basic Information blocker.")
    assert office.calls == 0


def test_application_form_write_back_blocks_gateway_critical_failure(
    tmp_path: Path,
) -> None:
    official = tmp_path / "DL-001" / "DL-001 Connector Qualification test"
    target = official / "Submitted Material" / "application.docx"
    _write_docx(target)
    output_store = _OutputStore()
    service = ProjectApplicationFormWriteBackService(
        project_store=_ProjectStore(),
        workspace_store=_WorkspaceStore(official),
        application_form_store=_ApplicationFormStore(),
        file_asset_store=_FileAssetStore(target),
        basic_information_reader=_BasicInformationReader(_basic_information()),
        output_record_service=output_store,
        office=_FailingOffice("Application Form header LTR location not found."),
    )

    try:
        service.write_back("P1")
    except ProjectApplicationFormWriteBackError as exc:
        assert "header LTR" in str(exc)
    else:
        raise AssertionError("Expected gateway critical failure blocker.")
    assert not output_store.commands


def test_application_form_write_back_uses_selected_request_material_target(
    tmp_path: Path,
) -> None:
    official = tmp_path / "DL-001" / "DL-001 Connector Qualification test"
    submitted = official / "Submitted Material"
    selected_target = submitted / "selected request.docx"
    other_target = submitted / "other request.docx"
    _write_docx(selected_target)
    _write_docx(other_target)
    office = _CapturingOffice()
    service = ProjectApplicationFormWriteBackService(
        project_store=_ProjectStore(),
        workspace_store=_WorkspaceStore(official),
        application_form_store=_ApplicationFormStore(),
        file_asset_store=_FileAssetStore(other_target),
        request_material_collection_store=_CollectionStore(selected_target),
        basic_information_reader=_BasicInformationReader(_basic_information()),
        output_record_service=_OutputStore(),
        office=office,
    )

    result = service.write_back("P1")

    assert result.target_path == selected_target
    assert office.calls == 1
    assert office.owned_session_calls == 1
    assert office.default_calls == 0


def test_application_form_write_back_allows_rebuilt_target_restored_to_source(
    tmp_path: Path,
) -> None:
    official = tmp_path / "DL-001" / "DL-001 Connector Qualification test"
    target = official / "Submitted Material" / "application.docx"
    _write_docx(target)
    source_sha = compute_sha256(target)
    output_store = _OutputStore(
        items=(
            ProjectOutputStatusItem(
                output_kind=ProjectOutputKind.SECTION2_WRITE_BACK,
                status=ProjectOutputStatus.CURRENT,
                output_path=str(target),
                source=ProjectOutputSource.SYSTEM_GENERATED,
                draft_id="D1",
                draft_version=1,
                reason="current",
                updated_at="2026-06-20T00:00:00+00:00",
                output_sha256="old-managed-write-back-sha",
                output_size_bytes=1,
                source_context_signature="application-form:F1|basic:1@old",
            ),
        )
    )
    office = _CapturingOffice()
    service = ProjectApplicationFormWriteBackService(
        project_store=_ProjectStore(),
        workspace_store=_WorkspaceStore(official),
        application_form_store=_ApplicationFormStore(),
        file_asset_store=_FileAssetStore(target),
        request_material_collection_store=_CollectionStore(target, source_sha),
        basic_information_reader=_BasicInformationReader(_basic_information()),
        output_record_service=output_store,
        office=office,
    )

    result = service.write_back("P1")

    assert result.target_path == target
    assert office.calls == 1


def _write_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for label in (
        "LTR Number",
        "Lab Performing the Tests",
        "Lab Personnel Assigned",
        "Date Lab Received Samples",
        "Estimated Completion Date",
        "Condition of Samples when Received",
    ):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = ""
    document.save(path)


def _read_table_values(path: Path) -> dict[str, str]:
    document = Document(path)
    values = {
        row.cells[0].text.strip(): row.cells[1].text.strip()
        for row in document.tables[0].rows
    }
    return values


class _ProjectStore:
    def get(self, project_id: str) -> Project | None:
        return Project(
            project_id=project_id,
            project_no="DL-001",
            product_name="Connector",
            requestor="MP Cao",
            status=ProjectStatus.FOLDER_CREATED,
            business_unit="BU",
        )


class _WorkspaceStore:
    def __init__(self, official: Path) -> None:
        self.official = official

    def get_by_project(self, project_id: str) -> OfficialWorkspaceRecord:
        return OfficialWorkspaceRecord(
            workspace_id="W1",
            project_id=project_id,
            dl_number="DL-001",
            local_workspace_path=self.official.parent,
            source_book_path=self.official.parent / "Source Book",
            official_folder_path=self.official,
            manifest_path=self.official.parent / ".connlab" / "manifest.json",
            template_source_path=Path("D:/Source/Template/DL-XXXX-YY-ZZZ project"),
            created_at="2026-06-19T00:00:00Z",
        )


class _ApplicationFormStore:
    def __init__(self, *, assigned_personnel: str = "") -> None:
        self.assigned_personnel = assigned_personnel

    def list_by_project(self, project_id: str) -> list[ApplicationForm]:
        return [
            ApplicationForm(
                form_id="F1",
                project_id=project_id,
                form_no="E-3718",
                revision="H",
                requester="MP Cao",
                email="mp@example.test",
                requested_testing="Qualification Testing",
                received_date="2026-06-01",
                assigned_personnel=self.assigned_personnel,
            )
        ]


class _FileAssetStore:
    def __init__(self, target: Path) -> None:
        self.target = target

    def list_by_project(self, project_id: str) -> list[FileAsset]:
        return [
            FileAsset(
                asset_id="A1",
                project_id=project_id,
                asset_type=FileAssetType.APPLICATION_FORM,
                path=self.target,
                original_name="application.docx",
                source_role="selected_application_form",
            )
        ]


class _CollectionStore:
    def __init__(self, target: Path, source_sha: str | None = None) -> None:
        self.target = target
        self.source_sha = source_sha or compute_sha256(target)

    def latest_by_project(self, project_id: str):
        return SimpleNamespace(collection_id="C1")

    def list_items(self, collection_id: str):
        return (
            SimpleNamespace(
                source_asset_id="A1",
                source_asset_type="application_form",
                source_role="selected_application_form",
                source_path=self.target,
                target_area="submitted_material",
                target_path=self.target,
                sha256=self.source_sha,
            ),
        )


class _OutputStore:
    def __init__(
        self,
        items: tuple[ProjectOutputStatusItem, ...] = tuple(),
    ) -> None:
        self.commands = []
        self.items = items

    def get_status_summary(self, project_id: str):
        return type(
            "_Summary",
            (),
            {"active_draft_id": "D1", "items": self.items},
        )()

    def register_output(self, command):
        self.commands.append(command)

        class _Record:
            output_record_id = "O1"

        return _Record()


class _BasicInformationReader:
    def __init__(self, snapshot: ConfirmedBasicInformationSnapshot | None) -> None:
        self.snapshot = snapshot

    def get_latest_confirmed(
        self, project_id: str
    ) -> ConfirmedBasicInformationSnapshot | None:
        return self.snapshot


class _RejectingOffice:
    def __init__(self) -> None:
        self.calls = 0

    def write_word_application_form_fields(self, source_path: Path, fields: dict[str, str]):
        self.calls += 1
        raise AssertionError("Office writer should not be called.")

    def write_word_application_form_fields_with_owned_session(
        self,
        source_path: Path,
        fields: dict[str, str],
    ):
        self.calls += 1
        raise AssertionError("Office writer should not be called.")


class _CapturingOffice:
    def __init__(self) -> None:
        self.calls = 0
        self.default_calls = 0
        self.owned_session_calls = 0
        self.fields: dict[str, str] = {}

    def write_word_application_form_fields(self, source_path: Path, fields: dict[str, str]):
        self.default_calls += 1
        return self._write(fields)

    def write_word_application_form_fields_with_owned_session(
        self,
        source_path: Path,
        fields: dict[str, str],
    ):
        self.owned_session_calls += 1
        return self._write(fields)

    def _write(self, fields: dict[str, str]):
        self.calls += 1
        self.fields = dict(fields)

        class _Result:
            changed_fields = tuple()
            unchanged_fields = tuple()
            warnings = tuple()

        return _Result()


class _FailingOffice:
    def __init__(self, message: str) -> None:
        self.message = message

    def write_word_application_form_fields(self, source_path: Path, fields: dict[str, str]):
        raise ValueError(self.message)

    def write_word_application_form_fields_with_owned_session(
        self,
        source_path: Path,
        fields: dict[str, str],
    ):
        raise ValueError(self.message)


def _basic_information() -> ConfirmedBasicInformationSnapshot:
    return ConfirmedBasicInformationSnapshot(
        project_id="P1",
        version=2,
        values={
            "dl_number": "DL-001",
            "project_number": "PN-1",
            "product_description": "Connector from Basic Info",
            "test_item": "BI Qualification Test",
            "requested_by": "Requester BI",
            "requestor_email": "requester@example.test",
            "date_lab_received_samples": "20 Jun 2026",
            "estimated_completion_date": "30 Jun 2026",
            "condition_of_samples_when_received": "Acceptable",
            "lab_performing_tests": "Dongguan",
            "project_leader": "BI Leader",
        },
        source_signature='{"dl_number":"DL-001"}',
        confirmed_at="2026-06-20T00:00:00+00:00",
        confirmed_by="Lab User",
    )


def _basic_information_with_conflicting_sources() -> ConfirmedBasicInformationSnapshot:
    return ConfirmedBasicInformationSnapshot(
        project_id="P1",
        version=3,
        values={
            "dl_number": "DL-001",
            "description_pn": "101-BI",
            "product_description": "",
            "test_item": "BI Test Item",
            "requested_by": "BI Requester",
            "location": "BI Dongguan",
            "project_leader": "BI Leader",
            "applicable_specifications": "BI Spec",
            "lab_performing_tests": "BI Lab",
            "date_lab_received_samples": "21 Jun 2026",
            "estimated_completion_date": "30 Jun 2026",
            "condition_of_samples_when_received": "Acceptable",
        },
        source_signature='{"dl_number":"DL-001","version":3}',
        confirmed_at="2026-06-21T00:00:00+00:00",
        confirmed_by="Lab User",
    )


def _basic_information_with_blank_project_leader() -> ConfirmedBasicInformationSnapshot:
    return ConfirmedBasicInformationSnapshot(
        project_id="P1",
        version=4,
        values={
            "dl_number": "DL-001",
            "project_leader": "",
            "lab_performing_tests": "Dongguan",
            "date_lab_received_samples": "20 Jun 2026",
            "estimated_completion_date": "30 Jun 2026",
            "condition_of_samples_when_received": "Acceptable",
        },
        source_signature='{"dl_number":"DL-001","version":4}',
        confirmed_at="2026-06-21T00:00:00+00:00",
        confirmed_by="Lab User",
    )


def _basic_information_missing_lab() -> ConfirmedBasicInformationSnapshot:
    return ConfirmedBasicInformationSnapshot(
        project_id="P1",
        version=5,
        values={
            "dl_number": "DL-001",
            "project_leader": "BI Leader",
            "date_lab_received_samples": "20 Jun 2026",
            "estimated_completion_date": "30 Jun 2026",
            "condition_of_samples_when_received": "Acceptable",
        },
        source_signature='{"dl_number":"DL-001","version":5}',
        confirmed_at="2026-06-21T00:00:00+00:00",
        confirmed_by="Lab User",
    )
