from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from backend.modules.intake import ApplicationFormParser
from tests.fixtures.real_style_application_forms import build_real_style_application_form


def test_application_form_parser_extracts_fields_and_sample_row(tmp_path: Path) -> None:
    docx_path = tmp_path / "application-form.docx"
    document = Document()
    document.add_paragraph("Form No.: E-3718")
    document.add_paragraph("Form Rev: H")
    _add_key_value_table(
        document,
        [
            ("Reference Doc", "LAB-QA-001"),
            ("Lab Test Request Number", "LTR-001"),
            ("Requested By", "Alice"),
            ("Phone", "555-0100"),
            ("Date", "2026-04-26"),
            ("Email", "alice@example.com"),
            ("Business Unit", "BU-1"),
            ("Mfg. Site", "Plant 1"),
            ("Project #", "PRJ-001"),
            ("Requested Testing Completion Date", "2026-05-01"),
            ("Results Format", "PDF"),
            ("Test Type", "Qualification"),
            ("Sample Status", "Prototype"),
            ("Project Type", "New Product"),
            ("Post-testing disposition", "Return samples"),
            ("Description of Requested Testing", "Salt spray test"),
            ("Confidential", "Yes"),
            ("Subcontract", "No"),
            ("Additional Information", "Handle with care"),
            ("Send Copies", "bob@example.com"),
            ("Lab", "Connector Lab"),
            ("Assigned Personnel", "Charlie"),
            ("Received Date", "2026-04-27"),
            ("Estimated Completion Date", "2026-05-02"),
            ("Sample Condition", "Good"),
        ],
    )
    sample_table = document.add_table(rows=2, cols=8)
    headers = [
        "Product Name",
        "Part Number",
        "Revision",
        "Lot/Traceability",
        "Material",
        "Plating",
        "Housing Material",
        "Quantity",
    ]
    values = [
        "Connector",
        "PN-001",
        "A",
        "LOT-1",
        "Copper",
        "Tin",
        "LCP",
        "12",
    ]
    for index, header in enumerate(headers):
        sample_table.cell(0, index).text = header
        sample_table.cell(1, index).text = values[index]
    document.save(docx_path)

    parsed = ApplicationFormParser().parse(docx_path)

    assert parsed.form_no == "E-3718"
    assert parsed.form_rev == "H"
    assert parsed.reference_doc == "LAB-QA-001"
    assert parsed.lab_test_request_number == "LTR-001"
    assert parsed.requested_by == "Alice"
    assert parsed.project_number == "PRJ-001"
    assert parsed.requested_testing_description == "Salt spray test"
    assert parsed.lab_section.assigned_personnel == "Charlie"
    assert len(parsed.samples) == 1
    assert parsed.samples[0].part_number == "PN-001"
    assert parsed.samples[0].quantity == "12"


def test_application_form_parser_preserves_repeated_sample_placeholder_columns(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "sample-repeated-placeholders.docx"
    document = Document()
    sample_table = document.add_table(rows=3, cols=8)
    headers = [
        "Product Name",
        "Part Number / Revision",
        "Traceability Manufacturing Lot Info",
        "Contact Base Material",
        "Contact Plating",
        "Contact Lubricant",
        "Housing Material",
        "Quantity",
    ]
    first_row = ["EK340", "10178799-001LF", "/", "C10070", "Silver Over Ni", "/", "PA9T", "36"]
    repeated_placeholder_row = [
        "MATING BUSBAR",
        "10158889-32",
        "/",
        "C1100R-1/2H",
        "Silver Over Ni",
        "/",
        "/",
        "5",
    ]
    for column, header in enumerate(headers):
        sample_table.cell(0, column).text = header
        sample_table.cell(1, column).text = first_row[column]
        sample_table.cell(2, column).text = repeated_placeholder_row[column]
    document.save(docx_path)

    parsed = ApplicationFormParser().parse(docx_path)

    assert len(parsed.samples) == 2
    assert parsed.samples[0].lubricant == "/"
    assert parsed.samples[0].housing_material == "PA9T"
    assert parsed.samples[0].quantity == "36"
    assert parsed.samples[1].lubricant == "/"
    assert parsed.samples[1].housing_material == "/"
    assert parsed.samples[1].quantity == "5"


def test_application_form_parser_tolerates_missing_optional_fields(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "minimal-form.docx"
    document = Document()
    _add_key_value_table(
        document,
        [
            ("Form No.", "E-3718"),
            ("Requested By", "Alice"),
        ],
    )
    document.save(docx_path)

    parsed = ApplicationFormParser().parse(docx_path)

    assert parsed.form_no == "E-3718"
    assert parsed.requested_by == "Alice"
    assert parsed.email is None
    assert parsed.samples == ()
    assert parsed.lab_section.lab is None


def test_application_form_parser_extracts_value_after_neighbor_label(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "neighbor-label-value.docx"
    document = Document()
    table = document.add_table(rows=1, cols=5)
    values = ["Email:", "kris.li@example.com", "Business Unit:", "Mfg. Site:", "Nantong"]
    for index, value in enumerate(values):
        table.cell(0, index).text = value
    document.save(docx_path)

    parsed = ApplicationFormParser().parse(docx_path)

    assert parsed.email == "kris.li@example.com"
    assert parsed.manufacturing_site == "Nantong"


def test_application_form_parser_extracts_real_style_applicant_fixture(
    tmp_path: Path,
) -> None:
    docx_path = build_real_style_application_form(tmp_path / "real-style-applicant.docx")

    parsed = ApplicationFormParser().parse(docx_path)

    assert parsed.form_no == "E-3718"
    assert parsed.form_rev == "H"
    assert parsed.requested_by == "Alice Requestor"
    assert parsed.phone == "555-0100"
    assert parsed.request_date == "2026-04-27"
    assert parsed.email == "alice.requestor@example.com"
    assert parsed.business_unit == "Industrial"
    assert parsed.manufacturing_site == "Plant 7"
    assert parsed.project_number == "PRJ-038-A"
    assert parsed.project_type == "Qualification"
    assert parsed.subcontract == "No"
    assert parsed.requested_testing_description == "Thermal cycling and contact resistance"
    assert parsed.lab_section.lab == "Connector Lab"
    assert parsed.lab_section.assigned_personnel == "Charlie Tester"
    assert parsed.lab_section.received_date == "2026-04-28"
    assert parsed.lab_section.estimated_completion_date == "2026-05-15"
    assert parsed.lab_section.sample_condition == "Good"
    assert len(parsed.samples) == 1
    assert parsed.samples[0].product_name == "Synthetic Connector"
    assert parsed.samples[0].part_number == "PN-038-A"
    assert parsed.samples[0].revision == "A"
    assert parsed.samples[0].lot_or_traceability == "LOT-038"
    assert parsed.samples[0].quantity == "24"


def test_application_form_parser_extracts_yes_no_content_controls(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "content-control-form.docx"
    document = Document()
    _add_content_control_row(document, "Confidential tests or samples?", "Confidential", "No")
    _add_content_control_row(document, "Can testing be subcontracted?", "Subcontracted", "Yes")
    document.save(docx_path)

    parsed = ApplicationFormParser().parse(docx_path)

    assert parsed.confidential == "No"
    assert parsed.subcontract == "Yes"


def test_application_form_parser_extracts_two_column_requested_testing_and_additional_information(
    tmp_path: Path,
) -> None:
    """Real DOCX structure: two-column requested-testing table plus dedicated Additional Information block."""
    docx_path = tmp_path / "real-structure-form.docx"
    document = Document()
    document.add_paragraph("Form No.: E-3718")
    document.add_paragraph("Form Rev: H")
    _add_key_value_table(
        document,
        [
            ("Requested By", "Alice"),
            ("Email", "alice@example.com"),
        ],
    )
    sample_table = document.add_table(rows=2, cols=8)
    headers = [
        "Product Name",
        "Part Number",
        "Revision",
        "Lot/Traceability",
        "Material",
        "Plating",
        "Housing Material",
        "Quantity",
    ]
    values = ["Connector", "PN-001", "A", "LOT-1", "Copper", "Tin", "LCP", "12"]
    for index, header in enumerate(headers):
        sample_table.cell(0, index).text = header
        sample_table.cell(1, index).text = values[index]

    rt_table = document.add_table(rows=3, cols=2)
    rt_table.cell(0, 0).text = "Tests to be Performed"
    rt_table.cell(0, 1).text = "Applicable Specifications"
    rt_table.cell(1, 0).text = "依附件表格要求"
    rt_table.cell(1, 1).text = "GS-12-2652-22"
    rt_table.cell(2, 0).text = "T-rise"
    rt_table.cell(2, 1).text = "EIA-364-70D"

    _add_content_control_row(document, "Confidential tests or samples?", "Confidential", "No")
    _add_content_control_row(document, "Can testing be subcontracted?", "Subcontracted", "Yes")

    document.add_paragraph("Additional Information")
    ai_table = document.add_table(rows=1, cols=1)
    ai_table.cell(0, 0).text = "依附件EVE客户表格要求..."

    sc_table = document.add_table(rows=1, cols=2)
    sc_table.cell(0, 0).text = "Send copies of test results/reports to:"
    sc_table.cell(0, 1).text = "Mike.Rao@fci.com;Yang.Fu@fci.com"

    document.save(docx_path)

    parsed = ApplicationFormParser().parse(docx_path)

    assert len(parsed.requested_testing_rows) == 2
    assert parsed.requested_testing_rows[0].test_to_be_performed == "依附件表格要求"
    assert parsed.requested_testing_rows[0].applicable_specification == "GS-12-2652-22"
    assert parsed.requested_testing_rows[1].test_to_be_performed == "T-rise"
    assert parsed.requested_testing_rows[1].applicable_specification == "EIA-364-70D"
    assert parsed.requested_testing_description == "依附件表格要求\nT-rise"
    assert parsed.additional_information == "依附件EVE客户表格要求..."
    assert parsed.confidential == "No"
    assert parsed.subcontract == "Yes"
    assert parsed.send_copies_recipients == "Mike.Rao@fci.com;Yang.Fu@fci.com"


def test_application_form_parser_extracts_comparable_tester_modified_fixture(
    tmp_path: Path,
) -> None:
    applicant_path = build_real_style_application_form(
        tmp_path / "real-style-applicant.docx",
    )
    tester_path = build_real_style_application_form(
        tmp_path / "real-style-tester.docx",
        tester_modified=True,
    )

    applicant = ApplicationFormParser().parse(applicant_path)
    tester = ApplicationFormParser().parse(tester_path)

    assert applicant.form_no == tester.form_no == "E-3718"
    assert applicant.form_rev == tester.form_rev == "H"
    assert applicant.requested_by == tester.requested_by == "Alice Requestor"
    assert applicant.requested_testing_description == tester.requested_testing_description
    assert applicant.project_number == "PRJ-038-A"
    assert tester.project_number == "PRJ-038-T"
    assert applicant.samples[0].part_number == "PN-038-A"
    assert tester.samples[0].part_number == "PN-038-T"
    assert tester.lab_section.sample_condition == "Good, tester reviewed"


def test_application_form_parser_calibrates_section1_content_controls(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "section1-content-controls.docx"
    document = Document()
    _add_key_value_table(
        document,
        [
            ("Requested By:", "Neo Xu"),
            ("Phone #:", "0513-80167327"),
            ("Date:", ""),
            ("Email:", "Neo.Xu@fci.com"),
            ("Business Unit:", "Mfg. Site:"),
            ("Project #:", "CoolPowerHD 2X3.5"),
            ("Results Format:", "Requested Testing Completion Date:"),
        ],
    )
    for value in [
        "10/11/2024",
        "Power Solutions",
        "Nantong",
        "Formal Report (Customer)",
        "11/15/2024",
        "Customer Specific Testing",
        "Production",
        "New Product Development",
        "Keep in the Lab",
        "No",
        "Yes",
        "Dongguan",
        "10/17/2024",
        "10/31/2024",
        "Acceptable",
    ]:
        _add_content_control_paragraph(document, value)
    document.save(docx_path)

    parsed = ApplicationFormParser().parse(docx_path)

    assert parsed.phone == "0513-80167327"
    assert parsed.request_date == "10/11/2024"
    assert parsed.business_unit == "Power Solutions"
    assert parsed.manufacturing_site == "Nantong"
    assert parsed.results_format == "Formal Report (Customer)"
    assert parsed.requested_completion_date == "11/15/2024"
    assert parsed.sample_status == "Production"
    assert parsed.project_type == "New Product Development"
    assert parsed.post_testing_disposition == "Keep in the Lab"


def test_application_form_parser_keeps_section1_placeholders_in_content_control_order(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "section1-placeholder-content-controls.docx"
    document = Document()
    _add_key_value_table(
        document,
        [
            ("Date:", ""),
            ("Business Unit:", "Mfg. Site:"),
            ("Results Format:", "Requested Testing Completion Date:"),
        ],
    )
    for value in [
        "Click here to enter a date.",
        "Power Solutions",
        "Dongguan",
        "Formal Report (Internal)",
        "4/23/2025",
        "Product/Process Qualification",
        "Pre-production",
        "New Product Development",
        "Keep in the Lab",
        "No",
        "Yes",
    ]:
        _add_content_control_paragraph(document, value)
    document.save(docx_path)

    parsed = ApplicationFormParser().parse(docx_path)

    assert parsed.request_date is None
    assert parsed.business_unit == "Power Solutions"
    assert parsed.manufacturing_site == "Dongguan"
    assert parsed.results_format == "Formal Report (Internal)"
    assert parsed.requested_completion_date == "4/23/2025"
    assert parsed.test_type == "Product/Process Qualification"
    assert parsed.sample_status == "Pre-production"
    assert parsed.project_type == "New Product Development"
    assert parsed.post_testing_disposition == "Keep in the Lab"


def test_application_form_parser_does_not_shift_disposition_into_project_type(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "missing-project-type-content-control.docx"
    document = Document()
    _add_key_value_table(
        document,
        [
            ("Date:", ""),
            ("Business Unit:", "Mfg. Site:"),
            ("Results Format:", "Requested Testing Completion Date:"),
        ],
    )
    for value in [
        "4/23/2025",
        "Power Solutions",
        "Dongguan",
        "Formal Report (Internal)",
        "5/23/2025",
        "Product/Process Qualification",
        "Prototype",
        "Send Back to Requestor",
        "No",
        "Yes",
    ]:
        _add_content_control_paragraph(document, value)
    document.save(docx_path)

    parsed = ApplicationFormParser().parse(docx_path)

    assert parsed.project_type is None
    assert parsed.post_testing_disposition == "Send Back to Requestor"


def test_application_form_parser_handles_section1_without_mfg_site_control(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "section1-no-mfg-site-control.docx"
    document = Document()
    _add_key_value_table(
        document,
        [
            ("Date:", ""),
            ("Business Unit:", "Mfg. Site:"),
            ("Results Format:", "Requested Testing Completion Date:"),
        ],
    )
    for value in [
        "5/19/2023",
        "Power",
        "Formal Report (Customer)",
        "6/19/2023",
        "Product/Process Qualification",
        "Production",
        "New Product Development",
        "Send Back to Requestor",
        "No",
        "Yes",
    ]:
        _add_content_control_paragraph(document, value)
    document.save(docx_path)

    parsed = ApplicationFormParser().parse(docx_path)

    assert parsed.request_date == "5/19/2023"
    assert parsed.business_unit == "Power"
    assert parsed.manufacturing_site is None
    assert parsed.results_format == "Formal Report (Customer)"
    assert parsed.requested_completion_date == "6/19/2023"
    assert parsed.test_type == "Product/Process Qualification"
    assert parsed.sample_status == "Production"
    assert parsed.project_type == "New Product Development"
    assert parsed.post_testing_disposition == "Send Back to Requestor"


def test_application_form_parser_keeps_blank_fields_blank_when_neighbor_is_label(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "blank-neighbor-label.docx"
    document = Document()
    _add_key_value_table(
        document,
        [
            ("Business Unit:", "Mfg. Site:"),
            ("Results Format:", "Requested Testing Completion Date:"),
        ],
    )
    document.save(docx_path)

    parsed = ApplicationFormParser().parse(docx_path)

    assert parsed.business_unit is None
    assert parsed.results_format is None


def _add_key_value_table(document: Document, pairs: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=len(pairs), cols=2)
    for row_index, (label, value) in enumerate(pairs):
        table.cell(row_index, 0).text = label
        table.cell(row_index, 1).text = value


def _add_content_control_row(
    document: Document,
    label: str,
    tag: str,
    value: str,
) -> None:
    table = document.add_table(rows=1, cols=1)
    row = table.rows[0]._tr
    table.cell(0, 0).text = label
    row.append(_content_control_cell(tag, value))


def _add_content_control_paragraph(document: Document, value: str) -> None:
    """Add a content control paragraph with a visible value."""
    paragraph = document.add_paragraph()
    paragraph._p.append(_content_control_run(value))


def _content_control_cell(tag: str, value: str):
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), tag)
    tag_element = OxmlElement("w:tag")
    tag_element.set(qn("w:val"), tag)
    sdt_pr.append(alias)
    sdt_pr.append(tag_element)
    sdt_content = OxmlElement("w:sdtContent")
    cell = OxmlElement("w:tc")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = value
    run.append(text)
    paragraph.append(run)
    cell.append(paragraph)
    sdt_content.append(cell)
    sdt.append(sdt_pr)
    sdt.append(sdt_content)
    return sdt


def _content_control_run(value: str):
    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = value
    run.append(text)
    paragraph.append(run)
    sdt_content.append(paragraph)
    sdt.append(sdt_content)
    return sdt
