from __future__ import annotations

from dataclasses import replace
from pathlib import Path

import pytest
from docx import Document

from backend.application.application_form_eligibility_service import (
    ApplicationFormEligibility,
)
from backend.application.intake_form_selection_service import (
    IntakeFormSelectionService,
    IntakeSelectionError,
    IntakeSelectionNotFoundError,
)
from backend.domain import (
    IntakeAsset,
    IntakeAssetRole,
    IntakeCase,
    IntakeCaseStatus,
    IntakeDraft,
    IntakePackage,
    IntakePackageSourceType,
    IntakePackageStatus,
    LtrRecord,
    LtrStatus,
)


class PackageStore:
    def __init__(self, packages: list[IntakePackage]) -> None:
        self.packages = {package.package_id: package for package in packages}

    def get(self, package_id: str) -> IntakePackage | None:
        return self.packages.get(package_id)

    def list(self) -> list[IntakePackage]:
        return list(self.packages.values())

    def delete(self, package_id: str) -> bool:
        return self.packages.pop(package_id, None) is not None


class AssetStore:
    def __init__(self, assets: list[IntakeAsset]) -> None:
        self.assets = {asset.asset_id: asset for asset in assets}

    def get(self, asset_id: str) -> IntakeAsset | None:
        return self.assets.get(asset_id)

    def list_by_package(self, package_id: str) -> list[IntakeAsset]:
        return [asset for asset in self.assets.values() if asset.package_id == package_id]

    def update(self, asset: IntakeAsset) -> IntakeAsset:
        self.assets[asset.asset_id] = asset
        return asset

    def delete_by_package(self, package_id: str) -> int:
        keys = [key for key, asset in self.assets.items() if asset.package_id == package_id]
        for key in keys:
            del self.assets[key]
        return len(keys)


class CaseStore:
    def __init__(self, cases: list[IntakeCase] | None = None) -> None:
        self.cases = {case.case_id: case for case in cases or []}

    def create(self, case: IntakeCase) -> IntakeCase:
        self.cases[case.case_id] = case
        return case

    def list_by_package(self, package_id: str) -> list[IntakeCase]:
        return [case for case in self.cases.values() if case.package_id == package_id]

    def update(self, case: IntakeCase) -> IntakeCase:
        self.cases[case.case_id] = case
        return case

    def delete_by_package(self, package_id: str) -> int:
        keys = [key for key, case in self.cases.items() if case.package_id == package_id]
        for key in keys:
            del self.cases[key]
        return len(keys)


class DraftStore:
    def __init__(self, drafts: list[IntakeDraft] | None = None) -> None:
        self.drafts = {draft.draft_id: draft for draft in drafts or []}

    def create(self, draft: IntakeDraft) -> IntakeDraft:
        self.drafts[draft.draft_id] = draft
        return draft

    def get_by_case(self, case_id: str) -> IntakeDraft | None:
        return next((draft for draft in self.drafts.values() if draft.case_id == case_id), None)

    def update(self, draft: IntakeDraft) -> IntakeDraft:
        self.drafts[draft.draft_id] = draft
        return draft

    def delete_by_case(self, case_id: str) -> int:
        keys = [key for key, draft in self.drafts.items() if draft.case_id == case_id]
        for key in keys:
            del self.drafts[key]
        return len(keys)

    def delete_by_package(self, package_id: str) -> int:
        case_ids = {
            draft.case_id
            for draft in self.drafts.values()
            if draft.case_id.startswith(package_id)
        }
        keys = [key for key, draft in self.drafts.items() if draft.case_id in case_ids]
        for key in keys:
            del self.drafts[key]
        return len(keys)


class LtrStore:
    def __init__(self, records: list[LtrRecord] | None = None) -> None:
        self.records = records or []

    def list_by_project(self, project_id: str) -> list[LtrRecord]:
        return [record for record in self.records if record.project_id == project_id]


class FakeEligibilityValidator:
    def __init__(self, result: ApplicationFormEligibility | None = None) -> None:
        self.result = result or ApplicationFormEligibility(
            eligible=True,
            reason_code="ok",
            message="Application form is ready for Precheck.",
            observed_header_cell="Laboratory Testing Request",
        )

    def evaluate(self, asset: IntakeAsset) -> ApplicationFormEligibility:
        return self.result


def _package() -> IntakePackage:
    return IntakePackage(
        package_id="pkg-1",
        source_type=IntakePackageSourceType.OUTLOOK_MSG,
        status=IntakePackageStatus.READY_FOR_REVIEW,
        source_original_name="request.msg",
        source_stored_path=Path("data/intake/pkg-1/source/request.msg"),
    )


def _asset(asset_id: str, role: IntakeAssetRole, extension: str = ".docx") -> IntakeAsset:
    return IntakeAsset(
        asset_id=asset_id,
        package_id="pkg-1",
        original_name=f"{asset_id}{extension}",
        stored_path=Path(f"data/intake/pkg-1/attachments/{asset_id}{extension}"),
        extension=extension,
        mime_type="application/octet-stream",
        size_bytes=100,
        sha256=asset_id * 64,
        asset_role=role,
    )


def _email_asset(
    asset_id: str = "email",
    package_id: str = "pkg-1",
    size: int = 1000,
    *,
    original_name: str = "request.msg",
    sha256: str = "e" * 64,
) -> IntakeAsset:
    return IntakeAsset(
        asset_id=asset_id,
        package_id=package_id,
        original_name=original_name,
        stored_path=Path(f"data/intake/{package_id}/source/{original_name}"),
        extension=".msg",
        mime_type="application/vnd.ms-outlook",
        size_bytes=size,
        sha256=sha256,
        asset_role=IntakeAssetRole.EMAIL_SOURCE,
    )


def _asset_with_path(
    asset_id: str,
    role: IntakeAssetRole,
    stored_path: Path,
) -> IntakeAsset:
    return IntakeAsset(
        asset_id=asset_id,
        package_id="pkg-1",
        original_name=stored_path.name,
        stored_path=stored_path,
        extension=stored_path.suffix,
        mime_type="application/octet-stream",
        size_bytes=stored_path.stat().st_size,
        sha256=asset_id * 64,
        asset_role=role,
    )


def _service(
    assets: list[IntakeAsset],
    cases: list[IntakeCase] | None = None,
    drafts: list[IntakeDraft] | None = None,
    eligibility: ApplicationFormEligibility | None = None,
    packages: list[IntakePackage] | None = None,
    ltrs: list[LtrRecord] | None = None,
) -> tuple[IntakeFormSelectionService, AssetStore, CaseStore, DraftStore]:
    asset_store = AssetStore(assets)
    case_store = CaseStore(cases)
    draft_store = DraftStore(drafts)
    return (
        IntakeFormSelectionService(
            PackageStore(packages or [_package()]),
            asset_store,
            case_store,
            draft_store,
            eligibility_validator=FakeEligibilityValidator(eligibility),
            ltr_store=LtrStore(ltrs),
        ),
        asset_store,
        case_store,
        draft_store,
    )


def test_select_candidate_creates_case_and_empty_draft() -> None:
    service, asset_store, case_store, draft_store = _service(
        [_asset("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE)]
    )

    result = service.select_form_asset("pkg-1", "asset-a")

    assert result.selected_asset.asset_role is IntakeAssetRole.SELECTED_APPLICATION_FORM
    assert asset_store.assets["asset-a"].asset_role is IntakeAssetRole.SELECTED_APPLICATION_FORM
    assert result.case.status is IntakeCaseStatus.NEEDS_REVIEW
    assert result.case.confirmed_project_id is None
    assert result.draft.parsed_fields_json == "{}"
    assert len(case_store.cases) == 1
    assert len(draft_store.drafts) == 1


def test_select_candidate_parses_selected_docx_into_draft(tmp_path: Path) -> None:
    docx_path = tmp_path / "selected-application.docx"
    document = Document()
    table = document.add_table(rows=4, cols=2)
    for row_index, (label, value) in enumerate(
        [
            ("Form No.", "E-3718"),
            ("Requested By", "Alice Requestor"),
            ("Email", "alice@example.com"),
            ("Description of Requested Testing", "Thermal cycling"),
        ]
    ):
        table.cell(row_index, 0).text = label
        table.cell(row_index, 1).text = value
    sample_table = document.add_table(rows=2, cols=4)
    for index, header in enumerate(["Product Name", "Part Number", "Revision", "Quantity"]):
        sample_table.cell(0, index).text = header
    for index, value in enumerate(["Connector A", "PN-073", "A", "12"]):
        sample_table.cell(1, index).text = value
    document.save(docx_path)
    service, _, _, _ = _service(
        [_asset_with_path("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE, docx_path)]
    )

    result = service.select_form_asset("pkg-1", "asset-a")

    assert '"requester": "Alice Requestor"' in result.draft.parsed_fields_json
    assert '"product_name": "Connector A"' in result.draft.parsed_fields_json
    assert '"requested_testing": "Thermal cycling"' in result.draft.parsed_fields_json
    assert result.draft.parser_warnings_json == "[]"


def test_select_word_asset_without_candidate_role_is_allowed_for_human_override() -> None:
    service, _, _, _ = _service([_asset("asset-a", IntakeAssetRole.SUPPORTING_ATTACHMENT)])

    result = service.select_form_asset("pkg-1", "asset-a")

    assert result.selected_asset.asset_role is IntakeAssetRole.SELECTED_APPLICATION_FORM


def test_selection_updates_existing_case_and_draft() -> None:
    existing_case = IntakeCase(
        case_id="case-1",
        package_id="pkg-1",
        selected_form_asset_id="old-asset",
        status=IntakeCaseStatus.DRAFT_CREATED,
    )
    existing_draft = IntakeDraft(
        draft_id="draft-1",
        case_id="case-1",
        parsed_fields_json='{"old":true}',
    )
    service, _, case_store, draft_store = _service(
        [_asset("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE)],
        cases=[existing_case],
        drafts=[existing_draft],
    )

    result = service.select_form_asset("pkg-1", "asset-a")

    assert result.case.case_id != "case-1"
    assert result.case.selected_form_asset_id == "asset-a"
    assert result.case.confirmed_project_id is None
    assert case_store.cases["case-1"].selected_form_asset_id == "old-asset"
    assert draft_store.drafts["draft-1"].parsed_fields_json == '{"old":true}'


def test_selection_preserves_manual_overrides_for_same_selected_asset() -> None:
    existing_case = IntakeCase(
        case_id="case-1",
        package_id="pkg-1",
        selected_form_asset_id="asset-a",
        status=IntakeCaseStatus.NEEDS_REVIEW,
    )
    existing_draft = IntakeDraft(
        draft_id="draft-1",
        case_id="case-1",
        parsed_fields_json='{"requester":"Parsed"}',
        manual_overrides_json='{"requester":"Corrected"}',
    )
    service, _, _, draft_store = _service(
        [_asset("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE)],
        cases=[existing_case],
        drafts=[existing_draft],
    )

    result = service.select_form_asset("pkg-1", "asset-a")

    assert result.case.case_id == "case-1"
    assert draft_store.drafts["draft-1"].manual_overrides_json == '{"requester":"Corrected"}'


def test_selection_replaces_manual_overrides_when_explicitly_requested() -> None:
    """TASK_103 confirmed import replacement clears current editor overrides."""
    existing_case = IntakeCase(
        case_id="case-1",
        package_id="pkg-1",
        selected_form_asset_id="asset-a",
        status=IntakeCaseStatus.NEEDS_REVIEW,
    )
    existing_draft = IntakeDraft(
        draft_id="draft-1",
        case_id="case-1",
        parsed_fields_json='{"requester":"Parsed"}',
        manual_overrides_json='{"requester":"Corrected"}',
    )
    service, _, _, draft_store = _service(
        [_asset("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE)],
        cases=[existing_case],
        drafts=[existing_draft],
    )

    result = service.select_form_asset("pkg-1", "asset-a", replace_existing=True)

    assert result.case.case_id == "case-1"
    assert draft_store.drafts["draft-1"].manual_overrides_json is None


def test_selection_clears_manual_overrides_when_rebinding_reusable_case() -> None:
    reusable_case = IntakeCase(
        case_id="case-1",
        package_id="pkg-1",
        selected_form_asset_id="old-asset",
        status=IntakeCaseStatus.DRAFT_CREATED,
    )
    existing_draft = IntakeDraft(
        draft_id="draft-1",
        case_id="case-1",
        parsed_fields_json='{"requester":"Old"}',
        manual_overrides_json='{"requester":"Old corrected"}',
    )
    service, _, _, draft_store = _service(
        [_asset("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE)],
        cases=[reusable_case],
        drafts=[existing_draft],
    )

    result = service.select_form_asset("pkg-1", "asset-a")

    assert result.case.case_id != "case-1"
    assert result.case.selected_form_asset_id == "asset-a"
    assert draft_store.drafts["draft-1"].manual_overrides_json == '{"requester":"Old corrected"}'


def test_selection_rebinds_unconfirmed_case_when_package_has_other_form_case() -> None:
    existing_case = IntakeCase(
        case_id="case-1",
        package_id="pkg-1",
        selected_form_asset_id="asset-a",
        status=IntakeCaseStatus.NEEDS_REVIEW,
    )
    service, _, case_store, _ = _service(
        [
            _asset("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE),
            _asset("asset-b", IntakeAssetRole.APPLICATION_FORM_CANDIDATE),
        ],
        cases=[existing_case],
    )

    result = service.select_form_asset("pkg-1", "asset-b")

    assert result.case.case_id != "case-1"
    assert result.case.selected_form_asset_id == "asset-b"
    assert len(case_store.cases) == 2
    assert case_store.cases["case-1"].selected_form_asset_id == "asset-a"


def test_selection_creates_new_case_when_existing_case_is_confirmed() -> None:
    confirmed_case = IntakeCase(
        case_id="case-1",
        package_id="pkg-1",
        selected_form_asset_id="asset-a",
        status=IntakeCaseStatus.CONFIRMED,
        confirmed_project_id="project-1",
    )
    service, _, case_store, _ = _service(
        [
            _asset("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE),
            _asset("asset-b", IntakeAssetRole.APPLICATION_FORM_CANDIDATE),
        ],
        cases=[confirmed_case],
    )

    result = service.select_form_asset("pkg-1", "asset-b")

    assert result.case.case_id != "case-1"
    assert result.case.selected_form_asset_id == "asset-b"
    assert len(case_store.cases) == 2
    assert case_store.cases["case-1"].confirmed_project_id == "project-1"


def test_selection_requires_resolution_for_existing_application_draft() -> None:
    from backend.application.intake_form_selection_service import (
        IntakeDraftDuplicateResolutionRequiredError,
    )

    existing_package = replace(_package(), package_id="pkg-existing")
    incoming_package = _package()
    existing_form = replace(
        _asset("asset-existing", IntakeAssetRole.SELECTED_APPLICATION_FORM),
        package_id="pkg-existing",
        original_name="application.docx",
    )
    incoming_form = replace(
        _asset("asset-incoming", IntakeAssetRole.APPLICATION_FORM_CANDIDATE),
        original_name="application.docx",
    )
    existing_case = IntakeCase(
        case_id="case-existing",
        package_id="pkg-existing",
        selected_form_asset_id="asset-existing",
        status=IntakeCaseStatus.NEEDS_REVIEW,
    )
    service, _, _, _ = _service(
        [
            _email_asset("email-existing", "pkg-existing", 2048),
            _email_asset("email-incoming", "pkg-1", 2048),
            existing_form,
            incoming_form,
        ],
        cases=[existing_case],
        drafts=[
            IntakeDraft(
                draft_id="draft-existing",
                case_id="case-existing",
                parsed_fields_json="{}",
            )
        ],
        packages=[existing_package, incoming_package],
    )

    with pytest.raises(IntakeDraftDuplicateResolutionRequiredError) as exc_info:
        service.select_form_asset("pkg-1", "asset-incoming")

    check = exc_info.value.check
    assert check.classification == "exact_existing_application_draft"
    assert check.existing_case_id == "case-existing"
    assert check.incoming_application_form_name == "application.docx"


def test_selection_duplicate_hides_replace_when_existing_package_has_confirmed_case() -> None:
    from backend.application.intake_form_selection_service import (
        IntakeDraftDuplicateResolutionRequiredError,
    )

    existing_package = replace(_package(), package_id="pkg-existing")
    incoming_package = _package()
    existing_form = replace(
        _asset("asset-existing", IntakeAssetRole.SELECTED_APPLICATION_FORM),
        package_id="pkg-existing",
        original_name="application.docx",
    )
    incoming_form = replace(
        _asset("asset-incoming", IntakeAssetRole.APPLICATION_FORM_CANDIDATE),
        original_name="application.docx",
    )
    reusable_case = IntakeCase(
        case_id="case-existing",
        package_id="pkg-existing",
        selected_form_asset_id="asset-existing",
        status=IntakeCaseStatus.NEEDS_REVIEW,
    )
    confirmed_case = IntakeCase(
        case_id="case-confirmed",
        package_id="pkg-existing",
        selected_form_asset_id=None,
        status=IntakeCaseStatus.CONFIRMED,
        confirmed_project_id="project-existing",
    )
    service, _, _, _ = _service(
        [
            _email_asset("email-existing", "pkg-existing", 2048),
            _email_asset("email-incoming", "pkg-1", 2048),
            existing_form,
            incoming_form,
        ],
        cases=[reusable_case, confirmed_case],
        drafts=[
            IntakeDraft(
                draft_id="draft-existing",
                case_id="case-existing",
                parsed_fields_json="{}",
            )
        ],
        packages=[existing_package, incoming_package],
    )

    with pytest.raises(IntakeDraftDuplicateResolutionRequiredError) as exc_info:
        service.select_form_asset("pkg-1", "asset-incoming")

    assert exc_info.value.check.allowed_actions == ("open_existing",)


def test_selection_duplicate_uses_email_hash_not_display_filename() -> None:
    from backend.application.intake_form_selection_service import (
        IntakeDraftDuplicateResolutionRequiredError,
    )

    existing_package = replace(_package(), package_id="pkg-existing")
    incoming_package = replace(
        _package(),
        source_original_name="连接器主板对busbar对接测试副本.msg",
    )
    existing_form = replace(
        _asset("asset-existing", IntakeAssetRole.SELECTED_APPLICATION_FORM),
        package_id="pkg-existing",
        original_name="application.docx",
    )
    incoming_form = replace(
        _asset("asset-incoming", IntakeAssetRole.APPLICATION_FORM_CANDIDATE),
        original_name="application.docx",
    )
    service, _, _, _ = _service(
        [
            _email_asset(
                "email-existing",
                "pkg-existing",
                2048,
                original_name="_busbar_.msg",
                sha256="5" * 64,
            ),
            _email_asset(
                "email-incoming",
                "pkg-1",
                2048,
                original_name="连接器主板对busbar对接测试副本.msg",
                sha256="5" * 64,
            ),
            existing_form,
            incoming_form,
        ],
        cases=[
            IntakeCase(
                case_id="case-existing",
                package_id="pkg-existing",
                selected_form_asset_id="asset-existing",
                status=IntakeCaseStatus.NEEDS_REVIEW,
            )
        ],
        drafts=[
            IntakeDraft(
                draft_id="draft-existing",
                case_id="case-existing",
                parsed_fields_json="{}",
            )
        ],
        packages=[existing_package, incoming_package],
    )

    with pytest.raises(IntakeDraftDuplicateResolutionRequiredError) as exc_info:
        service.select_form_asset("pkg-1", "asset-incoming")

    check = exc_info.value.check
    assert check.existing_source_original_name == "_busbar_.msg"
    assert check.incoming_source_original_name == "连接器主板对busbar对接测试副本.msg"


def test_selection_allows_same_source_when_existing_project_already_confirmed() -> None:
    existing_package = replace(_package(), package_id="pkg-existing")
    incoming_package = replace(
        _package(),
        source_original_name="连接器主板对busbar对接测试副本.msg",
    )
    existing_form = replace(
        _asset("asset-existing", IntakeAssetRole.SELECTED_APPLICATION_FORM),
        package_id="pkg-existing",
        original_name="application.docx",
    )
    incoming_form = replace(
        _asset("asset-incoming", IntakeAssetRole.APPLICATION_FORM_CANDIDATE),
        original_name="application.docx",
    )
    service, _, case_store, draft_store = _service(
        [
            _email_asset(
                "email-existing",
                "pkg-existing",
                2048,
                original_name="_busbar_.msg",
                sha256="5" * 64,
            ),
            _email_asset(
                "email-incoming",
                "pkg-1",
                2048,
                original_name="连接器主板对busbar对接测试副本.msg",
                sha256="5" * 64,
            ),
            existing_form,
            incoming_form,
        ],
        cases=[
            IntakeCase(
                case_id="case-existing",
                package_id="pkg-existing",
                selected_form_asset_id="asset-existing",
                status=IntakeCaseStatus.CONFIRMED,
                confirmed_project_id="project-existing",
            )
        ],
        packages=[existing_package, incoming_package],
        ltrs=[
            LtrRecord(
                ltr_id="ltr-existing",
                project_id="project-existing",
                ltr_number="DL-2026-05-001",
                status=LtrStatus.REGISTERED,
            )
        ],
    )

    result = service.select_form_asset("pkg-1", "asset-incoming")

    assert result.case.package_id == "pkg-1"
    assert result.case.selected_form_asset_id == "asset-incoming"
    assert result.case.case_id != "case-existing"
    assert len(case_store.cases) == 2
    assert len(draft_store.drafts) == 1


def test_selection_returns_existing_case_when_reselecting_same_asset() -> None:
    """When user re-clicks an already selected asset, return existing case directly.

    This prevents unnecessary duplicate confirmation dialogs (Phase 1.2 fix).
    """
    asset_a = _asset("asset-a", IntakeAssetRole.SELECTED_APPLICATION_FORM)
    asset_b = _asset("asset-b", IntakeAssetRole.SELECTED_APPLICATION_FORM)
    case_a = IntakeCase(
        case_id="case-a",
        package_id="pkg-1",
        selected_form_asset_id="asset-a",
        status=IntakeCaseStatus.NEEDS_REVIEW,
    )
    case_b = IntakeCase(
        case_id="case-b",
        package_id="pkg-1",
        selected_form_asset_id="asset-b",
        status=IntakeCaseStatus.NEEDS_REVIEW,
    )
    draft_a = IntakeDraft(
        draft_id="draft-a",
        case_id="case-a",
        parsed_fields_json="{}",
        manual_overrides_json='{"test_item":"Edited"}',
    )
    service, _, _, _ = _service(
        [_email_asset(), asset_a, asset_b],
        cases=[case_a, case_b],
        drafts=[
            draft_a,
            IntakeDraft(
                draft_id="draft-b",
                case_id="case-b",
                parsed_fields_json="{}",
            ),
        ],
    )

    # Re-selecting asset-a should return existing case-a directly, no exception
    result = service.select_form_asset("pkg-1", "asset-a")

    assert result.case.case_id == "case-a"
    assert result.draft.draft_id == "draft-a"
    assert result.selected_asset.asset_id == "asset-a"


def test_selection_returns_existing_case_for_single_form_repeat() -> None:
    """When user repeats selection of the only form, return existing case directly.

    This prevents unnecessary duplicate confirmation dialogs (Phase 1.2 fix).
    """
    service, _, _, _ = _service(
        [
            _email_asset(),
            _asset("asset-a", IntakeAssetRole.SELECTED_APPLICATION_FORM),
        ],
        cases=[
            IntakeCase(
                case_id="case-a",
                package_id="pkg-1",
                selected_form_asset_id="asset-a",
                status=IntakeCaseStatus.NEEDS_REVIEW,
            )
        ],
        drafts=[
            IntakeDraft(
                draft_id="draft-a",
                case_id="case-a",
                parsed_fields_json="{}",
            )
        ],
    )

    # Re-selecting the same asset should return existing case directly, no exception
    result = service.select_form_asset("pkg-1", "asset-a")

    assert result.case.case_id == "case-a"
    assert result.draft.draft_id == "draft-a"
    assert result.selected_asset.asset_id == "asset-a"


def test_selection_reinitialize_same_package_form_clears_manual_overrides() -> None:
    asset_a = _asset("asset-a", IntakeAssetRole.SELECTED_APPLICATION_FORM)
    asset_b = _asset("asset-b", IntakeAssetRole.SELECTED_APPLICATION_FORM)
    service, _, _, draft_store = _service(
        [_email_asset(), asset_a, asset_b],
        cases=[
            IntakeCase(
                case_id="case-a",
                package_id="pkg-1",
                selected_form_asset_id="asset-a",
                status=IntakeCaseStatus.NEEDS_REVIEW,
            ),
            IntakeCase(
                case_id="case-b",
                package_id="pkg-1",
                selected_form_asset_id="asset-b",
                status=IntakeCaseStatus.NEEDS_REVIEW,
            ),
        ],
        drafts=[
            IntakeDraft(
                draft_id="draft-a",
                case_id="case-a",
                parsed_fields_json="{}",
                manual_overrides_json='{"test_item":"Edited"}',
            ),
            IntakeDraft(
                draft_id="draft-b",
                case_id="case-b",
                parsed_fields_json="{}",
            ),
        ],
    )

    result = service.select_form_asset(
        "pkg-1",
        "asset-a",
        resolution_action="replace_existing",
        resolution_case_id="case-a",
    )

    assert result.case.case_id == "case-a"
    rebuilt = draft_store.get_by_case("case-a")
    assert rebuilt is not None
    assert rebuilt.manual_overrides_json is None
    assert rebuilt.draft_id != "draft-a"
    assert "draft-a" not in draft_store.drafts


def test_selection_creates_separate_case_for_same_email_different_form() -> None:
    existing_package = replace(_package(), package_id="pkg-existing")
    incoming_package = _package()
    existing_form = replace(
        _asset("asset-existing", IntakeAssetRole.SELECTED_APPLICATION_FORM),
        package_id="pkg-existing",
        original_name="old-application.docx",
    )
    incoming_form = replace(
        _asset("asset-incoming", IntakeAssetRole.APPLICATION_FORM_CANDIDATE),
        original_name="new-application.docx",
    )
    existing_case = IntakeCase(
        case_id="case-existing",
        package_id="pkg-existing",
        selected_form_asset_id="asset-existing",
        status=IntakeCaseStatus.NEEDS_REVIEW,
    )
    service, _, case_store, _ = _service(
        [
            _email_asset("email-existing", "pkg-existing", 2048),
            _email_asset("email-incoming", "pkg-1", 2048),
            existing_form,
            incoming_form,
        ],
        cases=[existing_case],
        packages=[existing_package, incoming_package],
    )

    result = service.select_form_asset("pkg-1", "asset-incoming")

    assert result.case.package_id == "pkg-1"
    assert result.case.case_id != "case-existing"
    assert len(case_store.cases) == 2


def test_selection_rejects_non_word_non_candidate_asset() -> None:
    service, _, _, _ = _service([_asset("asset-a", IntakeAssetRole.UNKNOWN, ".pdf")])

    with pytest.raises(IntakeSelectionError):
        service.select_form_asset("pkg-1", "asset-a")


def test_selection_rejects_doc_asset_even_when_candidate() -> None:
    service, _, _, _ = _service(
        [_asset("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE, ".doc")],
        eligibility=ApplicationFormEligibility(
            eligible=False,
            reason_code="not_docx",
            message="Select a .docx Laboratory Testing Request form to continue.",
        ),
    )

    with pytest.raises(IntakeSelectionError) as exc_info:
        service.select_form_asset("pkg-1", "asset-a")

    assert "Select a .docx Laboratory Testing Request form" in str(exc_info.value)


def test_selection_rejects_docx_when_header_gate_fails() -> None:
    service, _, _, _ = _service(
        [_asset("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE)],
        eligibility=ApplicationFormEligibility(
            eligible=False,
            reason_code="header_cell_mismatch",
            message=(
                "Selected document is not recognized as Laboratory Testing Request. "
                'Header table cell (1,2): "Connector Test Request"'
            ),
            observed_header_cell="Connector Test Request",
        ),
    )

    with pytest.raises(IntakeSelectionError) as exc_info:
        service.select_form_asset("pkg-1", "asset-a")

    assert 'Header table cell (1,2): "Connector Test Request"' in str(exc_info.value)


def test_selection_rejects_missing_package_or_asset() -> None:
    service, _, _, _ = _service([_asset("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE)])

    with pytest.raises(IntakeSelectionNotFoundError):
        service.select_form_asset("missing", "asset-a")
    with pytest.raises(IntakeSelectionNotFoundError):
        service.select_form_asset("pkg-1", "missing")


def test_selection_rejects_asset_from_another_package() -> None:
    foreign_asset = replace(
        _asset("asset-a", IntakeAssetRole.APPLICATION_FORM_CANDIDATE),
        package_id="pkg-2",
    )
    service, _, _, _ = _service([foreign_asset])

    with pytest.raises(IntakeSelectionError):
        service.select_form_asset("pkg-1", "asset-a")
