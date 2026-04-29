"""Generated DOCX fixtures that mirror real Phase 7 application form layouts."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def build_real_style_application_form(path: Path, *, tester_modified: bool = False) -> Path:
    """Create a sanitized real-style application form fixture."""
    document = Document()
    _add_header(document)
    document.add_paragraph("Laboratory Test Request")
    _add_request_table(document, tester_modified=tester_modified)
    _add_sample_table(document, tester_modified=tester_modified)
    _add_lab_table(document, tester_modified=tester_modified)
    _add_footer(document)
    document.save(path)
    return path


def _add_header(document: Document) -> None:
    """Add a minimal real-style document header."""
    header = document.sections[0].header
    table = header.add_table(rows=1, cols=1, width=0)
    table.cell(0, 0).text = "Connector Laboratory Test Request"


def _add_footer(document: Document) -> None:
    """Add footer metadata in the compact format seen in real forms."""
    footer = document.sections[0].footer
    footer.paragraphs[0].text = "Form No. E-3718 Rev H"


def _add_request_table(document: Document, *, tester_modified: bool) -> None:
    """Add real-style request rows with multiple label/value pairs per row."""
    rows = [
        ["Requested By", "Alice Requestor", "Phone", "555-0100", "Date", "2026-04-27"],
        [
            "Email",
            "alice.requestor@example.com",
            "Business Unit",
            "Industrial",
            "Manufacturing Site",
            "Plant 7",
        ],
        [
            "Project #",
            "PRJ-038-T" if tester_modified else "PRJ-038-A",
            "Project Type",
            "Qualification",
            "Subcontract",
            "No",
        ],
        [
            "Requested Testing",
            "Thermal cycling and contact resistance",
            "Test Type",
            "Validation",
            "Sample Status",
            "Prototype",
        ],
    ]
    table = document.add_table(rows=len(rows), cols=6)
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            table.cell(row_index, col_index).text = value


def _add_sample_table(document: Document, *, tester_modified: bool) -> None:
    """Add a sample table that uses the real multi-column shape."""
    table = document.add_table(rows=2, cols=8)
    headers = [
        "Product Name",
        "Part Number",
        "Rev",
        "Lot Traceability",
        "Material",
        "Plating",
        "Housing Material",
        "Qty",
    ]
    values = [
        "Synthetic Connector",
        "PN-038-T" if tester_modified else "PN-038-A",
        "B" if tester_modified else "A",
        "LOT-038",
        "Copper Alloy",
        "Tin",
        "LCP",
        "24",
    ]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
        table.cell(1, index).text = values[index]


def _add_lab_table(document: Document, *, tester_modified: bool) -> None:
    """Add a real-style lab section table."""
    rows = [
        ["Lab", "Connector Lab", "Assigned Personnel", "Charlie Tester"],
        ["Received Date", "2026-04-28", "Estimated Completion Date", "2026-05-15"],
        ["Sample Condition", "Good" if not tester_modified else "Good, tester reviewed", "", ""],
    ]
    table = document.add_table(rows=len(rows), cols=4)
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            table.cell(row_index, col_index).text = value
