from __future__ import annotations

from collections.abc import Generator
from pathlib import Path

from fastapi.testclient import TestClient
from docx import Document
from docx.shared import Inches
from sqlalchemy.orm import Session

from backend.api.dependencies import get_session, get_settings
from backend.api.main import app
from backend.infrastructure.storage.database import (
    create_database_engine,
    create_session_factory,
    init_db,
)
from backend.infrastructure.storage.repositories.intake_package import (
    IntakeAssetRepository,
    IntakeCaseRepository,
    IntakeDraftRepository,
    IntakePackageRepository,
)
from backend.shared.config import Settings
from backend.domain import (
    IntakeAsset,
    IntakeAssetRole,
    IntakePackage,
    IntakePackageSourceType,
    IntakePackageStatus,
)


def test_msg_package_import_api_persists_package_and_assets(tmp_path: Path) -> None:
    """The API imports one manual `.msg` package through application service boundaries."""
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        response = client.post(
            "/api/intake-packages/import-msg",
            files={
                "file": (
                    "request.msg",
                    _msg_bytes(
                        [
                            "Subject: Connector qualification request",
                            "From: Jane Engineer <jane@example.com>",
                            "Attachment: E-3718 Application Form.docx; content=docx bytes",
                            "Attachment: drawing.pdf; content=pdf bytes",
                        ]
                    ),
                    "application/vnd.ms-outlook",
                )
            },
        )

        assert response.status_code == 201
        payload = response.json()
        assert payload["source_type"] == "outlook_msg"
        assert payload["package_status"] == "ready_for_review"
        assert payload["subject"] == "Connector qualification request"
        assert payload["sender_email"] == "jane@example.com"
        assert payload["asset_count"] == 3
        assert payload["candidate_count"] == 1
        assert payload["next_action"] == "review_application_form_candidates"

        detail_response = client.get(f"/api/intake-packages/{payload['package_id']}")
        assert detail_response.status_code == 200
        detail = detail_response.json()
        assert detail["package_id"] == payload["package_id"]
        assert detail["source_stored"] is True
        assert detail["asset_count"] == 3
        assert detail["candidate_count"] == 1
        assert detail["case_count"] == 0
        assert detail["next_action"] == "create_review_cases"
        assert [asset["original_name"] for asset in detail["candidate_assets"]] == [
            "E-3718 Application Form.docx"
        ]

        with session_factory() as session:
            package = IntakePackageRepository(session).get(payload["package_id"])
            assets = IntakeAssetRepository(session).list_by_package(payload["package_id"])
            assert package is not None
            assert package.source_stored_path.is_file()
            assert len(assets) == 3
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_intake_package_detail_api_returns_404_for_missing_package(
    tmp_path: Path,
) -> None:
    """The package detail endpoint reports missing package ids clearly."""
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        response = client.get("/api/intake-packages/missing")

        assert response.status_code == 404
        assert "Intake package not found" in response.json()["detail"]
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_select_form_api_binds_selected_docx_to_precheck_case(tmp_path: Path) -> None:
    """Selecting one Word asset creates the case opened by Precheck with parsed fields."""
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        docx_path = _create_application_docx(tmp_path / "application.docx")
        with session_factory() as session:
            package_repo = IntakePackageRepository(session)
            asset_repo = IntakeAssetRepository(session)
            package_repo.create(
                IntakePackage(
                    package_id="pkg-select",
                    source_type=IntakePackageSourceType.OUTLOOK_MSG,
                    status=IntakePackageStatus.READY_FOR_REVIEW,
                    source_original_name="request.msg",
                    source_stored_path=tmp_path / "request.msg",
                )
            )
            asset_repo.create(
                IntakeAsset(
                    asset_id="asset-selected",
                    package_id="pkg-select",
                    original_name=docx_path.name,
                    stored_path=docx_path,
                    extension=".docx",
                    mime_type="application/octet-stream",
                    size_bytes=docx_path.stat().st_size,
                    sha256="a" * 64,
                    asset_role=IntakeAssetRole.APPLICATION_FORM_CANDIDATE,
                    candidate_score=90,
                )
            )
            asset_repo.create(
                IntakeAsset(
                    asset_id="asset-other",
                    package_id="pkg-select",
                    original_name="other.docx",
                    stored_path=docx_path,
                    extension=".docx",
                    mime_type="application/octet-stream",
                    size_bytes=docx_path.stat().st_size,
                    sha256="b" * 64,
                    asset_role=IntakeAssetRole.APPLICATION_FORM_CANDIDATE,
                    candidate_score=80,
                )
            )
            session.commit()

        response = client.post(
            "/api/intake-packages/pkg-select/select-form",
            json={"asset_id": "asset-selected"},
        )

        assert response.status_code == 200
        selected = response.json()
        assert selected["selected_form_asset_id"] == "asset-selected"
        assert selected["case_id"]

        review_response = client.get("/api/intake-packages/pkg-select/case-review")
        assert review_response.status_code == 200
        cases = review_response.json()["cases"]
        assert len(cases) == 1
        assert cases[0]["case_id"] == selected["case_id"]
        assert cases[0]["selected_form_asset_id"] == "asset-selected"
        field_values = {field["key"]: field["value"] for field in cases[0]["fields"]}
        assert field_values["requester"] == "Alice Requestor"
        assert field_values["product_name"] == "Connector A"
        assert field_values["form_no"] == "E-3718"
        assert cases[0]["sample_rows"][0]["product_name"] == "Connector A"
        assert cases[0]["sample_rows"][0]["part_number"] == "PN-073"

        with session_factory() as session:
            assert (
                IntakeCaseRepository(session).list_by_package("pkg-select")[0].selected_form_asset_id
                == "asset-selected"
            )
            draft = IntakeDraftRepository(session).get_by_case(selected["case_id"])
            assert draft is not None
            assert "Alice Requestor" in draft.parsed_fields_json

        second_response = client.post(
            "/api/intake-packages/pkg-select/select-form",
            json={"asset_id": "asset-other"},
        )

        assert second_response.status_code == 200
        second = second_response.json()
        assert second["case_id"] == selected["case_id"]
        assert second["selected_form_asset_id"] == "asset-other"

        second_review_response = client.get("/api/intake-packages/pkg-select/case-review")
        assert second_review_response.status_code == 200
        second_cases = second_review_response.json()["cases"]
        assert len(second_cases) == 1
        assert second_cases[0]["case_id"] == selected["case_id"]
        assert second_cases[0]["selected_form_asset_id"] == "asset-other"

        with session_factory() as session:
            persisted_cases = IntakeCaseRepository(session).list_by_package("pkg-select")
            assert len(persisted_cases) == 1
            assert persisted_cases[0].selected_form_asset_id == "asset-other"
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_email_package_without_form_accepts_supplemental_application_form(
    tmp_path: Path,
) -> None:
    """A no-form email package can continue after uploading a Word form into it."""
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        import_response = client.post(
            "/api/intake-packages/import-msg",
            files={
                "file": (
                    "request.msg",
                    _msg_bytes(
                        [
                            "Subject: Connector qualification request",
                            "From: Jane Engineer <jane@example.com>",
                            "Attachment: drawing.pdf; content=pdf bytes",
                        ]
                    ),
                    "application/vnd.ms-outlook",
                )
            },
        )
        assert import_response.status_code == 201
        imported = import_response.json()
        assert imported["package_status"] == "needs_application_form_selection"
        assert imported["candidate_count"] == 0
        assert imported["next_action"] == "resolve_missing_application_form"

        docx_path = _create_application_docx(tmp_path / "supplemental-form.docx")
        with docx_path.open("rb") as handle:
            supplemental_response = client.post(
                f"/api/intake-packages/{imported['package_id']}/application-form",
                files={
                    "file": (
                        docx_path.name,
                        handle,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        assert supplemental_response.status_code == 201
        selected = supplemental_response.json()
        assert selected["package_id"] == imported["package_id"]
        assert selected["selected_form_asset_id"]
        assert selected["case_id"]
        assert selected["next_action"] == "review_selected_application_form"

        detail_response = client.get(f"/api/intake-packages/{imported['package_id']}")
        assert detail_response.status_code == 200
        detail = detail_response.json()
        assert detail["source_type"] == "outlook_msg"
        assert detail["source_original_name"] == "request.msg"
        assert detail["subject"] == "Connector qualification request"
        assert detail["sender_email"] == "jane@example.com"
        assert detail["asset_count"] == 3
        assert detail["case_count"] == 1
        assert selected["selected_form_asset_id"] in [
            asset["asset_id"] for asset in detail["assets"]
        ]
        assert any(asset["asset_role"] == "email_source" for asset in detail["assets"])
        assert any(asset["original_name"] == "drawing.pdf" for asset in detail["assets"])
        supplemental_asset = next(
            asset
            for asset in detail["assets"]
            if asset["asset_id"] == selected["selected_form_asset_id"]
        )
        assert supplemental_asset["original_name"] == "supplemental-form.docx"
        assert supplemental_asset["asset_role"] == "selected_application_form"

        review_response = client.get(
            f"/api/intake-packages/{imported['package_id']}/case-review"
        )
        assert review_response.status_code == 200
        cases = review_response.json()["cases"]
        assert len(cases) == 1
        assert cases[0]["case_id"] == selected["case_id"]
        assert cases[0]["selected_form_asset_id"] == selected["selected_form_asset_id"]

        with session_factory() as session:
            assets = IntakeAssetRepository(session).list_by_package(imported["package_id"])
            selected_asset = next(
                asset for asset in assets if asset.asset_id == selected["selected_form_asset_id"]
            )
            assert settings.data_dir / "intake" in selected_asset.stored_path.parents
            assert selected_asset.stored_path.is_file()
            assert IntakeDraftRepository(session).get_by_case(selected["case_id"]) is not None
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_select_form_rejects_docx_with_mismatched_header_gate(tmp_path: Path) -> None:
    """Selected-form API blocks docx files whose header marker does not match."""
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        docx_path = _create_application_docx(
            tmp_path / "wrong-form.docx",
            header_text="Connector Test Request",
        )
        with session_factory() as session:
            package_repo = IntakePackageRepository(session)
            asset_repo = IntakeAssetRepository(session)
            package_repo.create(
                IntakePackage(
                    package_id="pkg-bad-header",
                    source_type=IntakePackageSourceType.OUTLOOK_MSG,
                    status=IntakePackageStatus.READY_FOR_REVIEW,
                    source_original_name="request.msg",
                    source_stored_path=tmp_path / "request.msg",
                )
            )
            asset_repo.create(
                IntakeAsset(
                    asset_id="asset-bad-header",
                    package_id="pkg-bad-header",
                    original_name=docx_path.name,
                    stored_path=docx_path,
                    extension=".docx",
                    mime_type="application/octet-stream",
                    size_bytes=docx_path.stat().st_size,
                    sha256="c" * 64,
                    asset_role=IntakeAssetRole.APPLICATION_FORM_CANDIDATE,
                    candidate_score=90,
                )
            )
            session.commit()

        validate_response = client.post(
            "/api/intake-assets/asset-bad-header/application-form/validate"
        )
        assert validate_response.status_code == 200
        validation = validate_response.json()
        assert validation["eligible"] is False
        assert validation["reason_code"] == "header_cell_mismatch"
        assert validation["observed_header_cell"] == "Connector Test Request"

        select_response = client.post(
            "/api/intake-packages/pkg-bad-header/select-form",
            json={"asset_id": "asset-bad-header"},
        )
        assert select_response.status_code == 400
        assert 'Header table cell (1,2): "Connector Test Request"' in (
            select_response.json()["detail"]
        )
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_email_package_supplemental_application_form_rejects_non_word(
    tmp_path: Path,
) -> None:
    """The supplemental path rejects non-Word files without creating a case."""
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        with session_factory() as session:
            IntakePackageRepository(session).create(
                IntakePackage(
                    package_id="pkg-no-form",
                    source_type=IntakePackageSourceType.OUTLOOK_MSG,
                    status=IntakePackageStatus.NEEDS_APPLICATION_FORM_SELECTION,
                    source_original_name="request.msg",
                    source_stored_path=tmp_path / "request.msg",
                )
            )
            session.commit()

        response = client.post(
            "/api/intake-packages/pkg-no-form/application-form",
            files={"file": ("drawing.pdf", b"%PDF-1.4", "application/pdf")},
        )

        assert response.status_code == 400
        assert "accepts only .docx files" in response.json()["detail"]
        with session_factory() as session:
            assert IntakeCaseRepository(session).list_by_package("pkg-no-form") == []
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_email_package_supplemental_application_form_rejects_bad_header(
    tmp_path: Path,
) -> None:
    """Supplemental `.docx` uploads return a business error when the header gate fails."""
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        with session_factory() as session:
            IntakePackageRepository(session).create(
                IntakePackage(
                    package_id="pkg-bad-supplemental-form",
                    source_type=IntakePackageSourceType.OUTLOOK_MSG,
                    status=IntakePackageStatus.NEEDS_APPLICATION_FORM_SELECTION,
                    source_original_name="request.msg",
                    source_stored_path=tmp_path / "request.msg",
                )
            )
            session.commit()

        docx_path = _create_application_docx(
            tmp_path / "wrong-form.docx",
            header_text="Connector Test Request",
        )
        with docx_path.open("rb") as handle:
            response = client.post(
                "/api/intake-packages/pkg-bad-supplemental-form/application-form",
                files={
                    "file": (
                        docx_path.name,
                        handle,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        assert response.status_code == 400
        assert 'Header table cell (1,2): "Connector Test Request"' in (
            response.json()["detail"]
        )
        with session_factory() as session:
            assert IntakeCaseRepository(session).list_by_package(
                "pkg-bad-supplemental-form"
            ) == []
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_intake_asset_preview_api_returns_docx_preview_without_paths(
    tmp_path: Path,
) -> None:
    """The preview endpoint returns structured DOCX content without local paths."""
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        docx_path = _create_application_docx(tmp_path / "preview-application.docx")
        pdf_path = tmp_path / "drawing.pdf"
        pdf_path.write_bytes(b"%PDF-1.4")
        image_path = tmp_path / "photo.png"
        image_path.write_bytes(b"\x89PNG\r\n\x1a\n")
        with session_factory() as session:
            package_repo = IntakePackageRepository(session)
            asset_repo = IntakeAssetRepository(session)
            package_repo.create(
                IntakePackage(
                    package_id="pkg-preview",
                    source_type=IntakePackageSourceType.OUTLOOK_MSG,
                    status=IntakePackageStatus.READY_FOR_REVIEW,
                    source_original_name="request.msg",
                    source_stored_path=tmp_path / "request.msg",
                )
            )
            asset_repo.create(
                IntakeAsset(
                    asset_id="asset-docx-preview",
                    package_id="pkg-preview",
                    original_name=docx_path.name,
                    stored_path=docx_path,
                    extension=".docx",
                    mime_type="application/octet-stream",
                    size_bytes=docx_path.stat().st_size,
                    sha256="c" * 64,
                    asset_role=IntakeAssetRole.APPLICATION_FORM_CANDIDATE,
                    candidate_score=90,
                )
            )
            asset_repo.create(
                IntakeAsset(
                    asset_id="asset-pdf-preview",
                    package_id="pkg-preview",
                    original_name=pdf_path.name,
                    stored_path=pdf_path,
                    extension=".pdf",
                    mime_type="application/pdf",
                    size_bytes=pdf_path.stat().st_size,
                    sha256="d" * 64,
                    asset_role=IntakeAssetRole.SUPPORTING_ATTACHMENT,
                )
            )
            asset_repo.create(
                IntakeAsset(
                    asset_id="asset-image-preview",
                    package_id="pkg-preview",
                    original_name=image_path.name,
                    stored_path=image_path,
                    extension=".png",
                    mime_type="image/png",
                    size_bytes=image_path.stat().st_size,
                    sha256="e" * 64,
                    asset_role=IntakeAssetRole.SUPPORTING_ATTACHMENT,
                )
            )
            session.commit()

        response = client.get("/api/intake-assets/asset-docx-preview/preview")

        assert response.status_code == 200
        payload = response.json()
        assert payload["kind"] == "docx_application_form"
        assert payload["metadata"]["original_name"] == "preview-application.docx"
        assert "stored_path" not in str(payload)
        fields = {field["label"]: field["value"] for field in payload["fields"]}
        assert fields["Requested By"] == "Alice Requestor"
        assert any(table["title"] == "Test Sample Information" for table in payload["tables"])

        unsupported = client.get("/api/intake-assets/asset-pdf-preview/preview")
        assert unsupported.status_code == 200
        assert unsupported.json()["kind"] == "metadata_only"
        assert unsupported.json()["fields"][1]["value"] == "PDF"

        image = client.get("/api/intake-assets/asset-image-preview/preview")
        assert image.status_code == 200
        image_payload = image.json()
        assert image_payload["kind"] == "image"
        assert image_payload["image_data_url"].startswith("data:image/png;base64,")

        missing = client.get("/api/intake-assets/missing/preview")
        assert missing.status_code == 404
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_intake_asset_download_api_returns_file_with_original_name(
    tmp_path: Path,
) -> None:
    """The download endpoint returns a stored file with the original name."""
    file_content = b"test file content for download"
    stored_file = tmp_path / "original-application.docx"
    stored_file.write_bytes(file_content)
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        with session_factory() as session:
            asset_repo = IntakeAssetRepository(session)
            asset_repo.create(
                IntakeAsset(
                    asset_id="asset-download",
                    package_id="pkg-download",
                    original_name="Original Application Form.docx",
                    stored_path=stored_file,
                    extension=".docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    size_bytes=len(file_content),
                    sha256="f" * 64,
                    asset_role=IntakeAssetRole.APPLICATION_FORM_CANDIDATE,
                    candidate_score=90,
                )
            )
            session.commit()

        response = client.get("/api/intake-assets/asset-download/download")

        assert response.status_code == 200
        assert response.content == file_content
        content_disposition = response.headers["content-disposition"].lower()
        assert "filename" in content_disposition
        assert "original" in content_disposition
        assert "application" in content_disposition
        assert ".docx" in content_disposition

        missing_response = client.get("/api/intake-assets/missing/download")
        assert missing_response.status_code == 404
        assert "Intake asset not found" in missing_response.json()["detail"]

        with session_factory() as session:
            asset_repo = IntakeAssetRepository(session)
            asset = asset_repo.get("asset-download")
            assert asset is not None
            stored_file.unlink()
            session.commit()

        missing_file_response = client.get(
            "/api/intake-assets/asset-download/download"
        )
        assert missing_file_response.status_code == 400
        assert "Stored intake asset file is missing" in missing_file_response.json()[
            "detail"
        ]
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_intake_asset_download_api_for_msg_uses_binary_download_headers(
    tmp_path: Path,
) -> None:
    """The download endpoint forces `.msg` files to download as binary files."""
    file_content = b"test msg file content"
    stored_file = tmp_path / "source.msg"
    stored_file.write_bytes(file_content)
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        with session_factory() as session:
            asset_repo = IntakeAssetRepository(session)
            asset_repo.create(
                IntakeAsset(
                    asset_id="asset-msg-download",
                    package_id="pkg-msg-download",
                    original_name="Original Source.msg",
                    stored_path=stored_file,
                    extension=".msg",
                    mime_type="application/vnd.ms-outlook",
                    size_bytes=len(file_content),
                    sha256="f" * 64,
                    asset_role=IntakeAssetRole.EMAIL_SOURCE,
                )
            )
            session.commit()

        response = client.get("/api/intake-assets/asset-msg-download/download")

        assert response.status_code == 200
        assert response.content == file_content
        assert response.headers["content-type"].lower().startswith("application/octet-stream")
        content_disposition = response.headers["content-disposition"].lower()
        assert "filename" in content_disposition
        assert "filename*" in content_disposition
        assert "source.msg" in content_disposition
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_msg_package_import_api_rejects_non_msg_upload(tmp_path: Path) -> None:
    """The API rejects non-msg files before intake records are created."""
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        response = client.post(
            "/api/intake-packages/import-msg",
            files={
                "file": (
                    "request.docx",
                    b"word",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        assert response.status_code == 400
        assert "accepts only .msg files" in response.json()["detail"]
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def _msg_bytes(lines: list[str]) -> bytes:
    """Build a fixture-style `.msg` byte stream."""
    return "\n".join(lines).encode("utf-8")


def _create_application_docx(
    path: Path,
    header_text: str = "Laboratory Testing Request",
) -> Path:
    """Create a small application-form document for selected-form API tests."""
    document = Document()
    header_table = document.sections[0].header.add_table(rows=1, cols=2, width=Inches(6))
    header_table.cell(0, 1).text = header_text
    table = document.add_table(rows=3, cols=2)
    for row_index, (label, value) in enumerate(
        [
            ("Form No.", "E-3718"),
            ("Requested By", "Alice Requestor"),
            ("Email", "alice@example.com"),
        ]
    ):
        table.cell(row_index, 0).text = label
        table.cell(row_index, 1).text = value
    sample_table = document.add_table(rows=2, cols=4)
    for index, header in enumerate(["Product Name", "Part Number", "Revision", "Quantity"]):
        sample_table.cell(0, index).text = header
    for index, value in enumerate(["Connector A", "PN-073", "A", "12"]):
        sample_table.cell(1, index).text = value
    document.save(path)
    return path
