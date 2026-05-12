from __future__ import annotations

from collections.abc import Generator
from datetime import date
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from backend.api.dependencies import get_session, get_settings
from backend.api.main import app
from backend.domain import Project, ProjectStatus
from backend.infrastructure.storage.database import (
    create_database_engine,
    create_session_factory,
    init_db,
)
from backend.infrastructure.storage.repositories import ProjectRepository
from backend.shared.config import Settings


def test_section2_write_back_api_updates_docx_and_creates_backup(tmp_path: Path) -> None:
    client, engine = _client(tmp_path)
    try:
        _create_project("P1", tmp_path)
        target = _section2_docx(tmp_path / "request.docx")
        draft_id = client.post(
            "/api/projects/P1/test-plan/drafts",
            json=_create_draft_payload(),
        ).json()["draft_id"]

        response = client.post(
            f"/api/projects/P1/test-plan/drafts/{draft_id}/section2-write-back",
            json={
                "target_application_form_path": str(target),
                "received_date": "2026-05-12",
                "lab": "Connector Lab",
                "assigned_personnel": "White",
                "sample_condition": "Good condition",
                "sample_preparation_days": 1,
                "test_group_scheduling_buffer_days": 1,
                "report_drafting_days": 3,
                "review_days": 1,
                "operator": "Alice",
            },
        )

        assert response.status_code == 200
        body = response.json()
        assert Path(body["backup_path"]).is_file()
        assert body["operator"] == "Alice"
        assert {item["field_key"] for item in body["changed_fields"]} == {
            "lab",
            "assigned_personnel",
            "received_date",
            "estimated_completion_date",
            "sample_condition",
        }
        values = _table_values(target)
        assert values["Lab"] == "Connector Lab"
        assert values["Estimated Completion Date"] == "2026-05-19"
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_section2_write_back_api_rejects_missing_section2_location(
    tmp_path: Path,
) -> None:
    client, engine = _client(tmp_path)
    try:
        _create_project("P1", tmp_path)
        target = _section2_docx(tmp_path / "request.docx", include_sample_condition=False)
        draft_id = client.post(
            "/api/projects/P1/test-plan/drafts",
            json=_create_draft_payload(),
        ).json()["draft_id"]

        response = client.post(
            f"/api/projects/P1/test-plan/drafts/{draft_id}/section2-write-back",
            json={
                "target_application_form_path": str(target),
                "received_date": "2026-05-12",
                "sample_condition": "Good condition",
            },
        )

        assert response.status_code == 400
        assert "sample_condition" in response.json()["detail"]
        assert _table_values(target)["Lab"] == ""
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def _client(tmp_path: Path) -> tuple[TestClient, object]:
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    return TestClient(app), engine


def _create_project(project_id: str, tmp_path: Path) -> None:
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    session_factory = create_session_factory(engine)
    with session_factory() as session:
        ProjectRepository(session).create(
            Project(
                project_id=project_id,
                project_no=f"DL-2026-05-{project_id}",
                product_name="Connector",
                requestor="Alice",
                status=ProjectStatus.LTR_REGISTERED,
                created_on=date(2026, 5, 12),
            )
        )
        session.commit()
    engine.dispose()


def _create_draft_payload() -> dict[str, object]:
    return {
        "source_document_path": "C:/spec.docx",
        "source_document_name": "spec.docx",
        "source_format": ".docx",
        "payload": {
            "groups": [
                {
                    "group_key": "group_1",
                    "group_label": "Group 1",
                    "steps": [
                        {"sequence": 1, "test_item": "Examination"},
                        {
                            "sequence": 2,
                            "test_item": "LLCR",
                            "estimated_duration_days": 1,
                        },
                    ],
                }
            ],
            "warnings": [],
            "blockers": [],
        },
    }


def _section2_docx(path: Path, *, include_sample_condition: bool = True) -> Path:
    document = Document()
    labels = [
        "Lab",
        "Assigned Personnel",
        "Received Date",
        "Estimated Completion Date",
    ]
    if include_sample_condition:
        labels.append("Sample Condition")
    table = document.add_table(rows=len(labels), cols=2)
    for index, label in enumerate(labels):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = ""
    document.save(path)
    return path


def _table_values(path: Path) -> dict[str, str]:
    document = Document(path)
    table = document.tables[0]
    return {
        row.cells[0].text.strip(): row.cells[1].text.strip()
        for row in table.rows
    }
