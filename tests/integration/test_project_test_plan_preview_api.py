from __future__ import annotations

from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from backend.api.dependencies import get_project_test_plan_matrix_preview_service
from backend.api.main import app
from backend.application.project_test_plan_matrix_preview_service import (
    MatrixPreviewFromPathCommand,
    ProjectTestPlanMatrixPreview,
)
from backend.infrastructure.office import OfficeAutomationUnavailable


def test_matrix_preview_api_extracts_docx_matrix(tmp_path: Path) -> None:
    docx_path = tmp_path / "product-spec.docx"
    _write_product_spec_docx(docx_path)
    client = TestClient(app)

    response = client.post(
        "/api/test-plan/matrix-preview-from-path",
        json={"source_path": str(docx_path), "project_id": "P1"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["project_id"] == "P1"
    assert payload["source_format"] == ".docx"
    assert payload["capability_status"] == "supported"
    assert payload["selected_table_index"] == 1
    assert payload["blockers"] == []
    assert [group["group_label"] for group in payload["groups"]] == [
        "Group 1",
        "Group 2",
    ]
    group_1 = payload["groups"][0]
    assert [step["sequence"] for step in group_1["steps"]] == [1, 2, 5, 8, 10]
    assert group_1["steps"][1]["test_item"] == "Contact Resistance (Low Level)"
    assert group_1["steps"][1]["source_section"] == "6.1"


def test_matrix_preview_path_api_passes_locator_fields_unchanged() -> None:
    fake_service = _FakeUploadPreviewService()
    app.dependency_overrides[get_project_test_plan_matrix_preview_service] = lambda: fake_service
    client = TestClient(app)
    try:
        response = client.post(
            "/api/test-plan/matrix-preview-from-path",
            json={
                "source_path": "D:/project/Submitted Material/spec.pdf",
                "project_id": "P-locator",
                "page_number": 7,
                "page_table_index": 2,
                "table_text_query": "qualification matrix",
            },
        )
    finally:
        app.dependency_overrides.pop(get_project_test_plan_matrix_preview_service, None)

    assert response.status_code == 200
    assert fake_service.previewed_locator == (7, 2, "qualification matrix")


def test_matrix_preview_api_extracts_cross_page_pdf_details(tmp_path: Path) -> None:
    pdf_path = tmp_path / "cross-page-mfg.pdf"
    _write_cross_page_mfg_pdf(pdf_path)
    client = TestClient(app)

    response = client.post(
        "/api/test-plan/matrix-preview-from-path",
        json={"source_path": str(pdf_path), "project_id": "P-pdf-cross-page"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["capability_status"] == "supported"
    mfg = next(row for row in payload["rows"] if row["test_item"] == "MFG")
    assert mfg["method"] == "EIA-364-65"
    assert mfg["condition"] == "Class IIA; unmated 224 hours; mated 112 hours"


def test_matrix_preview_api_reports_doc_as_deferred_and_no_text_pdf_as_unsupported(tmp_path: Path) -> None:
    doc_path = tmp_path / "spec.doc"
    pdf_path = tmp_path / "spec.pdf"
    doc_path.write_bytes(b"legacy-doc-placeholder")
    _write_blank_pdf(pdf_path)
    client = TestClient(app)

    doc_response = client.post(
        "/api/test-plan/matrix-preview-from-path",
        json={"source_path": str(doc_path)},
    )
    pdf_response = client.post(
        "/api/test-plan/matrix-preview-from-path",
        json={"source_path": str(pdf_path)},
    )

    assert doc_response.status_code == 200
    assert doc_response.json()["capability_status"] == "deferred"
    assert ".doc product specifications require" in doc_response.json()["blockers"][0]
    assert pdf_response.status_code == 200
    assert pdf_response.json()["capability_status"] == "unsupported"
    assert "no extractable text" in pdf_response.json()["blockers"][0].lower()


def test_matrix_preview_api_reports_missing_source(tmp_path: Path) -> None:
    client = TestClient(app)

    response = client.post(
        "/api/test-plan/matrix-preview-from-path",
        json={"source_path": str(tmp_path / "missing.docx")},
    )

    assert response.status_code == 404
    assert "not found" in response.json()["detail"]


def test_matrix_preview_api_maps_variant_marker_note_with_path_text(tmp_path: Path) -> None:
    docx_path = tmp_path / "product-spec-marker-variant.docx"
    _write_product_spec_docx_with_variant_note(docx_path)
    client = TestClient(app)

    response = client.post(
        "/api/test-plan/matrix-preview-from-path",
        json={"source_path": str(docx_path), "project_id": "P2"},
    )

    assert response.status_code == 200
    payload = response.json()
    group_1 = payload["groups"][0]
    step_1 = group_1["steps"][0]
    assert step_1["raw_token"] == "1(a)"
    assert step_1["source_note"] is not None
    assert "GS-12-2113" in step_1["source_note"]
    assert "Rev7.doc" in step_1["source_note"]


def test_matrix_preview_api_uses_last_contiguous_note_block_for_marker_mapping(tmp_path: Path) -> None:
    docx_path = tmp_path / "product-spec-note-block-scope.docx"
    _write_product_spec_docx_with_conflicting_note_blocks(docx_path)
    client = TestClient(app)

    response = client.post(
        "/api/test-plan/matrix-preview-from-path",
        json={"source_path": str(docx_path), "project_id": "P3"},
    )

    assert response.status_code == 200
    payload = response.json()
    group_1 = payload["groups"][0]
    notes_by_token = {step["raw_token"]: step.get("source_note") for step in group_1["steps"]}
    assert notes_by_token["3(a)"] == "(a) Precondition specimens with 20 durability cycles;"
    assert notes_by_token["10(c)"] == "(c) Energize at current for 18℃ temperature rise;"
    assert group_1["sample_note"] == "(e) Test with different 5 samples for solder ability and Resistance to solder heat, respectively"


def test_matrix_preview_api_rejects_test_record_like_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "test-record-like.docx"
    _write_test_record_like_docx(docx_path)
    client = TestClient(app)

    response = client.post(
        "/api/test-plan/matrix-preview-from-path",
        json={"source_path": str(docx_path), "project_id": "P4"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["groups"] == []
    assert payload["selected_table_index"] is None
    assert payload["capability_status"] == "unsupported"
    assert "No Matrix table" in payload["blockers"][0]


def test_matrix_preview_upload_accepts_doc_by_converting_to_temp_docx(tmp_path: Path) -> None:
    fake_service = _FakeUploadPreviewService()
    app.dependency_overrides[get_project_test_plan_matrix_preview_service] = lambda: fake_service
    client = TestClient(app)
    try:
        response = client.post(
            "/api/test-plan/matrix-preview-from-upload",
            files={"file": ("legacy-spec.doc", b"legacy word", "application/msword")},
            data={"project_id": "P-doc"},
        )
    finally:
        app.dependency_overrides.pop(get_project_test_plan_matrix_preview_service, None)

    assert response.status_code == 200
    payload = response.json()
    assert payload["preview_pdf_token"] is not None
    assert payload["project_id"] == "P-doc"
    assert payload["source_document_name"] == "legacy-spec.doc"
    assert payload["source_format"] == ".doc"
    assert payload["source_document_path"] == "legacy-spec.doc"
    assert not payload["source_document_path"].endswith(".docx")
    assert fake_service.office.converted_source is not None
    assert fake_service.office.converted_source.suffix == ".doc"
    assert fake_service.office.converted_output is not None
    assert fake_service.office.converted_output.suffix == ".docx"
    assert fake_service.previewed_source is not None
    assert fake_service.previewed_source == fake_service.office.converted_output
    assert fake_service.office.word_locations_requested == [fake_service.office.converted_output]
    assert len(fake_service.office.word_pdf_exports) == 1
    assert fake_service.office.word_pdf_exports[0][0] == fake_service.office.converted_output
    assert not fake_service.office.converted_source.exists()
    assert not fake_service.office.converted_output.exists()
    preview_response = client.get(
        f"/api/test-plan/matrix-preview-pdf/{payload['preview_pdf_token']}"
    )
    assert preview_response.status_code == 200
    assert preview_response.headers["content-type"].startswith("application/pdf")


def test_matrix_preview_upload_doc_conversion_failure_is_readable() -> None:
    fake_service = _FakeUploadPreviewService(conversion_error=OfficeAutomationUnavailable("Word COM automation requires pywin32."))
    app.dependency_overrides[get_project_test_plan_matrix_preview_service] = lambda: fake_service
    client = TestClient(app)
    try:
        response = client.post(
            "/api/test-plan/matrix-preview-from-upload",
            files={"file": ("legacy-spec.doc", b"legacy word", "application/msword")},
        )
    finally:
        app.dependency_overrides.pop(get_project_test_plan_matrix_preview_service, None)

    assert response.status_code == 400
    assert "Cannot convert legacy .doc for Matrix import" in response.json()["detail"]
    assert fake_service.previewed_source is None
    assert fake_service.office.converted_source is not None
    assert not fake_service.office.converted_source.exists()
    if fake_service.office.converted_output is not None:
        assert not fake_service.office.converted_output.exists()


def test_matrix_preview_upload_accepts_pdf_with_preview_token() -> None:
    fake_service = _FakeUploadPreviewService()
    app.dependency_overrides[get_project_test_plan_matrix_preview_service] = lambda: fake_service
    client = TestClient(app)
    try:
        response = client.post(
            "/api/test-plan/matrix-preview-from-upload",
            files={"file": ("spec.pdf", b"%PDF-1.4\n% fake", "application/pdf")},
            data={"project_id": "P-pdf", "page_number": "2", "page_table_index": "1"},
        )
        assert response.status_code == 200
        preview_response = client.get(
            f"/api/test-plan/matrix-preview-pdf/{response.json()['preview_pdf_token']}"
        )
    finally:
        app.dependency_overrides.pop(get_project_test_plan_matrix_preview_service, None)

    payload = response.json()
    assert payload["project_id"] == "P-pdf"
    assert payload["source_document_name"] == "spec.pdf"
    assert payload["source_format"] == ".pdf"
    assert fake_service.previewed_source is not None
    assert fake_service.previewed_source.suffix == ".pdf"
    assert fake_service.previewed_locator == (2, 1, None)
    assert fake_service.office.word_locations_requested == []
    assert fake_service.office.word_pdf_exports == []
    assert not fake_service.previewed_source.exists()
    assert preview_response.status_code == 200
    assert preview_response.headers["content-type"].startswith("application/pdf")


def test_matrix_preview_path_api_supports_docx_preview_pdf_token(
    tmp_path: Path,
) -> None:
    fake_service = _FakeUploadPreviewService()
    app.dependency_overrides[get_project_test_plan_matrix_preview_service] = lambda: fake_service
    docx_path = tmp_path / "direct-spec.docx"
    _write_product_spec_docx(docx_path)
    client = TestClient(app)
    try:
        response = client.post(
            "/api/test-plan/matrix-preview-from-path",
            json={"source_path": str(docx_path), "project_id": "P-direct"},
        )
        assert response.status_code == 200
        preview_response = client.get(
            f"/api/test-plan/matrix-preview-pdf/{response.json()['preview_pdf_token']}"
        )
    finally:
        app.dependency_overrides.pop(get_project_test_plan_matrix_preview_service, None)

    payload = response.json()
    assert payload["project_id"] == "P-direct"
    assert payload["source_document_name"] == docx_path.name
    assert payload["source_document_path"] == docx_path.name
    assert payload["source_format"] == ".docx"
    assert payload["preview_pdf_token"] is not None
    assert fake_service.office.word_locations_requested == [docx_path]
    assert fake_service.previewed_source == docx_path
    assert len(fake_service.office.word_pdf_exports) == 1
    assert preview_response.status_code == 200
    assert preview_response.headers["content-type"].startswith("application/pdf")


def _write_product_spec_docx(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=4, cols=4)
    rows = [
        ["test Items", "Section", "Group 1", "Group 2"],
        ["Examination of Product", "5.4", "1,10", "1,13"],
        ["Contact Resistance (Low Level)", "6.1", "2,5,8", "2,5,10"],
        ["Durability", "7.1", "", "3"],
    ]
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value
    document.save(path)


def _write_blank_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=letter)
    c.showPage()
    c.save()


def _write_cross_page_mfg_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=letter)
    c.drawString(72, 740, "8.2 MFG Reference - EIA-364-65")
    c.drawString(72, 720, "Mixed gas conditions refer to Clause 4.8 Industrial Mixed Gas.")
    c.drawString(72, 700, "Test Condition: CLASS IIA")
    c.showPage()
    c.drawString(72, 740, "Expose the connector in unmated condition for 224h.")
    c.drawString(72, 720, "Expose the connector in mated condition for 112h.")
    c.drawString(72, 700, "8.3 Voltage surge Power Pin 10 kA.")
    c.showPage()

    rows = [
        ["test Items", "Section", "Group 1"],
        ["MFG", "8.2", "1"],
    ]
    x0 = 72
    y0 = 700
    col_widths = [180, 70, 80]
    row_height = 26
    for row_index, row in enumerate(rows):
        y = y0 - row_index * row_height
        x = x0
        for col_index, value in enumerate(row):
            width = col_widths[col_index]
            c.rect(x, y - row_height, width, row_height)
            c.drawString(x + 4, y - 17, value)
            x += width
    c.showPage()
    c.save()


def _write_product_spec_docx_with_variant_note(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=3)
    rows = [
        ["test Items", "Section", "Group 1"],
        ["Examination", "5.5", "1(a)"],
    ]
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value
    document.add_paragraph(
        "a) C:\\Users\\White\\Desktop\\AI information\\Spec\\GS-12-2113 CoolPower HDF 3.40mm product specification_20251219_Rev7.doc"
    )
    document.save(path)


def _write_product_spec_docx_with_conflicting_note_blocks(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=4, cols=3)
    rows = [
        ["test Items", "Section", "Group 1"],
        ["Durability", "8.11", "3(a)"],
        ["Vibration Random", "8.9", "10(c)"],
        ["Samples Quantity (PCS)", "", "5+(5e)"],
    ]
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value
    document.add_paragraph("(a) Precondition Category E Test")
    document.add_paragraph("(c) Minimum solder coverage: 95 %")
    document.add_paragraph("Table 5: Qualification Test Table")
    document.add_paragraph("Precondition specimens with 20 durability cycles;")
    document.add_paragraph("Precondition specimens with 212 hours high temperature life;")
    document.add_paragraph("Energize at current for 18℃ temperature rise;")
    document.add_paragraph("5pcs for LLCR test another 5pcs loose connector for DWV test.")
    document.add_paragraph("(e) Test with different 5 samples for solder ability and Resistance to solder heat, respectively")
    document.save(path)


def _write_test_record_like_docx(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=4, cols=5)
    rows = [
        ["Test Item", "Requirement", "Result", "Judgement", "Record"],
        ["1", "As spec", "Pass", "OK", "notes"],
        ["2", "As spec", "Pass", "OK", "notes"],
        ["3", "As spec", "Pass", "OK", "notes"],
    ]
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value
    document.save(path)


class _FakeUploadPreviewOffice:
    def __init__(self, conversion_error: Exception | None = None) -> None:
        self.conversion_error = conversion_error
        self.converted_source: Path | None = None
        self.converted_output: Path | None = None
        self.word_locations_requested: list[Path] = []
        self.word_pdf_exports: list[tuple[Path, Path]] = []

    def convert_legacy_doc_to_docx(self, source_path: Path, output_path: Path) -> Path:
        self.converted_source = Path(source_path)
        self.converted_output = Path(output_path)
        if self.conversion_error is not None:
            raise self.conversion_error
        output_path.write_bytes(b"converted docx")
        return output_path

    def read_word_table_locations(self, source_path: Path) -> tuple:
        self.word_locations_requested.append(Path(source_path))
        assert source_path.suffix == ".docx"
        return ()

    def export_word_preview_pdf(self, source_path: Path, output_pdf_path: Path) -> Path:
        self.word_pdf_exports.append((Path(source_path), Path(output_pdf_path)))
        assert source_path.suffix == ".docx"
        output_pdf_path.write_bytes(b"%PDF-1.4")
        return output_pdf_path


class _FakeUploadPreviewService:
    def __init__(self, conversion_error: Exception | None = None) -> None:
        self._office = _FakeUploadPreviewOffice(conversion_error)
        self.previewed_source: Path | None = None
        self.previewed_locator: tuple[int | None, int | None, str | None] | None = None

    @property
    def office(self) -> _FakeUploadPreviewOffice:
        return self._office

    def convert_legacy_doc_to_docx(self, source_path: Path, output_path: Path) -> Path:
        return self._office.convert_legacy_doc_to_docx(source_path, output_path)

    def read_word_table_locations(self, source_path: Path) -> tuple:
        return self._office.read_word_table_locations(source_path)

    def export_word_preview_pdf(self, source_path: Path, output_pdf_path: Path) -> Path:
        return self._office.export_word_preview_pdf(source_path, output_pdf_path)

    def preview_from_path(
        self,
        command: MatrixPreviewFromPathCommand,
        *,
        preview_pdf_token: str | None = None,
        table_locations: tuple | None = None,
    ) -> ProjectTestPlanMatrixPreview:
        self.previewed_source = Path(command.source_path)
        self.previewed_locator = (
            command.page_number,
            command.page_table_index,
            command.table_text_query,
        )
        assert self.previewed_source.suffix in {".docx", ".pdf"}
        return ProjectTestPlanMatrixPreview(
            project_id=command.project_id,
            source_document_path=self.previewed_source,
            source_document_name=self.previewed_source.name,
            source_format=self.previewed_source.suffix,
            capability_status="supported",
            generated_at="2026-07-04T00:00:00+00:00",
            preview_pdf_token=preview_pdf_token,
        )
