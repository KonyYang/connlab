from __future__ import annotations

from pathlib import Path
from datetime import date
import json

from docx import Document

from backend.api.dependencies import get_settings
from backend.api.main import app
from backend.domain import IntakeCase, IntakeDraft, LtrRecord
from backend.domain.enums import IntakeCaseStatus, LtrStatus
from backend.infrastructure.storage.database import create_database_engine, create_session_factory
from backend.infrastructure.storage.repositories import LtrRecordRepository
from backend.infrastructure.storage.repositories.intake_package import (
    IntakeCaseRepository,
    IntakeDraftRepository,
)
from backend.shared.config import Settings, TestRecordSettings
from tests.integration.test_confirmed_matrix_test_record_preview_api import (
    _client,
    _seed_project,
    _seed_source_import,
)


def test_confirmed_matrix_test_record_generation_api_downloads_docx(tmp_path: Path) -> None:
    client, engine, _ = _client(tmp_path)
    try:
        template_path = _build_template(tmp_path / "template.docx")
        app.dependency_overrides[get_settings] = lambda: Settings(
            data_dir=tmp_path / "data",
            projects_dir=tmp_path / "projects",
            templates_dir=tmp_path / "templates",
            database_path=tmp_path / "connlab.sqlite3",
            test_record=TestRecordSettings(template_path=template_path),
        )
        _seed_project("P1", tmp_path)
        _seed_header_metadata_sources("P1", tmp_path)
        source_import_id = _seed_source_import("P1", tmp_path)
        draft = client.post(
            "/api/projects/P1/matrix-drafts",
            json={"source_import_id": source_import_id, "selected_group_keys": ["g1"]},
        )
        draft_id = draft.json()["record"]["project_matrix_draft_id"]
        confirm = client.post(
            f"/api/projects/P1/matrix-drafts/{draft_id}/confirm",
            json={"confirmed_by": "operator"},
        )
        assert confirm.status_code == 201

        response = client.post("/api/projects/P1/confirmed-matrix/test-record-draft/generate")

        assert response.status_code == 200
        assert response.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        output = tmp_path / "downloaded.docx"
        output.write_bytes(response.content)
        document = Document(output)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        header_text = "\n".join(
            cell.text
            for table in document.sections[0].header.tables
            for row in table.rows
            for cell in row.cells
        )
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        assert (
            "Group Number 组别编号: 1 ;   "
            "Sample Quantity & Number 样品数量及编号: 5 sets (Group1-1#~5#)"
        ) in text
        assert "Visual" in table_text
        assert "LLCR" in table_text
        assert "G2" not in text
        assert "DL-2026-05-003" in header_text
        assert "Coolpower HDF 3.40mm pin" in header_text
        assert "GS-12-1507" in header_text
        assert document.sections[0].header.tables[1].cell(0, 5).text == ""
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_confirmed_matrix_test_record_generation_api_returns_404_without_active_matrix(
    tmp_path: Path,
) -> None:
    client, engine, _ = _client(tmp_path)
    try:
        template_path = _build_template(tmp_path / "template.docx")
        app.dependency_overrides[get_settings] = lambda: Settings(
            data_dir=tmp_path / "data",
            projects_dir=tmp_path / "projects",
            templates_dir=tmp_path / "templates",
            database_path=tmp_path / "connlab.sqlite3",
            test_record=TestRecordSettings(template_path=template_path),
        )
        _seed_project("P1", tmp_path)
        response = client.post("/api/projects/P1/confirmed-matrix/test-record-draft/generate")
        assert response.status_code == 404
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_confirmed_matrix_test_record_generation_api_returns_422_without_template_config(
    tmp_path: Path,
) -> None:
    client, engine, _ = _client(tmp_path)
    try:
        _seed_project("P1", tmp_path)
        response = client.post("/api/projects/P1/confirmed-matrix/test-record-draft/generate")
        assert response.status_code == 422
        assert response.json()["detail"] == "Test Record template path is not configured."
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def _build_template(path: Path) -> Path:
    document = Document()
    header = document.sections[0].header
    table0 = header.add_table(rows=1, cols=3, width=7 * 914400)
    table0.cell(0, 2).text = "Lab Test Request Number:\n实验室测试项目编号："
    table1 = header.add_table(rows=1, cols=6, width=7 * 914400)
    table1.cell(0, 0).text = "Product Description\n产品描述"
    table1.cell(0, 2).text = "Applicable Specification\n适用的规范"
    table1.cell(0, 4).text = "Estimated Completion Date\n预计完成日期"
    document.add_paragraph("Group Number: PLACEHOLDER")
    step_table = document.add_table(rows=1, cols=9)
    step_table.rows[0].cells[0].text = "Step"
    document.add_paragraph("EQUIPMENT USED 使用的设备:")
    equipment_table = document.add_table(rows=1, cols=7)
    equipment_table.rows[0].cells[0].text = "Equipment"
    document.save(path)
    return path


def _seed_header_metadata_sources(project_id: str, tmp_path: Path) -> None:
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    session_factory = create_session_factory(engine)
    with session_factory() as session:
        LtrRecordRepository(session).create(
            LtrRecord(
                ltr_id="ltr-1",
                project_id=project_id,
                ltr_number="DL-2026-05-003",
                status=LtrStatus.REGISTERED,
                registered_on=date(2026, 5, 30),
                notes=json.dumps({"sample_description": "Coolpower HDF 3.40mm pin"}),
            )
        )
        IntakeCaseRepository(session).create(
            IntakeCase(
                case_id="case-1",
                package_id="pkg-1",
                selected_form_asset_id=None,
                status=IntakeCaseStatus.CONFIRMED,
                confirmed_project_id=project_id,
                created_at="2026-05-30T09:00:00+00:00",
                updated_at="2026-05-30T09:00:00+00:00",
            )
        )
        IntakeDraftRepository(session).create(
            IntakeDraft(
                draft_id="draft-1",
                case_id="case-1",
                parsed_fields_json="{}",
                sample_rows_json="[]",
                requested_testing_json=json.dumps(
                    [
                        {
                            "test_to_be_performed": "Visual",
                            "applicable_specification": "GS-12-1507",
                        }
                    ]
                ),
                field_confidence_json="{}",
                parser_warnings_json="[]",
                manual_overrides_json=None,
                updated_at="2026-05-30T09:00:00+00:00",
            )
        )
        session.commit()
    engine.dispose()
