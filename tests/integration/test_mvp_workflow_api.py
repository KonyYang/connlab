from collections.abc import Generator
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from backend.api.dependencies import get_session
from backend.api.main import app
from backend.domain import ProjectStatus
from backend.infrastructure.storage.database import (
    create_database_engine,
    create_session_factory,
    init_db,
)
from backend.infrastructure.storage.repositories import ProjectRepository
from backend.shared.config import Settings


def test_full_mvp_workflow_project_form_precheck_ltr_folder(tmp_path: Path) -> None:
    """A project can move through the full MVP backend workflow."""
    engine = create_database_engine(
        Settings(
            data_dir=tmp_path / "data",
            projects_dir=tmp_path / "projects",
            templates_dir=tmp_path / "templates",
            database_path=tmp_path / "connlab.sqlite3",
        )
    )
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    client = TestClient(app)

    try:
        project_response = client.post(
            "/api/projects",
            json={
                "project_no": "PRJ-014",
                "product_name": "MVP Connector",
                "requestor": "Dana",
                "business_unit": "BU-14",
            },
        )
        assert project_response.status_code == 201
        project_id = project_response.json()["project_id"]

        application_docx = _create_application_docx(tmp_path / "application.docx")
        with application_docx.open("rb") as handle:
            upload_response = client.post(
                f"/api/projects/{project_id}/application-form",
                files={
                    "file": (
                        "application.docx",
                        handle,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        assert upload_response.status_code == 201
        form_id = upload_response.json()["form_id"]

        precheck_response = client.post(
            f"/api/application-forms/{form_id}/precheck/run"
        )
        assert precheck_response.status_code == 200
        assert precheck_response.json()["status"] in {"passed", "warning"}
        assert isinstance(precheck_response.json()["issues"], list)

        ltr_response = client.post(
            f"/api/projects/{project_id}/ltr",
            json={"ltr_number": "LTR-014", "requested_by": "Dana"},
        )
        assert ltr_response.status_code == 201
        assert ltr_response.json()["status"] == "registered"

        template = _create_template(tmp_path / "{PROJECT_NO}_{PRODUCT_NAME}")
        target_root = tmp_path / "generated"
        target_root.mkdir()
        folder_payload = {
            "template_path": str(template),
            "target_root": str(target_root),
            "dl_number": "DL-014",
            "plan_date": "2026-04-26",
        }

        preview_response = client.post(
            f"/api/projects/{project_id}/folder/preview",
            json=folder_payload,
        )
        assert preview_response.status_code == 200
        assert preview_response.json()["conflict"] is False

        generate_response = client.post(
            f"/api/projects/{project_id}/folder/generate",
            json=folder_payload,
        )
        assert generate_response.status_code == 201
        assert Path(generate_response.json()["project_folder_path"]).is_dir()

        refreshed = client.get(f"/api/projects/{project_id}")
        assert refreshed.status_code == 200
        assert refreshed.json()["status"] == "folder_created"

        with session_factory() as session:
            project = ProjectRepository(session).get(project_id)
            assert project is not None
            assert project.status is ProjectStatus.FOLDER_CREATED
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def _create_template(template: Path) -> Path:
    (template / "request").mkdir(parents=True)
    (template / "request" / "form.txt").write_text("template", encoding="utf-8")
    return template


def _create_application_docx(path: Path) -> Path:
    document = Document()
    table = document.add_table(rows=10, cols=2)
    pairs = [
        ("Form No.", "E-3718"),
        ("Form Rev", "H"),
        ("Requested By", "Dana"),
        ("Phone", "555-0140"),
        ("Date", "2026-04-26"),
        ("Email", "dana@example.com"),
        ("Business Unit", "BU-14"),
        ("Mfg. Site", "Plant 14"),
        ("Project #", "PRJ-014"),
        ("Description of Requested Testing", "Salt spray and visual inspection"),
    ]
    for index, (label, value) in enumerate(pairs):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value
    sample = document.add_table(rows=2, cols=7)
    headers = [
        "Product Name",
        "Part Number",
        "Lot/Traceability",
        "Material",
        "Plating",
        "Housing Material",
        "Quantity",
    ]
    values = ["MVP Connector", "PN-014", "LOT-014", "Copper", "Tin", "LCP", "12"]
    for index, header in enumerate(headers):
        sample.cell(0, index).text = header
        sample.cell(1, index).text = values[index]
    document.save(path)
    return path
