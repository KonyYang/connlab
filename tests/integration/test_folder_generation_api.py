from collections.abc import Generator
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from backend.api.dependencies import get_session
from backend.api.main import app
from backend.domain import ProjectStatus
from backend.infrastructure.storage.database import (
    create_database_engine,
    create_session_factory,
    init_db,
)
from backend.infrastructure.storage.repositories import (
    ProjectFolderRecordRepository,
    ProjectRepository,
)
from backend.shared.config import Settings


def test_folder_api_preview_generate_refuse_overwrite_and_persist(
    tmp_path: Path,
) -> None:
    engine = create_database_engine(
        Settings(
            data_dir=tmp_path / "data",
            projects_dir=tmp_path / "projects",
            templates_dir=tmp_path / "templates",
            database_path=tmp_path / "connlab.sqlite3",
        )
    )
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    client = TestClient(app)

    try:
        template = _create_template(tmp_path / "{PROJECT_NO}_{PRODUCT_NAME}")
        target_root = tmp_path / "generated"
        target_root.mkdir()
        project_response = client.post(
            "/api/projects",
            json={
                "project_no": "PRJ-001",
                "product_name": "Connector",
                "requestor": "Alice",
            },
        )
        project_id = project_response.json()["project_id"]
        application_docx = _create_application_docx(tmp_path / "application.docx")
        with application_docx.open("rb") as handle:
            upload_response = client.post(
                f"/api/projects/{project_id}/application-form",
                files={
                    "file": (
                        "application.docx",
                        handle,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
        )
        assert upload_response.status_code == 201
        _set_project_status(session_factory, project_id, ProjectStatus.LTR_REGISTERED)
        payload = {
            "template_path": str(template),
            "target_root": str(target_root),
            "plan_date": "2026-04-26",
        }

        preview_response = client.post(
            f"/api/projects/{project_id}/folder/preview",
            json=payload,
        )
        generate_response = client.post(
            f"/api/projects/{project_id}/folder/generate",
            json=payload,
        )
        overwrite_response = client.post(
            f"/api/projects/{project_id}/folder/generate",
            json=payload,
        )

        assert preview_response.status_code == 200
        assert preview_response.json()["conflict"] is False
        assert generate_response.status_code == 201
        assert overwrite_response.status_code == 400
        assert "already been created" in overwrite_response.json()["detail"]

        generated = generate_response.json()
        assert Path(generated["project_folder_path"]).is_dir()
        request_dir = Path(generated["project_folder_path"]) / "request"
        assert (request_dir / "form.txt").is_file()
        assert (request_dir / "application.docx").is_file()

        with session_factory() as session:
            folders = ProjectFolderRecordRepository(session).list_by_project(project_id)
            project = ProjectRepository(session).get(project_id)
            assert len(folders) == 1
            assert project is not None
            assert project.status is ProjectStatus.FOLDER_CREATED
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def _set_project_status(
    session_factory,
    project_id: str,
    status: ProjectStatus,
) -> None:
    with session_factory() as session:
        repository = ProjectRepository(session)
        project = repository.get(project_id)
        assert project is not None
        repository.update(project.with_status(status))
        session.commit()


def _create_template(template: Path) -> Path:
    (template / "request").mkdir(parents=True)
    (template / "request" / "form.txt").write_text("template", encoding="utf-8")
    return template


def _create_application_docx(path: Path) -> Path:
    document = Document()
    table = document.add_table(rows=10, cols=2)
    pairs = [
        ("Form No.", "E-3718"),
        ("Form Rev", "H"),
        ("Requested By", "Alice"),
        ("Phone", "555-0100"),
        ("Date", "2026-04-26"),
        ("Email", "alice@example.com"),
        ("Business Unit", "BU-1"),
        ("Mfg. Site", "Plant 1"),
        ("Project #", "PRJ-001"),
        ("Description of Requested Testing", "Salt spray test"),
    ]
    for index, (label, value) in enumerate(pairs):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value
    sample = document.add_table(rows=2, cols=7)
    headers = [
        "Product Name",
        "Part Number",
        "Lot/Traceability",
        "Material",
        "Plating",
        "Housing Material",
        "Quantity",
    ]
    values = ["Connector", "PN-001", "LOT-1", "Copper", "Tin", "LCP", "12"]
    for index, header in enumerate(headers):
        sample.cell(0, index).text = header
        sample.cell(1, index).text = values[index]
    document.save(path)
    return path
