from __future__ import annotations

from collections.abc import Generator
from datetime import date
from pathlib import Path

from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Inches

from backend.api.dependencies import get_session, get_settings
from backend.api.routes_intake import router as intake_router
from backend.api.routes_intake_review import router as intake_review_router
from backend.api.main import app
from backend.infrastructure.storage.database import (
    create_database_engine,
    create_session_factory,
    init_db,
)
from backend.infrastructure.storage.repositories.intake_package import (
    IntakeAssetRepository,
    IntakeCaseRepository,
    IntakeDraftRepository,
    IntakePackageRepository,
)
from backend.infrastructure.storage.repositories.intake import (
    ApplicationFormRepository,
    SampleInfoRepository,
)
from backend.infrastructure.storage.repositories.project import ProjectRepository
from backend.infrastructure.storage.repositories import LtrRecordRepository
from backend.domain import (
    IntakeAsset,
    IntakeAssetRole,
    IntakePackage,
    IntakePackageSourceType,
    IntakePackageStatus,
    LtrRecord,
    LtrStatus,
)
from backend.shared.config import Settings


def test_direct_word_intake_api_creates_package_and_candidate(tmp_path: Path) -> None:
    """Direct Word upload creates an intake package without an email source."""
    settings, engine, session_factory = _test_database(tmp_path)
    app.dependency_overrides[get_session] = _override_session(session_factory)
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)
    sample_path = tmp_path / "direct_application.docx"
    document = Document()
    document.add_paragraph("Laboratory Test Request")
    document.save(sample_path)

    try:
        with sample_path.open("rb") as handle:
            response = client.post(
                "/api/intake-packages/import-docx",
                files={
                    "file": (
                        sample_path.name,
                        handle,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        assert response.status_code == 201
        payload = response.json()
        assert payload["source_type"] == "direct_application_form"
        assert payload["source_original_name"] == sample_path.name
        assert payload["asset_count"] == 1
        assert payload["candidate_count"] == 1
        assert "received_at" in payload
        assert payload["received_at"] is None
        assert payload["next_action"] == "review_application_form_candidates"
        assert payload["assets"][0]["original_name"] == sample_path.name
        assert payload["assets"][0]["asset_role"] in {
            "application_form_candidate",
            "selected_application_form",
        }
    finally:
        app.dependency_overrides.clear()


def test_manual_intake_api_creates_review_case(tmp_path: Path) -> None:
    """The API stores no-email manual intake before project creation."""
    settings, engine, session_factory = _test_database(tmp_path)
    app.dependency_overrides[get_session] = _override_session(session_factory)
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        response = client.post(
            "/api/intake-packages/manual",
                json={
                    "form_no": "E-3718",
                    "revision": "H",
                    "product_name": "Connector sample",
                    "requester": "White",
                    "phone": "555-0100",
                    "request_date": "2026-05-03",
                    "email": "white@example.com",
                    "business_unit": "Power Solutions",
                    "project_no": "PRJ-1",
                    "requested_testing": "Qualification",
                    "sample": {
                        "product_name": "Connector sample",
                        "part_number": "PN-001",
                        "lot_or_traceability": "LOT-001",
                        "material": "Copper",
                        "plating": "Ag",
                        "housing_material": "PA10T",
                        "quantity": 3,
                    },
                },
            )

        assert response.status_code == 201
        payload = response.json()
        assert payload["package_status"] == "ready_for_review"
        assert payload["missing_required_fields"] == []
        assert payload["next_action"] == "review_manual_intake"

        detail_response = client.get(f"/api/intake-packages/{payload['package_id']}")
        assert detail_response.status_code == 200
        detail = detail_response.json()
        assert detail["source_type"] == "manual"
        assert detail["source_stored"] is True
        assert detail["case_count"] == 1
        assert detail["candidate_count"] == 1

        review_response = client.get(
            f"/api/intake-packages/{payload['package_id']}/case-review"
        )
        assert review_response.status_code == 200
        review = review_response.json()
        assert review["source_type"] == "manual"
        assert review["cases"][0]["case_id"] == payload["case_id"]
        assert review["cases"][0]["confirm_allowed"] is False
        assert review["cases"][0]["precheck_issues"]
        assert review["cases"][0]["missing_required_fields"]
        assert any(
            field["key"] == "product_name" and field["value"] == "Connector sample"
            for field in review["cases"][0]["fields"]
        )

        rejected_response = client.post(
            f"/api/intake-cases/{payload['case_id']}/confirm",
            json={"operator_confirmed": False},
        )
        assert rejected_response.status_code == 400

        confirm_response = client.post(
            f"/api/intake-cases/{payload['case_id']}/confirm",
            json={"operator_confirmed": True},
        )
        assert confirm_response.status_code == 400
        assert "SECTION 1 precheck blockers" in confirm_response.json()["detail"]

        with session_factory() as session:
            package = IntakePackageRepository(session).get(payload["package_id"])
            case = IntakeCaseRepository(session).get(payload["case_id"])
            draft = IntakeDraftRepository(session).get(payload["draft_id"])
            assert package is not None
            assert package.source_stored_path.is_file()
            assert case is not None
            assert case.confirmed_project_id is None
            assert draft is not None
            assert "Connector sample" in draft.parsed_fields_json
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_manual_intake_api_returns_missing_required_fields(tmp_path: Path) -> None:
    """Manual drafts can be saved while required field blockers remain visible."""
    settings, engine, session_factory = _test_database(tmp_path)
    app.dependency_overrides[get_session] = _override_session(session_factory)
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        response = client.post(
            "/api/intake-packages/manual",
            json={"product_name": "", "requester": ""},
        )

        assert response.status_code == 201
        payload = response.json()
        assert payload["missing_required_fields"] == ["product_name", "requester"]
        assert payload["next_action"] == "complete_required_manual_fields"

        review_response = client.get(
            f"/api/intake-packages/{payload['package_id']}/case-review"
        )
        assert review_response.status_code == 200
        review = review_response.json()
        assert review["cases"][0]["confirm_allowed"] is False
        assert "product_name" in review["cases"][0]["missing_required_fields"]
        assert "requester" in review["cases"][0]["missing_required_fields"]
        assert review["cases"][0]["precheck_issues"]

        update_response = client.patch(
            f"/api/intake-cases/{payload['case_id']}/review-fields",
            json={
                "fields": {
                    "form_no": "E-3718",
                    "revision": "H",
                    "product_name": "Corrected connector",
                    "requester": "White",
                    "phone": "555-0100",
                    "request_date": "2026-05-03",
                    "email": "white@example.com",
                    "business_unit": "Power Solutions",
                    "manufacturing_site": "Nantong",
                    "results_format": "Formal Report (Customer)",
                    "requested_completion_date": "2026-05-10",
                    "test_type": "Customer Specific Testing",
                    "sample_status": "Production",
                    "project_type": "New Product Development",
                    "requested_testing": "Bend testing",
                    "post_testing_disposition": "Keep in the Lab",
                    "confidential": "No",
                    "subcontract": "Yes",
                    "send_copies_recipients": "Neo Xu",
                },
                "sample_rows": [
                    {
                        "product_name": "Corrected connector",
                        "part_number": "PN-082",
                        "lot_or_traceability": "LOT-082",
                        "material": "Copper",
                        "plating": "Ag",
                        "housing_material": "PA10T",
                        "quantity": "4",
                    }
                ],
            },
        )
        assert update_response.status_code == 200
        updated_case = update_response.json()
        assert updated_case["confirm_allowed"] is True
        assert updated_case["missing_required_fields"] == []
        assert updated_case["sample_rows"][0]["part_number"] == "PN-082"

        confirm_after_update_response = client.post(
            f"/api/intake-cases/{payload['case_id']}/confirm",
            json={"operator_confirmed": True},
        )
        assert confirm_after_update_response.status_code == 200
        assert confirm_after_update_response.json()["project_id"]
        assert confirm_after_update_response.json()["sample_count"] == 1

        with session_factory() as session:
            draft = IntakeDraftRepository(session).get(payload["draft_id"])
            assert draft is not None
            assert "Corrected connector" in (draft.manual_overrides_json or "")
            project_id = confirm_after_update_response.json()["project_id"]
            project = ProjectRepository(session).get(project_id)
            forms = ApplicationFormRepository(session).list_by_project(project_id)
            samples = SampleInfoRepository(session).list_by_project(project_id)
            assert project is not None
            assert project.product_name == "Corrected connector"
            assert project.requestor == "White"
            assert project.business_unit == "Power Solutions"
            assert forms[0].requested_testing == "Bend testing"
            assert forms[0].requester == "White"
            assert samples[0].product_name == "Corrected connector"
            assert samples[0].part_number == "PN-082"
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_new_project_application_draft_endpoint_reuses_durable_case(tmp_path: Path) -> None:
    """TASK_102 endpoint prepares the single-page editor through API routing."""
    settings, engine, session_factory = _test_database(tmp_path)
    app.dependency_overrides[get_session] = _override_session(session_factory)
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        create_response = client.post(
            "/api/intake-packages/manual",
            json={"product_name": "", "requester": ""},
        )
        assert create_response.status_code == 201
        created = create_response.json()

        draft_response = client.post(
            f"/api/intake-packages/{created['package_id']}/application-draft"
        )

        assert draft_response.status_code == 200
        payload = draft_response.json()
        assert payload["package_id"] == created["package_id"]
        assert payload["case_id"] == created["case_id"]
        assert payload["draft_id"] == created["draft_id"]
        assert payload["next_action"] == "edit_application_information"
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_new_project_application_draft_endpoint_defaults_highest_ranked_candidate(
    tmp_path: Path,
) -> None:
    """The single-page editor opens with the strongest candidate form by default."""
    settings, engine, session_factory = _test_database(tmp_path)
    app.dependency_overrides[get_session] = _override_session(session_factory)
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        form_a = _create_application_docx(tmp_path / "candidate-a.docx")
        form_b = _create_application_docx(tmp_path / "candidate-b.docx")
        with session_factory() as session:
            package_repo = IntakePackageRepository(session)
            asset_repo = IntakeAssetRepository(session)
            package_repo.create(
                IntakePackage(
                    package_id="pkg-draft-default",
                    source_type=IntakePackageSourceType.OUTLOOK_MSG,
                    status=IntakePackageStatus.READY_FOR_REVIEW,
                    source_original_name="request.msg",
                    source_stored_path=tmp_path / "request.msg",
                )
            )
            asset_repo.create(
                IntakeAsset(
                    asset_id="asset-low",
                    package_id="pkg-draft-default",
                    original_name=form_a.name,
                    stored_path=form_a,
                    extension=".docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    size_bytes=form_a.stat().st_size,
                    sha256="a" * 64,
                    asset_role=IntakeAssetRole.APPLICATION_FORM_CANDIDATE,
                    candidate_score=60,
                )
            )
            asset_repo.create(
                IntakeAsset(
                    asset_id="asset-high",
                    package_id="pkg-draft-default",
                    original_name=form_b.name,
                    stored_path=form_b,
                    extension=".docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    size_bytes=form_b.stat().st_size,
                    sha256="b" * 64,
                    asset_role=IntakeAssetRole.APPLICATION_FORM_CANDIDATE,
                    candidate_score=90,
                )
            )
            session.commit()

        draft_response = client.post(
            "/api/intake-packages/pkg-draft-default/application-draft"
        )

        assert draft_response.status_code == 200
        payload = draft_response.json()
        assert payload["package_id"] == "pkg-draft-default"
        assert payload["selected_form_asset_id"] == "asset-high"
        assert payload["next_action"] == "edit_application_information"

        review_response = client.get("/api/intake-packages/pkg-draft-default/case-review")
        assert review_response.status_code == 200
        cases = review_response.json()["cases"]
        assert len(cases) == 1
        assert cases[0]["selected_form_asset_id"] == "asset-high"
        field_values = {field["key"]: field["value"] for field in cases[0]["fields"]}
        assert field_values["form_no"] == "E-3718"
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_unsaved_creation_draft_discard_removes_records_and_files(tmp_path: Path) -> None:
    """Exit without saving removes ConnLab-owned temporary intake state."""
    settings, engine, session_factory = _test_database(tmp_path)
    app.dependency_overrides[get_session] = _override_session(session_factory)
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        create_response = client.post(
            "/api/intake-packages/manual",
            json={"product_name": "Connector sample", "requester": "White"},
        )
        assert create_response.status_code == 201
        created = create_response.json()
        package_id = created["package_id"]
        package_dir = settings.data_dir / "intake" / package_id
        assert package_dir.exists()

        discard_response = client.post(
            f"/api/intake-packages/{package_id}/draft/discard"
        )

        assert discard_response.status_code == 200
        payload = discard_response.json()
        assert payload["action"] == "discard_unsaved"
        assert payload["deleted_package"] is True
        assert payload["deleted_assets"] == 1
        assert payload["deleted_cases"] == 1
        assert payload["deleted_drafts"] == 1
        assert payload["deleted_files"] is True
        assert "imported copies were removed" in payload["message"]
        assert not package_dir.exists()
        with session_factory() as session:
            assert IntakePackageRepository(session).get(package_id) is None
            assert IntakeAssetRepository(session).list_by_package(package_id) == []
            assert IntakeCaseRepository(session).list_by_package(package_id) == []
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def _create_application_docx(path: Path) -> Path:
    """Create a small real application form for endpoint tests."""
    document = Document()
    header = document.sections[0].header.add_table(rows=1, cols=2, width=Inches(6))
    header.cell(0, 1).text = "Laboratory Testing Request"
    footer = document.sections[0].footer
    footer.paragraphs[0].text = "Form No. E-3718 Rev H"
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Form No."
    table.cell(0, 1).text = "E-3718"
    table.cell(1, 0).text = "Requested By"
    table.cell(1, 1).text = "Alice Requestor"
    document.save(path)
    return path


def test_saved_creation_draft_is_not_removed_by_unsaved_discard(tmp_path: Path) -> None:
    """Save draft protects the package from the unsaved-session discard path."""
    settings, engine, session_factory = _test_database(tmp_path)
    app.dependency_overrides[get_session] = _override_session(session_factory)
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        create_response = client.post(
            "/api/intake-packages/manual",
            json={"product_name": "Connector sample", "requester": "White"},
        )
        assert create_response.status_code == 201
        package_id = create_response.json()["package_id"]
        package_dir = settings.data_dir / "intake" / package_id

        save_response = client.post(f"/api/intake-packages/{package_id}/draft/save")
        assert save_response.status_code == 200
        saved = save_response.json()
        assert saved["action"] == "save_draft"
        assert saved["package_status"] == "draft_saved"

        discard_response = client.post(
            f"/api/intake-packages/{package_id}/draft/discard"
        )

        assert discard_response.status_code == 400
        assert package_dir.exists()
        with session_factory() as session:
            package = IntakePackageRepository(session).get(package_id)
            assert package is not None
            assert package.status.value == "draft_saved"
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_saved_creation_drafts_list_returns_continue_metadata(tmp_path: Path) -> None:
    """Drafts / In Progress API lists saved drafts with continuation metadata."""
    settings, engine, _ = _test_database(tmp_path)
    test_app = FastAPI(title="Test App")
    test_app.include_router(intake_router)
    test_app.include_router(intake_review_router)
    session_factory = create_session_factory(engine)
    test_app.dependency_overrides[get_session] = _override_session(session_factory)
    test_app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(test_app)

    try:
        create_response = client.post(
            "/api/intake-packages/manual",
            json={"product_name": "Connector sample", "requester": "White"},
        )
        assert create_response.status_code == 201
        created = create_response.json()

        save_response = client.post(
            f"/api/intake-packages/{created['package_id']}/draft/save"
        )
        assert save_response.status_code == 200

        list_response = client.get("/api/project-creation-drafts")

        assert list_response.status_code == 200
        drafts = list_response.json()
        assert len(drafts) == 1
        draft = drafts[0]
        assert draft["package_id"] == created["package_id"]
        assert draft["product_name"] == "Connector sample"
        assert draft["requester"] == "White"
        assert draft["current_step"] == "precheck"
        assert draft["selected_form_asset_id"] == created["selected_form_asset_id"]
        assert draft["active_case_id"] == created["case_id"]
    finally:
        test_app.dependency_overrides.clear()
        engine.dispose()


def test_saved_creation_draft_discard_endpoint_removes_draft(tmp_path: Path) -> None:
    """Drafts / In Progress discard removes a saved draft through its own endpoint."""
    settings, engine, session_factory = _test_database(tmp_path)
    app.dependency_overrides[get_session] = _override_session(session_factory)
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        create_response = client.post(
            "/api/intake-packages/manual",
            json={"product_name": "Connector sample", "requester": "White"},
        )
        assert create_response.status_code == 201
        package_id = create_response.json()["package_id"]
        assert client.post(f"/api/intake-packages/{package_id}/draft/save").status_code == 200

        discard_response = client.post(
            f"/api/project-creation-drafts/{package_id}/discard"
        )

        assert discard_response.status_code == 200
        payload = discard_response.json()
        assert payload["action"] == "discard_saved_draft"
        assert "Saved creation draft discarded" in payload["message"]
        assert client.get("/api/project-creation-drafts").json() == []
        with session_factory() as session:
            assert IntakePackageRepository(session).get(package_id) is None
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def _test_database(tmp_path: Path):
    """Create an isolated test database."""
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    return settings, engine, create_session_factory(engine)


def test_review_fields_persists_requested_testing_rows(tmp_path: Path) -> None:
    """PATCH /review-fields with requested_testing_rows persists and returns rows."""
    settings, engine, session_factory = _test_database(tmp_path)

    from fastapi.testclient import TestClient

    # Use IntakeDraftRepository from top-level imports
    # No additional imports needed

    # Create a fresh app instance for this test
    test_app = FastAPI(title="Test App")
    test_app.include_router(intake_router)
    test_app.include_router(intake_review_router)
    test_app.dependency_overrides[get_session] = _override_session(session_factory)

    client = TestClient(test_app)

    try:
        create_response = client.post(
            "/api/intake-packages/manual",
            json={
                "product_name": "Connector sample",
                "requester": "Test user",
                "email": "test@example.com",
                "business_unit": "Power Solutions",
                "project_no": "PRJ-001",
            },
        )
        assert create_response.status_code == 201
        payload = create_response.json()

        # First update: complete all required fields
        update_response = client.patch(
            f"/api/intake-cases/{payload['case_id']}/review-fields",
            json={
                "fields": {
                    "form_no": "E-3718",
                    "revision": "H",
                    "product_name": "Connector sample",
                    "requester": "Test user",
                    "phone": "555-0100",
                    "request_date": "2026-05-03",
                    "email": "test@example.com",
                    "business_unit": "Power Solutions",
                    "manufacturing_site": "Nantong",
                    "results_format": "Formal Report (Customer)",
                    "requested_completion_date": "2026-05-10",
                    "test_type": "Customer Specific Testing",
                    "sample_status": "Production",
                    "project_type": "New Product Development",
                    "post_testing_disposition": "Keep in the Lab",
                    "confidential": "No",
                    "subcontract": "Yes",
                    "send_copies_recipients": "Team",
                },
            },
        )
        assert update_response.status_code == 200

        # Second update: add requested_testing_rows
        rows_update_response = client.patch(
            f"/api/intake-cases/{payload['case_id']}/review-fields",
            json={
                "fields": {},
                "requested_testing_rows": [
                    {
                        "test_to_be_performed": "Qualification test",
                        "applicable_specification": "GS-12-2652-22",
                    },
                    {
                        "test_to_be_performed": "Environmental test",
                        "applicable_specification": "QG-03-016E_Rev2",
                    },
                ],
            },
        )
        assert rows_update_response.status_code == 200
        updated_case = rows_update_response.json()

        # Verify response contains rows
        assert "requested_testing_rows" in updated_case
        assert len(updated_case["requested_testing_rows"]) == 2
        assert updated_case["requested_testing_rows"][0]["test_to_be_performed"] == "Qualification test"
        assert updated_case["requested_testing_rows"][0]["applicable_specification"] == "GS-12-2652-22"

        # Verify compatibility field is synced
        fields_dict = {f["key"]: f["value"] for f in updated_case["fields"]}
        assert "requested_testing" in fields_dict
        assert "Qualification test" in fields_dict["requested_testing"]
        assert "Environmental test" in fields_dict["requested_testing"]

        # Verify draft persistence
        with session_factory() as session:
            draft = IntakeDraftRepository(session).get(payload["draft_id"])
            assert draft is not None
            assert draft.manual_overrides_json is not None
            import json

            overrides = json.loads(draft.manual_overrides_json)
            assert "requested_testing_rows" in overrides
            assert len(overrides["requested_testing_rows"]) == 2
            assert overrides["requested_testing_rows"][0]["test_to_be_performed"] == "Qualification test"
            # Compatibility field should also be in overrides
            assert "requested_testing" in overrides
    finally:
        test_app.dependency_overrides.clear()
        engine.dispose()


def test_review_fields_returns_conflict_after_registered_ltr(tmp_path: Path) -> None:
    """PATCH /review-fields rejects frozen base edits after registered LTR."""
    settings, engine, session_factory = _test_database(tmp_path)

    test_app = FastAPI(title="Test App")
    test_app.include_router(intake_router)
    test_app.include_router(intake_review_router)
    test_app.dependency_overrides[get_session] = _override_session(session_factory)
    test_app.dependency_overrides[get_settings] = lambda: settings

    client = TestClient(test_app)

    try:
        create_response = client.post(
            "/api/intake-packages/manual",
            json={
                "product_name": "Connector sample",
                "requester": "Test user",
                "email": "test@example.com",
                "business_unit": "Power Solutions",
                "project_no": "PRJ-001",
            },
        )
        assert create_response.status_code == 201
        created = create_response.json()
        update_response = client.patch(
            f"/api/intake-cases/{created['case_id']}/review-fields",
            json={
                "fields": {
                    "form_no": "E-3718",
                    "revision": "H",
                    "product_name": "Connector sample",
                    "requester": "Test user",
                    "phone": "555-0100",
                    "request_date": "2026-05-03",
                    "email": "test@example.com",
                    "business_unit": "Power Solutions",
                    "manufacturing_site": "Nantong",
                    "results_format": "Formal Report (Customer)",
                    "requested_completion_date": "2026-05-10",
                    "test_type": "Customer Specific Testing",
                    "sample_status": "Production",
                    "project_type": "New Product Development",
                    "post_testing_disposition": "Keep in the Lab",
                    "requested_testing": "Bend testing",
                    "confidential": "No",
                    "subcontract": "Yes",
                    "send_copies_recipients": "Team",
                },
                "sample_rows": [
                    {
                        "product_name": "Connector sample",
                        "part_number": "PN-100",
                        "quantity": "20 pcs",
                    }
                ],
            },
        )
        assert update_response.status_code == 200
        confirm_response = client.post(
            f"/api/intake-cases/{created['case_id']}/confirm",
            json={"operator_confirmed": True},
        )
        assert confirm_response.status_code == 200
        project_id = confirm_response.json()["project_id"]
        with session_factory() as session:
            LtrRecordRepository(session).create(
                LtrRecord(
                    ltr_id="ltr-registered-1",
                    project_id=project_id,
                    ltr_number="DL-2026-05-001",
                    status=LtrStatus.REGISTERED,
                    registered_on=date(2026, 5, 7),
                )
            )
            session.commit()

        review_response = client.get(
            f"/api/intake-packages/{created['package_id']}/case-review"
        )
        assert review_response.status_code == 200
        review_case = review_response.json()["cases"][0]
        assert review_case["base_editing_frozen"] is True
        assert "product_name" in review_case["frozen_field_keys"]

        frozen_update_response = client.patch(
            f"/api/intake-cases/{created['case_id']}/review-fields",
            json={"fields": {"product_name": "Renamed connector"}},
        )

        assert frozen_update_response.status_code == 409
        assert "revise/exception" in frozen_update_response.json()["detail"]["message"]
        assert frozen_update_response.json()["detail"]["field_keys"] == ["product_name"]
    finally:
        test_app.dependency_overrides.clear()
        engine.dispose()


def _override_session(session_factory):
    """Return a FastAPI session dependency override."""

    def override_session() -> Generator[Session, None, None]:
        """Yield one test database session."""
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    return override_session
