from __future__ import annotations

from pathlib import Path

from docx import Document

from backend.api.dependencies import get_settings
from backend.api.main import app
from backend.application.official_project_workspace_service import OfficialWorkspaceRecord
from backend.domain import (
    ExternalResource,
    ExternalResourceType,
    ExternalResourceValidationStatus,
)
from backend.infrastructure.storage.repositories import ExternalResourceRepository
from backend.infrastructure.storage.repositories.official_workspace import (
    ProjectOfficialWorkspaceRepository,
)
from backend.shared.config import Settings, TestRecordSettings
from tests.integration.test_confirmed_matrix_test_record_generation_api import (
    _build_template,
    _seed_basic_information,
    _seed_header_metadata_sources,
)
from tests.integration.test_confirmed_matrix_test_record_preview_api import (
    _client,
    _seed_project,
)


def test_matrix_editor_test_record_generation_uses_current_ui_payload(
    tmp_path: Path,
) -> None:
    client, engine, _ = _client(tmp_path)
    try:
        template_path = _build_template(tmp_path / "template.docx")
        app.dependency_overrides[get_settings] = lambda: Settings(
            data_dir=tmp_path / "data",
            projects_dir=tmp_path / "projects",
            templates_dir=tmp_path / "templates",
            database_path=tmp_path / "connlab.sqlite3",
            test_record=TestRecordSettings(template_path=template_path),
        )
        _seed_project("P1", tmp_path)
        _seed_header_metadata_sources("P1", tmp_path)

        response = client.post(
            "/api/projects/P1/matrix-editor/test-record-draft/generate",
            json={
                "source": "matrix_editor_current_ui_state",
                "groups": [
                    {
                        "group_key": "g1",
                        "group_label": "1",
                        "sample_quantity_expression": "7",
                    }
                ],
                "rows": [
                    {
                        "test_item": "Unsaved Visual Check",
                        "section": "9.9",
                        "method": "Unsaved method from UI",
                        "condition": "Unsaved condition from UI",
                        "requirement": "Unsaved requirement from UI",
                        "is_sample_row": False,
                        "group_values": {"g1": "1"},
                    }
                ],
            },
        )

        assert response.status_code == 200
        assert "Preview" in response.headers["content-disposition"]
        assert "Unconfirmed" in response.headers["content-disposition"]
        output = tmp_path / "downloaded-preview.docx"
        output.write_bytes(response.content)
        document = Document(output)
        body_text = "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        assert "Unsaved Visual Check" in body_text
        assert "Unsaved method from UI" in body_text
        assert "Unsaved condition from UI" in body_text
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_matrix_editor_test_record_generation_uses_settings_template_folder(
    tmp_path: Path,
) -> None:
    client, engine, session_factory = _client(tmp_path)
    try:
        template_folder = tmp_path / "template-folder"
        template_folder.mkdir()
        _build_template(template_folder / "FDQF-E-036 Test Record Template-Even.docx")
        _seed_template_folder(session_factory, template_folder)
        _seed_project("P1", tmp_path)
        _seed_header_metadata_sources("P1", tmp_path)

        response = client.post(
            "/api/projects/P1/matrix-editor/test-record-draft/generate",
            json={
                "source": "matrix_editor_current_ui_state",
                "groups": [
                    {
                        "group_key": "g1",
                        "group_label": "1",
                        "sample_quantity_expression": "3",
                    }
                ],
                "rows": [
                    {
                        "test_item": "Visual Check",
                        "section": "5.1",
                        "method": "EIA-364-18B",
                        "condition": "Normal",
                        "requirement": "No defect",
                        "group_values": {"g1": "1"},
                    }
                ],
            },
        )

        assert response.status_code == 200
        assert "Preview" in response.headers["content-disposition"]
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_matrix_editor_test_record_generation_rejects_wrong_source(
    tmp_path: Path,
) -> None:
    client, engine, _ = _client(tmp_path)
    try:
        response = client.post(
            "/api/projects/P1/matrix-editor/test-record-draft/generate",
            json={"source": "saved_draft", "groups": [], "rows": []},
        )
        assert response.status_code == 422
        assert response.json()["detail"] == (
            "Matrix Editor Test Record preview requires current UI state payload."
        )
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_publication_preview_keeps_download_mode_without_project_folder(
    tmp_path: Path,
) -> None:
    client, engine, _ = _client(tmp_path)
    try:
        _seed_project("P1", tmp_path)
        response = client.post(
            "/api/projects/P1/matrix-editor/test-record-publication/preview",
            json=_draft_payload("Unsaved method"),
        )

        assert response.status_code == 200
        assert response.json()["mode"] == "download"
        assert response.json()["status"] == "ready"
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_publication_writes_current_ui_draft_with_authoritative_header(
    tmp_path: Path,
) -> None:
    client, engine, session_factory = _client(tmp_path)
    try:
        template_path = _build_template(tmp_path / "template.docx")
        app.dependency_overrides[get_settings] = lambda: Settings(
            data_dir=tmp_path / "data",
            projects_dir=tmp_path / "projects",
            templates_dir=tmp_path / "templates",
            database_path=tmp_path / "connlab.sqlite3",
            test_record=TestRecordSettings(template_path=template_path),
        )
        _seed_project("P1", tmp_path)
        _seed_basic_information("P1", tmp_path)
        workspace = _seed_workspace(session_factory, tmp_path)
        payload = _draft_payload("Unsaved method from Matrix Editor")

        preview = client.post(
            "/api/projects/P1/matrix-editor/test-record-publication/preview",
            json=payload,
        )
        assert preview.status_code == 200
        assert preview.json()["status"] == "ready"

        published = client.post(
            "/api/projects/P1/matrix-editor/test-record-publication/publish",
            json={
                **payload,
                "preview_token": preview.json()["preview_token"],
                "conflict_action": "none",
            },
        )

        assert published.status_code == 200
        target = workspace.official_folder_path / "Test results" / "DL-2026-05-003 Test Record.docx"
        assert Path(published.json()["target_path"]) == target
        assert target.is_file()
        document = Document(target)
        header_text = "\n".join(
            cell.text
            for section in document.sections
            for table in section.header.tables
            for row in table.rows
            for cell in row.cells
        )
        body_text = "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        assert "DL-2026-05-003" in header_text
        assert "Confirmed Coolpower HDF 3.40mm pin" in header_text
        assert "GS-12-9999" in header_text
        assert "Unsaved method from Matrix Editor" in body_text
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_publication_archives_existing_test_record_in_workspace_history(
    tmp_path: Path,
) -> None:
    client, engine, session_factory = _client(tmp_path)
    try:
        template_path = _build_template(tmp_path / "template.docx")
        app.dependency_overrides[get_settings] = lambda: Settings(
            data_dir=tmp_path / "data",
            projects_dir=tmp_path / "projects",
            templates_dir=tmp_path / "templates",
            database_path=tmp_path / "connlab.sqlite3",
            test_record=TestRecordSettings(template_path=template_path),
        )
        _seed_project("P1", tmp_path)
        _seed_basic_information("P1", tmp_path)
        workspace = _seed_workspace(session_factory, tmp_path)
        target = workspace.official_folder_path / "Test results" / "DL-2026-05-003 Test Record.docx"
        target.write_text("operator old record", encoding="utf-8")
        payload = _draft_payload("replacement method")

        preview = client.post(
            "/api/projects/P1/matrix-editor/test-record-publication/preview",
            json=payload,
        )
        assert preview.json()["status"] == "conflict"
        published = client.post(
            "/api/projects/P1/matrix-editor/test-record-publication/publish",
            json={
                **payload,
                "preview_token": preview.json()["preview_token"],
                "conflict_action": "archive",
            },
        )

        assert published.status_code == 200
        archive = Path(published.json()["archive_path"])
        assert archive.parent == workspace.local_workspace_path / "History" / "Test Record"
        assert archive.read_text(encoding="utf-8") == "operator old record"
        assert target.is_file()
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def _seed_template_folder(session_factory, template_folder: Path) -> None:
    with session_factory() as session:
        ExternalResourceRepository(session).upsert(
            ExternalResource(
                resource_id="template-folder",
                resource_type=ExternalResourceType.PROJECT_FOLDER_TEMPLATE,
                path=template_folder,
                active=True,
                validation_status=ExternalResourceValidationStatus.VALID,
            )
        )
        session.commit()


def _draft_payload(method: str) -> dict[str, object]:
    return {
        "source": "matrix_editor_current_ui_state",
        "groups": [
            {
                "group_key": "g1",
                "group_label": "1",
                "sample_quantity_expression": "5",
            }
        ],
        "rows": [
            {
                "test_item": "Visual Check",
                "section": "5.1",
                "method": method,
                "condition": "Normal",
                "requirement": "No defect",
                "group_values": {"g1": "1"},
            }
        ],
    }


def _seed_workspace(session_factory, tmp_path: Path) -> OfficialWorkspaceRecord:
    local = tmp_path / "DL-2026-05-003"
    official = local / "DL-2026-05-003 Connector Qualification test"
    (official / "Test results").mkdir(parents=True)
    record = OfficialWorkspaceRecord(
        workspace_id="workspace-1",
        project_id="P1",
        dl_number="DL-2026-05-003",
        local_workspace_path=local,
        source_book_path=local / "Source Book",
        official_folder_path=official,
        manifest_path=local / ".connlab" / "manifest.json",
        template_source_path=tmp_path / "template-source",
        created_at="2026-08-28T00:00:00+00:00",
    )
    with session_factory() as session:
        ProjectOfficialWorkspaceRepository(session).save(record)
        session.commit()
    return record
