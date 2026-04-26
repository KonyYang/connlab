from collections.abc import Generator
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from backend.api.dependencies import get_session, get_settings
from backend.api.main import app
from backend.infrastructure.storage.database import (
    create_database_engine,
    create_session_factory,
    init_db,
)
from backend.shared.config import Settings


def test_intake_precheck_api_flow(tmp_path: Path) -> None:
    settings = Settings(
        data_dir=tmp_path / "data",
        projects_dir=tmp_path / "projects",
        templates_dir=tmp_path / "templates",
        database_path=tmp_path / "connlab.sqlite3",
    )
    engine = create_database_engine(settings)
    init_db(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        with session_factory() as session:
            try:
                yield session
                session.commit()
            except Exception:
                session.rollback()
                raise

    app.dependency_overrides[get_session] = override_session
    app.dependency_overrides[get_settings] = lambda: settings
    client = TestClient(app)

    try:
        project_response = client.post(
            "/api/projects",
            json={
                "project_no": "PRJ-001",
                "product_name": "Connector",
                "requestor": "Alice",
            },
        )
        project_id = project_response.json()["project_id"]
        docx_path = _create_docx(tmp_path / "form.docx")

        with docx_path.open("rb") as handle:
            upload_response = client.post(
                f"/api/projects/{project_id}/application-form",
                files={
                    "file": (
                        "form.docx",
                        handle,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        assert upload_response.status_code == 201
        uploaded = upload_response.json()
        assert uploaded["form_no"] == "E-3718"
        assert uploaded["samples"][0]["part_number"] == "PN-001"

        precheck_response = client.post(
            f"/api/application-forms/{uploaded['form_id']}/precheck/run"
        )
        assert precheck_response.status_code == 200
        precheck = precheck_response.json()
        assert precheck["status"] == "warning"
        assert precheck["issues"][0]["resolved"] is False

        latest_response = client.get(f"/api/projects/{project_id}/prechecks/latest")
        assert latest_response.status_code == 200
        assert latest_response.json()["result_id"] == precheck["result_id"]

        issue_id = precheck["issues"][0]["issue_id"]
        resolve_response = client.patch(f"/api/precheck-issues/{issue_id}/resolve")
        assert resolve_response.status_code == 200
        assert resolve_response.json()["resolved"] is True
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def _create_docx(path: Path) -> Path:
    document = Document()
    _add_key_value_table(
        document,
        [
            ("Form No.", "E-3718"),
            ("Form Rev", "H"),
            ("Requested By", "Alice"),
            ("Phone", "555-0100"),
            ("Date", "2026-04-26"),
            ("Email", "alice@example.com"),
            ("Business Unit", "BU-1"),
            ("Mfg. Site", "Plant 1"),
            ("Project #", "PRJ-001"),
            ("Description of Requested Testing", "依附件"),
            ("Subcontract", "No"),
            ("Estimated Completion Date", "2026-05-01"),
        ],
    )
    sample_table = document.add_table(rows=2, cols=7)
    headers = [
        "Product Name",
        "Part Number",
        "Lot/Traceability",
        "Material",
        "Plating",
        "Housing Material",
        "Quantity",
    ]
    values = ["Connector", "PN-001", "LOT-1", "Copper", "Tin", "LCP", "12"]
    for index, header in enumerate(headers):
        sample_table.cell(0, index).text = header
        sample_table.cell(1, index).text = values[index]
    document.save(path)
    return path


def _add_key_value_table(document: Document, pairs: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=len(pairs), cols=2)
    for row_index, (label, value) in enumerate(pairs):
        table.cell(row_index, 0).text = label
        table.cell(row_index, 1).text = value
