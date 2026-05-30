"""Word gateway for writing generated test-record documents."""

from __future__ import annotations

import shutil
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Cell
from docx.table import Table
from docx.text.paragraph import Paragraph

from backend.application.confirmed_matrix_test_record_document_generation_service import (
    TestRecordHeaderMetadata,
)
from backend.infrastructure.office.models import TestRecordDocumentWriteResult


class TestRecordDocumentGateway:
    """Generate test-record `.docx` files through the infrastructure boundary."""

    def generate(
        self,
        *,
        template_path: Path,
        output_path: Path,
        source_document_name: str,
        groups: tuple,
        warnings: tuple[str, ...],
    ) -> TestRecordDocumentWriteResult:
        """Copy the template and append deterministic test-record content."""
        template = Path(template_path)
        target = Path(output_path)
        if template.suffix.lower() != ".docx":
            raise ValueError(f"Only .docx template is supported: {template}")
        if not template.is_file():
            raise FileNotFoundError(f"Template does not exist: {template}")
        if not target.parent.exists():
            raise FileNotFoundError(f"Output directory does not exist: {target.parent}")

        shutil.copy2(template, target)
        document = Document(target)
        document.add_paragraph("ConnLab Generated Test Record")
        document.add_paragraph(f"Source document: {source_document_name}")
        for group in groups:
            document.add_paragraph(f"Group: {group.group_label}")
            table = document.add_table(rows=1, cols=7)
            headers = table.rows[0].cells
            headers[0].text = "Seq"
            headers[1].text = "Test Item"
            headers[2].text = "Condition"
            headers[3].text = "Method"
            headers[4].text = "Reference"
            headers[5].text = "Judgement"
            headers[6].text = "Duration"
            for step in group.steps:
                row = table.add_row().cells
                row[0].text = "" if step.sequence is None else str(step.sequence)
                row[1].text = step.test_item or ""
                row[2].text = step.condition_summary or ""
                row[3].text = step.method_summary or ""
                row[4].text = step.reference_standard or ""
                row[5].text = step.judgement_criteria or ""
                row[6].text = step.duration_hint or ""
        if warnings:
            document.add_paragraph("Warnings:")
            for warning in warnings:
                document.add_paragraph(f"- {warning}")
        document.save(target)
        return TestRecordDocumentWriteResult(
            output_path=target,
            status="generated",
            group_count=len(groups),
            warning_count=len(warnings),
            warnings=(),
        )

    def generate_from_confirmed_matrix(
        self,
        *,
        template_path: Path,
        output_path: Path,
        project_id: str,
        project_no: str,
        product_description: str,
        applicable_specification: str,
        confirmed_matrix_id: str,
        groups: tuple,
        header_metadata: TestRecordHeaderMetadata,
    ) -> Path:
        """Generate a template-backed Test Record draft from active ConfirmedMatrix data."""
        template = Path(template_path)
        target = Path(output_path)
        if template.suffix.lower() != ".docx":
            raise ValueError(f"Only .docx template is supported: {template}")
        if not template.is_file():
            raise FileNotFoundError(f"Template does not exist: {template}")
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(template, target)
        document = Document(target)
        _fill_header_metadata(document, header_metadata)
        if len(document.tables) < 2:
            raise ValueError("Template must contain at least two tables.")
        if not groups:
            raise ValueError("No selected groups are available for Test Record generation.")

        template_step_table = document.tables[0]
        template_equipment_table = document.tables[1]
        template_group_paragraph = _paragraph_before_table(document, template_step_table)
        template_block = _collect_block_elements(
            start=template_group_paragraph._p,
            end=template_equipment_table._tbl,
        )
        template_block_clone = [deepcopy(element) for element in template_block]

        first_group = groups[0]
        self._fill_group_block(
            group_paragraph=template_group_paragraph,
            step_table=template_step_table,
            equipment_table=template_equipment_table,
            group=first_group,
        )

        insert_after = template_equipment_table._tbl
        for group in groups[1:]:
            page_break = _build_page_break_paragraph_element()
            insert_after.addnext(page_break)
            insert_after = page_break

            cloned_elements: list[object] = []
            for element in template_block_clone:
                cloned = deepcopy(element)
                insert_after.addnext(cloned)
                insert_after = cloned
                cloned_elements.append(cloned)

            group_paragraph_element = cloned_elements[0]
            table_elements = [
                element for element in cloned_elements if _element_tag_name(element) == "tbl"
            ]
            if len(table_elements) < 2:
                raise ValueError("Template block must include step and equipment tables.")
            step_table_element = table_elements[0]
            equipment_table_element = table_elements[1]

            group_paragraph = Paragraph(group_paragraph_element, document._body)
            step_table = Table(step_table_element, document._body)
            equipment_table = Table(equipment_table_element, document._body)
            self._fill_group_block(
                group_paragraph=group_paragraph,
                step_table=step_table,
                equipment_table=equipment_table,
                group=group,
            )
        document.save(target)
        return target

    def _fill_group_block(
        self,
        *,
        group_paragraph: Paragraph,
        step_table: Table,
        equipment_table: Table,
        group,
    ) -> None:
        group_label = str(getattr(group, "group_label", "") or "").strip() or "-"
        group_number = _normalize_group_number_label(group_label)
        sample_expression = str(
            getattr(group, "sample_quantity_expression", "") or ""
        ).strip()
        sample_display = _format_sample_quantity(group_label, sample_expression)
        group_line = (
            f"Group Number 组别编号: {group_number} ;   "
            f"Sample Quantity & Number 样品数量及编号: {sample_display}"
        )
        _update_group_paragraph(
            group_paragraph,
            group_line=group_line,
            group_number=group_number,
            sample_display=sample_display,
        )

        ordered_steps = sorted(getattr(group, "steps", ()), key=_step_sort_key)
        rows = _prepare_step_rows(step_table, len(ordered_steps))
        for row, step in zip(rows, ordered_steps, strict=True):
            if len(row) < 9:
                raise ValueError("Template step table must contain at least 9 columns.")
            _set_cell_text_preserve_format(row[0], _step_display_text(step))
            _set_cell_text_preserve_format(row[1], _as_text(getattr(step, "test_item", "")))
            _set_cell_text_preserve_format(row[2], _as_text(getattr(step, "method", "")))
            _set_cell_text_preserve_format(row[3], _as_text(getattr(step, "condition", "")))
            _set_cell_text_preserve_format(row[4], "")
            _set_cell_text_preserve_format(row[5], "")
            _set_cell_text_preserve_format(row[6], "")
            _set_cell_text_preserve_format(row[7], "")
            _set_cell_text_preserve_format(row[8], _as_text(getattr(step, "requirement", "")))


def _paragraph_before_table(document: Document, table: Table) -> Paragraph:
    table_element = table._tbl
    paragraph = table_element.getprevious()
    while paragraph is not None and paragraph.tag.rsplit("}", 1)[-1] != "p":
        paragraph = paragraph.getprevious()
    if paragraph is None:
        return document.paragraphs[0] if document.paragraphs else document.add_paragraph("")
    return Paragraph(paragraph, document._body)


def _collect_block_elements(*, start: object, end: object) -> list[object]:
    elements: list[object] = [start]
    current = start
    while current is not end:
        current = current.getnext()
        if current is None:
            raise ValueError("Template block boundaries are invalid.")
        elements.append(current)
    return elements


def _element_tag_name(element: object) -> str:
    return str(element.tag).rsplit("}", 1)[-1]


def _build_page_break_paragraph_element() -> object:
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run.append(page_break)
    paragraph.append(run)
    return paragraph


def _fill_header_metadata(document: Document, metadata: TestRecordHeaderMetadata) -> None:
    for section in document.sections:
        header = section.header
        _fill_lab_test_request_number(
            header_tables=header.tables,
            value=metadata.lab_test_request_number,
        )
        _fill_header_value_by_label(
            header_tables=header.tables,
            label_tokens=("Product Description", "产品描述"),
            value=metadata.product_description,
        )
        _fill_header_value_by_label(
            header_tables=header.tables,
            label_tokens=("Applicable Specification", "Applicable Specifications", "适用的规范"),
            value=metadata.applicable_specification,
        )
        _fill_header_value_by_label(
            header_tables=header.tables,
            label_tokens=("Estimated Completion Date", "预计完成日期"),
            value="",
        )


def _fill_lab_test_request_number(*, header_tables: list[Table], value: str) -> None:
    match = _find_label_cell(
        tables=header_tables,
        label_tokens=("Lab Test Request Number", "实验室测试项目编号"),
    )
    if match is None:
        return
    table_index, row_index, cell_index = match
    cell = header_tables[table_index].rows[row_index].cells[cell_index]
    _set_labeled_cell_value(cell, value=value)


def _fill_header_value_by_label(
    *,
    header_tables: list[Table],
    label_tokens: tuple[str, ...],
    value: str,
) -> None:
    match = _find_label_cell(tables=header_tables, label_tokens=label_tokens)
    if match is None:
        return
    table_index, row_index, cell_index = match
    row_cells = header_tables[table_index].rows[row_index].cells
    value_index = cell_index + 1
    if value_index >= len(row_cells):
        return
    _set_cell_text_preserve_format(row_cells[value_index], value.strip())


def _find_label_cell(
    *,
    tables: list[Table],
    label_tokens: tuple[str, ...],
) -> tuple[int, int, int] | None:
    for table_index, table in enumerate(tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                text = _normalize_match_text(cell.text)
                if not text:
                    continue
                if any(_normalize_match_text(token) in text for token in label_tokens):
                    return (table_index, row_index, cell_index)
    return None


def _set_labeled_cell_value(cell: _Cell, *, value: str) -> None:
    paragraphs = list(cell.paragraphs)
    label_indexes = [
        index
        for index, paragraph in enumerate(paragraphs)
        if _is_ltr_label_paragraph(paragraph.text)
    ]
    chinese_label_indexes = [
        index
        for index, paragraph in enumerate(paragraphs)
        if "实验室测试项目编号" in _normalize_match_text(paragraph.text)
    ]
    has_separate_labels = len(label_indexes) >= 2

    for index in label_indexes:
        paragraph = paragraphs[index]
        if has_separate_labels and _contains_ltr_number(paragraph.text):
            _set_paragraph_text_preserve_format(paragraph, "")

    target = _find_ltr_value_paragraph(
        paragraphs=paragraphs,
        label_indexes=label_indexes,
        chinese_label_indexes=chinese_label_indexes,
    )
    if target is None:
        target = cell.add_paragraph("")
    _set_paragraph_text_preserve_format(target, value.strip())
    _apply_ltr_value_format(target)


def _find_ltr_value_paragraph(
    *,
    paragraphs: list[Paragraph],
    label_indexes: list[int],
    chinese_label_indexes: list[int],
) -> Paragraph | None:
    # Preferred shape in real template: Chinese label paragraph then value paragraph.
    if chinese_label_indexes:
        candidate_index = max(chinese_label_indexes) + 1
        if candidate_index < len(paragraphs):
            return paragraphs[candidate_index]

    for index, paragraph in enumerate(paragraphs):
        if index in label_indexes:
            continue
        text = paragraph.text.strip()
        if not text or _contains_ltr_number(text):
            return paragraph
    return None


def _is_ltr_label_paragraph(text: str) -> bool:
    normalized = _normalize_match_text(text)
    return (
        "lab test request number" in normalized
        or "实验室测试项目编号" in normalized
    )


def _contains_ltr_number(text: str) -> bool:
    return re.search(r"\bDL-\d{4}-\d{2}-\d{3,}\b", text, flags=re.IGNORECASE) is not None


def _set_paragraph_text_preserve_format(paragraph: Paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.add_run(value)


def _apply_ltr_value_format(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.font.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        r_pr = run._r.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is not None:
            r_fonts.set(qn("w:ascii"), "Arial")
            r_fonts.set(qn("w:hAnsi"), "Arial")


def _normalize_match_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip().casefold()


def _clear_data_rows(table: Table) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)


def _prepare_step_rows(table: Table, row_count: int) -> list[list[_Cell]]:
    if row_count <= 0:
        _clear_data_rows(table)
        return []
    if len(table.rows) < 2:
        if not table.rows:
            raise ValueError("Template step table must contain at least one header row.")
        prototype = deepcopy(table.add_row()._tr)
    else:
        prototype = deepcopy(table.rows[1]._tr)
    _clear_data_rows(table)
    for _ in range(row_count):
        table._tbl.append(deepcopy(prototype))
    return [table.rows[index + 1].cells for index in range(row_count)]


def _set_cell_text_preserve_format(cell: _Cell, value: str) -> None:
    if not cell.paragraphs:
        cell.text = value
        return
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.add_run(value)


def _update_group_paragraph(
    paragraph: Paragraph,
    *,
    group_line: str,
    group_number: str,
    sample_display: str,
) -> None:
    if not paragraph.runs:
        paragraph.text = group_line
        return
    placeholders = sum(run.text.count("#") for run in paragraph.runs)
    if placeholders >= 2:
        replacements = [group_number, sample_display]
        replacement_index = 0
        for run in paragraph.runs:
            if replacement_index >= len(replacements) or "#" not in run.text:
                continue
            while "#" in run.text and replacement_index < len(replacements):
                run.text = run.text.replace("#", replacements[replacement_index], 1)
                replacement_index += 1
            if replacement_index >= len(replacements):
                break
        return
    paragraph.runs[0].text = group_line
    for run in paragraph.runs[1:]:
        run.text = ""


def _as_text(value: object) -> str:
    if value is None:
        return ""
    return str(value)


def _format_sample_quantity(group_label: str, expression: str) -> str:
    raw = expression.strip()
    match = re.search(r"\d+", raw)
    if match is None:
        return raw
    number = match.group(0)
    normalized = _normalize_group_number_label(group_label)
    quantity = f"{raw} sets" if re.fullmatch(r"\d+", raw) else raw
    return f"{quantity} (Group{normalized}-1#~{number}#)"


def _normalize_group_number_label(group_label: str) -> str:
    compact = re.sub(r"\s+", "", group_label).strip()
    if not compact:
        return "-"
    prefixed = re.match(r"^(?:group|g)(.+)$", compact, flags=re.IGNORECASE)
    value = prefixed.group(1) if prefixed else compact
    return value or "-"


def _step_display_text(step: object) -> str:
    raw_token = _as_text(getattr(step, "raw_token", "")).strip()
    if raw_token:
        return raw_token
    return _as_text(getattr(step, "sequence", None))


def _step_sort_key(step: object) -> tuple[int, int, str]:
    display = _step_display_text(step)
    match = re.search(r"\d+", display)
    if match is not None:
        return (0, int(match.group(0)), display.casefold())
    sequence = getattr(step, "sequence", None)
    if isinstance(sequence, int):
        return (0, sequence, display.casefold())
    return (1, 0, display.casefold())
