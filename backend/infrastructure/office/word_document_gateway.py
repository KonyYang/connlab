"""Word document reading gateway for the Office integration boundary."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from backend.infrastructure.office.application_form_word_gateway import (
    application_form_requires_com,
    drop_duplicate_application_form_aliases,
    normalize_label,
    write_application_form_fields_with_com,
)
from backend.infrastructure.office.application_form_word_mapping import (
    APPLICATION_FORM_CRITICAL_FIELDS,
    APPLICATION_FORM_FIELD_LABELS,
    APPLICATION_FORM_NEXT_ROW_FIELDS,
)
from backend.infrastructure.office.models import (
    WordDocumentSnapshot,
    WordHeaderCellResult,
    WordSection2FieldChange,
    WordSection2WriteResult,
    WordTableLocation,
)
from backend.infrastructure.office.office_lifecycle import OfficeAutomationUnavailable
from backend.infrastructure.office.word_numbering import paragraph_texts_with_numbering


SECTION2_FIELD_LABELS: dict[str, tuple[str, ...]] = {
    "lab": ("lab", "laboratory", "lab performing the tests"),
    "assigned_personnel": (
        "assigned personnel",
        "assigned engineer",
        "test engineer",
        "tested by",
    ),
    "received_date": ("received date", "sample received date"),
    "estimated_completion_date": (
        "estimated completion date",
        "estimated complete date",
    ),
    "sample_condition": ("sample condition", "sample received condition"),
}


class WordDocumentGateway:
    """Read Word documents into neutral snapshots without business parsing."""

    def read_word_document(self, source_path: Path) -> WordDocumentSnapshot:
        """Read a `.docx` file into paragraphs, tables, headers, and footers."""
        path = Path(source_path)
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Only .docx files are supported by the Word gateway: {path}")
        if not path.is_file():
            raise FileNotFoundError(f"Word document does not exist: {path}")

        document = Document(path)
        paragraphs = paragraph_texts_with_numbering(document)
        tables = [_table_rows(table) for table in document.tables]
        headers = _section_text(document, part_name="header")
        footers = _section_text(document, part_name="footer")
        raw_text = _raw_text(paragraphs, tables, headers, footers)
        return WordDocumentSnapshot(
            paragraphs=[text for text in paragraphs if text],
            tables=tables,
            headers=headers,
            footers=footers,
            raw_text=raw_text,
        )

    def read_table_locations(self, source_path: Path) -> tuple[WordTableLocation, ...]:
        """Read table page/location metadata through Microsoft Word COM."""
        path = Path(source_path)
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Only .docx files are supported by Word table location: {path}")
        if not path.is_file():
            raise FileNotFoundError(f"Word document does not exist: {path}")
        return _read_table_locations_with_com(path)

    def export_preview_pdf(self, source_path: Path, output_pdf_path: Path) -> Path:
        """Export one `.docx` source file to a PDF preview via Word COM."""
        path = Path(source_path)
        output = Path(output_pdf_path)
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Only .docx files are supported by Word PDF export: {path}")
        if not path.is_file():
            raise FileNotFoundError(f"Word document does not exist: {path}")
        if output.suffix.lower() != ".pdf":
            raise ValueError(f"Preview output must be a .pdf file: {output}")
        output.parent.mkdir(parents=True, exist_ok=True)
        _export_pdf_with_com(path, output)
        return output

    def read_header_table_cell(
        self,
        source_path: Path,
        row: int,
        column: int,
    ) -> WordHeaderCellResult:
        """Read a Word header table cell using COM first and python-docx fallback."""
        path = Path(source_path)
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Only .docx files are supported by the Word header gate: {path}")
        if not path.is_file():
            raise FileNotFoundError(f"Word document does not exist: {path}")
        value = _read_header_cell_with_python_docx(path, row, column)
        if value is not None:
            return WordHeaderCellResult(
                value=_clean(value or ""),
                gateway_mode="python_docx",
            )
        value = _read_header_cell_with_com(path, row, column)
        return WordHeaderCellResult(value=_clean(value or ""), gateway_mode="word_com")

    def write_section2_fields(
        self,
        source_path: Path,
        fields: dict[str, str],
    ) -> WordSection2WriteResult:
        """Write known Section 2 fields into adjacent Word table value cells."""
        path = Path(source_path)
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Only .docx files are supported by the Word gateway: {path}")
        if not path.is_file():
            raise FileNotFoundError(f"Word document does not exist: {path}")
        unknown = sorted(key for key in fields if key not in SECTION2_FIELD_LABELS)
        if unknown:
            raise ValueError(f"Unsupported Section 2 field(s): {', '.join(unknown)}")

        document = Document(path)
        locations = _locate_section2_fields(document, fields)
        missing = [key for key in fields if key not in locations]
        if missing:
            raise ValueError(
                "Section 2 field location(s) not found: " + ", ".join(sorted(missing))
            )

        changed: list[WordSection2FieldChange] = []
        unchanged: list[WordSection2FieldChange] = []
        for field_key, new_value in fields.items():
            table_index, row_index, label_column, value_column = locations[field_key]
            row = document.tables[table_index].rows[row_index]
            label = _clean(row.cells[label_column].text)
            cell = row.cells[value_column]
            old_value = _clean(cell.text)
            update = WordSection2FieldChange(
                field_key=field_key,
                label=label,
                old_value=old_value,
                new_value=new_value,
                location=(
                    f"table[{table_index}].row[{row_index}].cell[{value_column}]"
                ),
            )
            if old_value == new_value:
                unchanged.append(update)
            else:
                cell.text = new_value
                changed.append(update)

        document.save(path)
        return WordSection2WriteResult(
            changed_fields=tuple(changed),
            unchanged_fields=tuple(unchanged),
        )

    def write_application_form_fields(
        self,
        source_path: Path,
        fields: dict[str, str],
    ) -> WordSection2WriteResult:
        """Write known application fields into visible Word form/table targets."""
        path = Path(source_path)
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Only .docx files are supported by the Word gateway: {path}")
        if not path.is_file():
            raise FileNotFoundError(f"Word document does not exist: {path}")
        normalized_fields = {
            key: str(value)
            for key, value in fields.items()
            if key in APPLICATION_FORM_FIELD_LABELS and value is not None
        }
        drop_duplicate_application_form_aliases(normalized_fields)
        if application_form_requires_com(path):
            return write_application_form_fields_with_com(path, normalized_fields)

        document = Document(path)
        locations = _locate_labeled_fields(
            document,
            normalized_fields,
            APPLICATION_FORM_FIELD_LABELS,
            next_row_fields=APPLICATION_FORM_NEXT_ROW_FIELDS,
        )

        changed: list[WordSection2FieldChange] = []
        unchanged: list[WordSection2FieldChange] = []
        warnings: list[str] = []
        for field_key, new_value in normalized_fields.items():
            location = locations.get(field_key)
            if location is None:
                message = f"Application Form field location not found: {field_key}"
                if field_key in APPLICATION_FORM_CRITICAL_FIELDS:
                    raise ValueError(message)
                warnings.append(message)
                continue
            table_index, row_index, label_column, value_column = location
            row = document.tables[table_index].rows[row_index]
            label = _clean(row.cells[label_column].text)
            cell = row.cells[value_column]
            old_value = _clean(cell.text)
            update = WordSection2FieldChange(
                field_key=field_key,
                label=label,
                old_value=old_value,
                new_value=new_value,
                location=(
                    f"table[{table_index}].row[{row_index}].cell[{value_column}]"
                ),
            )
            if old_value == new_value:
                unchanged.append(update)
            else:
                cell.text = new_value
                changed.append(update)

        if changed:
            document.save(path)
        return WordSection2WriteResult(
            changed_fields=tuple(changed),
            unchanged_fields=tuple(unchanged),
            warnings=tuple(warnings),
        )

def _table_rows(table) -> list[list[str]]:
    return [[_clean(cell.text) for cell in row.cells] for row in table.rows]


def _section_text(document, *, part_name: str) -> list[str]:
    values: list[str] = []
    for section in document.sections:
        part = getattr(section, part_name)
        values.extend(_clean(paragraph.text) for paragraph in part.paragraphs)
        for table in part.tables:
            for row in _table_rows(table):
                values.extend(row)
    return [value for value in values if value]


def _read_header_cell_with_com(path: Path, row: int, column: int) -> str | None:
    """Read the first matching header table cell through Microsoft Word COM."""
    try:
        import win32com.client  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - depends on Windows host
        raise OfficeAutomationUnavailable("Word COM automation requires pywin32.") from exc

    word = win32com.client.DispatchEx("Word.Application")
    document = None
    try:
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(path),
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        for section in _com_iter(document.Sections):
            for header in _com_iter(section.Headers):
                tables = getattr(header, "Range", header).Tables
                for table in _com_iter(tables):
                    try:
                        return str(table.Cell(row, column).Range.Text)
                    except Exception:
                        continue
        return None
    finally:
        if document is not None:
            document.Close(SaveChanges=False)
        word.Quit()


def _read_table_locations_with_com(path: Path) -> tuple[WordTableLocation, ...]:
    """Return table layout metadata through Microsoft Word COM."""
    try:
        import pythoncom  # type: ignore[import-not-found]
        import win32com.client  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - depends on Windows host
        raise OfficeAutomationUnavailable("Word COM automation requires pywin32.") from exc

    wd_active_end_page_number = 3
    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(path),
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        page_counts: dict[int, int] = {}
        locations: list[WordTableLocation] = []
        for table_index, table in enumerate(_com_iter(document.Tables), start=1):
            page_number = _safe_int(table.Range.Information(wd_active_end_page_number))
            page_table_index = None
            if page_number is not None:
                page_counts[page_number] = page_counts.get(page_number, 0) + 1
                page_table_index = page_counts[page_number]
            locations.append(
                WordTableLocation(
                    table_index=table_index,
                    page_number=page_number,
                    page_table_index=page_table_index,
                    preceding_paragraph=_clean(_preceding_paragraph_text(table)),
                    text_preview=_clean(str(table.Range.Text).replace("\r", " "))[:240],
                    row_count=_safe_int(getattr(table.Rows, "Count", 0)) or 0,
                    column_count=_safe_int(getattr(table.Columns, "Count", 0)) or 0,
                )
            )
        return tuple(locations)
    finally:
        if document is not None:
            document.Close(SaveChanges=False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def _export_pdf_with_com(source_path: Path, output_pdf_path: Path) -> None:
    """Export one Word document to PDF through Microsoft Word COM."""
    try:
        import pythoncom  # type: ignore[import-not-found]
        import win32com.client  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - depends on Windows host
        raise OfficeAutomationUnavailable("Word COM automation requires pywin32.") from exc

    source = source_path.resolve()
    output = output_pdf_path.resolve()
    wd_export_format_pdf = 17
    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(source),
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        document.ExportAsFixedFormat(
            OutputFileName=str(output),
            ExportFormat=wd_export_format_pdf,
            OpenAfterExport=False,
        )
    finally:
        if document is not None:
            document.Close(SaveChanges=False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def _preceding_paragraph_text(table) -> str:
    try:
        paragraph_range = table.Range.Duplicate
        paragraph_range.Start = max(0, int(table.Range.Start) - 1)
        paragraph_range.End = int(table.Range.Start)
        paragraph_range.MoveStart(Unit=4, Count=-1)
        return str(paragraph_range.Text)
    except Exception:
        return ""


def _safe_int(value: object) -> int | None:
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _com_iter(collection) -> list[object]:
    count = int(getattr(collection, "Count", 0))
    return [collection.Item(index) for index in range(1, count + 1)]


def _read_header_cell_with_python_docx(path: Path, row: int, column: int) -> str | None:
    """Read the first matching header table cell through python-docx."""
    document = Document(path)
    row_index = row - 1
    column_index = column - 1
    for section in document.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            for table in header.tables:
                if len(table.rows) <= row_index:
                    continue
                cells = table.rows[row_index].cells
                if len(cells) <= column_index:
                    continue
                return cells[column_index].text
    return None


def _raw_text(
    paragraphs: list[str],
    tables: list[list[list[str]]],
    headers: list[str],
    footers: list[str],
) -> str:
    lines: list[str] = []
    lines.extend(headers)
    lines.extend(paragraphs)
    for table in tables:
        for row in table:
            lines.append(" | ".join(cell for cell in row if cell))
    lines.extend(footers)
    return "\n".join(line for line in lines if line)


def _clean(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\x07", " ")).strip()


def _locate_section2_fields(document, fields: dict[str, str]) -> dict[str, tuple[int, int, int, int]]:
    return _locate_labeled_fields(document, fields, SECTION2_FIELD_LABELS)


def _locate_labeled_fields(
    document,
    fields: dict[str, str],
    aliases_by_field: dict[str, tuple[str, ...]],
    *,
    next_row_fields: set[str] | None = None,
) -> dict[str, tuple[int, int, int, int]]:
    locations: dict[str, tuple[int, int, int, int]] = {}
    requested = set(fields)
    next_row_fields = next_row_fields or set()
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            cells = row.cells
            for column_index, cell in enumerate(cells):
                field_key = _field_key_for_label(
                    cell.text,
                    requested - locations.keys(),
                    aliases_by_field,
                )
                if field_key is None:
                    continue
                if field_key in next_row_fields:
                    if row_index + 1 >= len(table.rows):
                        continue
                    value_cells = table.rows[row_index + 1].cells
                    if column_index >= len(value_cells):
                        continue
                    locations[field_key] = (
                        table_index,
                        row_index + 1,
                        column_index,
                        column_index,
                    )
                    continue
                value_column = _value_column(cells, column_index)
                if value_column is None:
                    continue
                locations[field_key] = (
                    table_index,
                    row_index,
                    column_index,
                    value_column,
                )
    return locations


def _field_key_for_label(
    label: str,
    candidates: set[str],
    aliases_by_field: dict[str, tuple[str, ...]],
) -> str | None:
    normalized = normalize_label(label)
    if not normalized:
        return None
    for field_key in candidates:
        aliases = aliases_by_field[field_key]
        normalized_aliases = {normalize_label(alias) for alias in aliases}
        if normalized in normalized_aliases:
            return field_key
    return None


def _value_column(cells, label_column: int) -> int | None:
    if label_column + 1 < len(cells):
        return label_column + 1
    return None
