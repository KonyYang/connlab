"""E-3707_H Word adapter for initialization-report drafts."""

from __future__ import annotations

from collections import OrderedDict
from copy import deepcopy
from datetime import datetime
import os
from pathlib import Path
import re
import shutil
from uuid import uuid4

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from backend.application.confirmed_matrix_test_record_preview_service import (
    is_llcr_test_item,
)
from backend.application.test_report_draft_service import TestReportDraftData


_SAMPLE_HEADERS = ("Description", "Part #")
_METHOD_HEADERS = ("Item", "Test Method", "Condition", "Requirement")
_RESULT_HEADERS = (
    "Step",
    "Test",
    "Requirement",
    "Step Description",
    "Result",
    "Comment",
)
_EQUIPMENT_HEADERS = (
    "Item",
    "Manufacturer",
    "ID Number",
    "Last Cal.",
    "Cal. Due",
)
_REVISION_HEADERS = (
    "Revision Level",
    "Affected Pages",
    "Description",
    "Revision Date",
)
_REQUIRED_HEADINGS = (
    "1. PURPOSE",
    "2. CONCLUSIONS",
    "3. SAMPLE DESCRIPTION",
    "4. TEST DESCRIPTION",
    "5. TEST METHODS/REQUIREMENTS",
    "6. TEST RESULTS",
    "7. EQUIPMENTS",
    "8. REVISION RECORD",
    "*** End of Report ***",
)
_REQUIRED_FIRST_PAGE_HEADER_PLACEHOLDERS = (
    "WW-XXXX-YY-ZZZ",
    "DDMMMYYYY",
    "DDMMMYYYY-DDMMMYYYY",
    "Name(s)",
    "NAME",
    "PRODUCT NAME/TEST DESCRIPTION ",
)
_TABLE_FONT_NAME = "Arial"
_HEADER_FILL = "B2B2B2"
_SAMPLE_SIZE_FILL = "8DB3E2"
_PREFERRED_TEST_ITEM_WIDTH_DXA = 3024
_MIN_TEST_ITEM_WIDTH_DXA = 2500
_MIN_GROUP_WIDTH_DXA = 690


class TestReportDocumentGateway:
    """Populate a copied E-3707_H template through semantic document anchors."""

    def generate(
        self,
        *,
        template_path: Path,
        output_path: Path,
        report: TestReportDraftData,
    ) -> Path:
        """Write one initialization draft while retaining the approved source."""
        template = Path(template_path)
        target = Path(output_path)
        if template.suffix.lower() != ".docx":
            raise ValueError(f"Only .docx report templates are supported: {template}")
        if not template.is_file():
            raise FileNotFoundError(f"Test report template does not exist: {template}")
        if template.resolve() == target.resolve():
            raise ValueError("Approved report template cannot be used as the output path.")
        if not target.parent.is_dir():
            raise FileNotFoundError(f"Output directory does not exist: {target.parent}")
        if target.exists() and target.stat().st_size:
            raise FileExistsError(f"Output file already exists and will not be replaced: {target}")

        temporary = target.with_name(
            f".{target.stem}.{uuid4().hex}.tmp{target.suffix}"
        )
        try:
            shutil.copy2(template, temporary)
            document = Document(temporary)
            anchors = _validate_template_contract(document)
            _fill_headers(document, report)
            _fill_narrative(document, report)
            _fill_sample_table(anchors.sample_table, report)
            _fill_test_description_table(anchors.description_table, report)
            _fill_method_table(anchors.method_table, report)
            _fill_result_blocks(
                document,
                anchors.result_heading,
                anchors.result_table,
                report,
            )
            _keep_heading_with_following_content(document, "4. TEST DESCRIPTION")
            _insert_page_break_before_heading(document, "7. EQUIPMENTS")
            _fill_revision_table(anchors.revision_table, report)
            _set_document_table_font(document, _TABLE_FONT_NAME)
            document.save(temporary)
            _audit_generated_document(temporary, report)
            os.replace(temporary, target)
        finally:
            temporary.unlink(missing_ok=True)
        return target


class _TemplateAnchors:
    def __init__(
        self,
        *,
        sample_table: Table,
        description_table: Table,
        method_table: Table,
        result_heading: Paragraph,
        result_table: Table,
        revision_table: Table,
    ) -> None:
        self.sample_table = sample_table
        self.description_table = description_table
        self.method_table = method_table
        self.result_heading = result_heading
        self.result_table = result_table
        self.revision_table = revision_table


def _validate_template_contract(document, *, populated: bool = False) -> _TemplateAnchors:
    paragraphs = {_heading_key(paragraph.text): paragraph for paragraph in document.paragraphs}
    for heading in _REQUIRED_HEADINGS:
        if _heading_key(heading) not in paragraphs:
            raise ValueError(f"E-3707_H template heading is missing: {heading}")

    sample = _find_table(document, _SAMPLE_HEADERS, "Sample Description table")
    description = _find_table(document, ("Test Items",), "Test Description table")
    methods = _find_table(document, _METHOD_HEADERS, "Test Methods/Requirements table")
    results = _find_table(document, _RESULT_HEADERS, "Test Results table")
    _find_table(document, _EQUIPMENT_HEADERS, "Equipment table")
    revision = _find_table(document, _REVISION_HEADERS, "Revision Record table")
    result_heading = paragraphs.get(_heading_key("Group # Test Results"))
    if result_heading is None and populated:
        result_heading = next(
            (
                paragraph
                for paragraph in document.paragraphs
                if re.fullmatch(r"Group\s+.+\s+Test Results", _normalized(paragraph.text))
            ),
            None,
        )
    if result_heading is None:
        raise ValueError("E-3707_H template result-group heading is missing.")
    if not document.sections or not any(
        section.header.tables for section in document.sections
    ):
        raise ValueError("E-3707_H template report header is missing.")
    if not populated:
        first_page_text = "".join(
            _header_text_nodes(document.sections[0].first_page_header)
        )
        for placeholder in _REQUIRED_FIRST_PAGE_HEADER_PLACEHOLDERS:
            if placeholder not in first_page_text:
                raise ValueError(
                    "E-3707_H first-page header placeholder is missing: "
                    f"{placeholder}"
                )
    return _TemplateAnchors(
        sample_table=sample,
        description_table=description,
        method_table=methods,
        result_heading=result_heading,
        result_table=results,
        revision_table=revision,
    )


def _find_table(document, expected_headers: tuple[str, ...], label: str) -> Table:
    expected = tuple(_normalized(value) for value in expected_headers)
    for table in document.tables:
        if not table.rows:
            continue
        actual = tuple(_normalized(cell.text) for cell in table.rows[0].cells)
        if actual[: len(expected)] == expected:
            return table
    raise ValueError(f"E-3707_H {label} does not match the approved table contract.")


def _fill_headers(document, report: TestReportDraftData) -> None:
    report_number = report.report_number
    project_leader = report.project_leader or "[To be assigned]"
    replacements = {
        "DDMMMYYYY-DDMMMYYYY": _display_test_date_range(
            report.start_test_date,
            report.finish_test_date,
        ),
        "WW-XXXX-YY-ZZZ": report_number,
        "XX-YY-ZZZ": report_number,
        "DDMMMYYYY": report.generated_on.strftime("%d/%b/%Y"),
        "Name(s)": project_leader,
        "NAME": report.requestor or "[To be confirmed]",
        "PRODUCT NAME/TEST DESCRIPTION ": (
            f"{report.product_name} {report.test_description}"
        ),
        "Name": project_leader,
        "(s)": "",
    }
    for section in document.sections:
        for header in (
            section.header,
            section.first_page_header,
        ):
            seen_cells: set[object] = set()
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_identity = cell._tc
                        if cell_identity in seen_cells:
                            continue
                        seen_cells.add(cell_identity)
                        updated = _replace_header_text(cell.text, replacements)
                        if updated != cell.text:
                            _set_cell_text(cell, updated)
            for text_node in header._element.xpath(".//w:t"):
                value = text_node.text or ""
                text_node.text = _replace_header_text(value, replacements)


def _replace_header_text(value: str, replacements: dict[str, str]) -> str:
    if value in replacements:
        return replacements[value]
    updated = value
    for placeholder in (
        "DDMMMYYYY-DDMMMYYYY",
        "WW-XXXX-YY-ZZZ",
        "XX-YY-ZZZ",
        "DDMMMYYYY",
    ):
        updated = updated.replace(placeholder, replacements[placeholder])
    return updated


def _fill_narrative(document, report: TestReportDraftData) -> None:
    received_date = _display_date(report.received_samples_date)
    replacements = {
        "[TEST DESCRIPTION]": report.test_description,
        "[PRODUCT NAME]": report.product_name,
        "[GS-XX-XXXX (Rev.X, DATE)]": report.applicable_specification,
        "[RECEIVED SAMPLES DATE]": received_date,
        "DDMMMYYYY": received_date,
    }
    purpose_heading_seen = False
    conclusion_heading_seen = False
    for paragraph in document.paragraphs:
        normalized = _heading_key(paragraph.text)
        if normalized == _heading_key("1. PURPOSE"):
            purpose_heading_seen = True
            continue
        if normalized == _heading_key("2. CONCLUSIONS"):
            conclusion_heading_seen = True
            continue
        if purpose_heading_seen and normalized:
            _set_paragraph_text(
                paragraph,
                (
                    f"This report summarizes the {report.test_description} conducted on "
                    f"{report.product_name} to assess the conformance to AFCI product "
                    f"specification {report.applicable_specification}."
                ),
            )
            purpose_heading_seen = False
            continue
        if conclusion_heading_seen and normalized:
            _set_paragraph_text(
                paragraph,
                "DRAFT — Conclusions will be completed after testing and review.",
            )
            conclusion_heading_seen = False
            continue
        updated = paragraph.text
        for placeholder, value in replacements.items():
            updated = updated.replace(placeholder, value or "[To be confirmed]")
        if updated != paragraph.text:
            _set_paragraph_text(paragraph, updated)


def _fill_sample_table(table: Table, report: TestReportDraftData) -> None:
    sample_rows = report.sample_rows
    if not sample_rows:
        _resize_rows(table, 2)
        fallback_values = [
            report.product_name,
            report.description_part_number,
            "",
            "",
            "",
            "",
            "",
        ]
        for cell, value in zip(
            table.rows[1].cells,
            fallback_values,
            strict=False,
        ):
            _set_cell_text(cell, value)
        return

    _resize_rows(table, 1 + len(sample_rows))
    for table_row, sample in zip(table.rows[1:], sample_rows, strict=True):
        values = [
            sample.product_name,
            sample.part_number,
            sample.lot_or_traceability,
            sample.material,
            sample.plating,
            sample.lubricant,
            sample.housing_material,
        ]
        for cell, value in zip(table_row.cells, values, strict=False):
            _set_cell_text(cell, value)


def _fill_test_description_table(table: Table, report: TestReportDraftData) -> None:
    groups = report.groups
    table_width_dxa = _table_width_dxa(table)
    item_names: list[str] = []
    for group in groups:
        for step in group.steps:
            if step.test_item not in item_names:
                item_names.append(step.test_item)

    target_columns = len(groups) + 1
    _resize_columns(table, target_columns)
    _resize_rows(table, 2 + len(item_names) + 1)
    _set_cell_text(table.cell(0, 0), "Test Items")
    _set_cell_text(table.cell(1, 0), "")
    for column_index, group in enumerate(groups, start=1):
        _set_cell_text(
            table.cell(0, column_index),
            "Test Sequence" if column_index == 1 else "",
        )
        _set_cell_text(
            table.cell(1, column_index),
            _test_sequence_group_label(group.group_label),
        )
    if len(groups) > 1:
        table.cell(0, 1).merge(table.cell(0, len(groups)))
    table.cell(0, 0).merge(table.cell(1, 0))
    _set_cell_text(table.cell(0, 0), "Test Items")
    _set_cell_text(table.cell(0, 1), "Test Sequence")
    for cell in table.rows[0].cells:
        _set_cell_fill(cell, _HEADER_FILL)
    for cell in table.rows[1].cells[1:]:
        _set_cell_fill(cell, _HEADER_FILL)

    for row_index, item_name in enumerate(item_names, start=2):
        _set_cell_text(table.cell(row_index, 0), item_name)
        for column_index, group in enumerate(groups, start=1):
            tokens = [
                step.raw_token
                for step in group.steps
                if step.test_item == item_name
            ]
            _set_cell_text(table.cell(row_index, column_index), ",".join(tokens))
    final_row = table.rows[-1]
    _set_cell_text(final_row.cells[0], "Samples Size(sets)")
    for column_index, group in enumerate(groups, start=1):
        _set_cell_text(final_row.cells[column_index], group.sample_quantity_expression)
    for cell in final_row.cells:
        _set_cell_fill(cell, _SAMPLE_SIZE_FILL)
    _set_description_widths(table, len(groups), table_width_dxa)


def _fill_method_table(table: Table, report: TestReportDraftData) -> None:
    methods: OrderedDict[tuple[str, str, str], list[str]] = OrderedDict()
    for group in report.groups:
        for step in group.steps:
            key = (step.test_item, step.method, step.condition)
            requirements = methods.setdefault(key, [])
            if step.requirement and step.requirement not in requirements:
                requirements.append(step.requirement)
    _resize_rows(table, 1 + len(methods))
    for row_index, ((item, method, condition), requirements) in enumerate(
        methods.items(),
        start=1,
    ):
        values = (item, method, condition, "; ".join(requirements))
        for cell, value in zip(table.rows[row_index].cells, values, strict=True):
            _set_cell_text(cell, value)


def _fill_result_blocks(
    document,
    template_heading: Paragraph,
    template_table: Table,
    report: TestReportDraftData,
) -> None:
    heading_xml = deepcopy(template_heading._p)
    table_xml = deepcopy(template_table._tbl)
    _fill_result_block(template_heading, template_table, report.groups[0])
    cursor = template_table._tbl
    for group in report.groups[1:]:
        cloned_heading_xml = deepcopy(heading_xml)
        cursor.addnext(cloned_heading_xml)
        cursor = cloned_heading_xml
        cloned_table_xml = deepcopy(table_xml)
        cursor.addnext(cloned_table_xml)
        cursor = cloned_table_xml
        _fill_result_block(
            Paragraph(cloned_heading_xml, document._body),
            Table(cloned_table_xml, document._body),
            group,
        )


def _fill_result_block(heading: Paragraph, table: Table, group) -> None:
    group_label = _test_sequence_group_label(group.group_label)
    _set_paragraph_text(heading, f"Group {group_label} Test Results")
    heading.paragraph_format.keep_with_next = True
    _resize_rows(table, 1 + len(group.steps))
    llcr_indexes = tuple(
        index
        for index, candidate in enumerate(group.steps)
        if is_llcr_test_item(candidate.test_item)
    )
    for step_index, step in enumerate(group.steps):
        row_index = step_index + 1
        description = _step_description(
            group.steps,
            step_index,
            llcr_indexes,
        )
        display_requirement = _display_requirement(
            step.requirement,
            step_index=step_index,
            llcr_indexes=llcr_indexes,
        )
        result_requirement = _stage_result_requirement(
            display_requirement,
            step_index=step_index,
            llcr_indexes=llcr_indexes,
        )
        values = (
            step.raw_token,
            step.test_item,
            display_requirement,
            description,
            _default_result(result_requirement),
            "Pass",
        )
        for cell, value in zip(table.rows[row_index].cells, values, strict=True):
            _set_cell_text(cell, value)


def _fill_revision_table(table: Table, report: TestReportDraftData) -> None:
    _resize_rows(table, max(2, len(table.rows)))
    values = (
        "A",
        "All",
        "Initial draft - not released",
        report.generated_on.strftime("%d/%b/%Y"),
    )
    for cell, value in zip(table.rows[1].cells, values, strict=True):
        _set_cell_text(cell, value)


def _resize_rows(table: Table, target: int) -> None:
    while len(table.rows) < target:
        template_row = table.rows[-1]
        cloned = deepcopy(template_row._tr)
        table._tbl.append(cloned)
    while len(table.rows) > target:
        row = table.rows[-1]
        table._tbl.remove(row._tr)


def _resize_columns(table: Table, target: int) -> None:
    while len(table.columns) < target:
        grid_columns = table._tbl.tblGrid.gridCol_lst
        if not grid_columns:
            raise ValueError("E-3707_H Test Description table has no column grid.")
        table._tbl.tblGrid.append(deepcopy(grid_columns[-1]))
        for row in table.rows:
            row._tr.append(deepcopy(row._tr.tc_lst[-1]))
    while len(table.columns) > target:
        index = len(table.columns) - 1
        grid_columns = table._tbl.tblGrid.gridCol_lst
        if index < len(grid_columns):
            table._tbl.tblGrid.remove(grid_columns[index])
        for row in table.rows:
            row._tr.remove(row.cells[index]._tc)


def _table_width_dxa(table: Table) -> int:
    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    if table_width is not None and table_width.get(qn("w:type")) == "dxa":
        width = int(table_width.get(qn("w:w"), "0"))
        if width > 0:
            return width
    grid_widths = [
        int(column.get(qn("w:w"), "0"))
        for column in table._tbl.tblGrid.gridCol_lst
    ]
    total_width = sum(grid_widths)
    if total_width <= 0:
        raise ValueError("E-3707_H Test Description table has no usable width.")
    return total_width


def _set_description_widths(
    table: Table,
    group_count: int,
    total_width_dxa: int,
) -> None:
    if group_count < 1:
        raise ValueError("E-3707_H Test Description table requires a group column.")
    first_width = min(
        _PREFERRED_TEST_ITEM_WIDTH_DXA,
        max(
            _MIN_TEST_ITEM_WIDTH_DXA,
            total_width_dxa - group_count * _MIN_GROUP_WIDTH_DXA,
        ),
    )
    group_area_width = total_width_dxa - first_width
    if group_area_width < group_count:
        raise ValueError("E-3707_H Test Description table is too narrow for its groups.")
    group_width, remainder = divmod(group_area_width, group_count)
    widths = [first_width]
    widths.extend(
        group_width + (1 if index < remainder else 0)
        for index in range(group_count)
    )

    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    if table_width is None:
        raise ValueError("E-3707_H Test Description table has no width contract.")
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(total_width_dxa))
    table.autofit = False

    grid_columns = table._tbl.tblGrid.gridCol_lst
    if len(grid_columns) != len(widths):
        raise ValueError("E-3707_H Test Description table grid is inconsistent.")
    for column, width in zip(grid_columns, widths, strict=True):
        column.set(qn("w:w"), str(width))

    for row in table.rows:
        grid_index = 0
        for cell_xml in row._tr.tc_lst:
            cell_properties = cell_xml.get_or_add_tcPr()
            span_xml = cell_properties.find(qn("w:gridSpan"))
            span = 1 if span_xml is None else int(span_xml.get(qn("w:val"), "1"))
            next_grid_index = grid_index + span
            if next_grid_index > len(widths):
                raise ValueError("E-3707_H Test Description cell span is inconsistent.")
            cell_width = cell_properties.get_or_add_tcW()
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(
                qn("w:w"),
                str(sum(widths[grid_index:next_grid_index])),
            )
            grid_index = next_grid_index
        if grid_index != len(widths):
            raise ValueError("E-3707_H Test Description row grid is inconsistent.")


def _insert_page_break_before_heading(document, heading: str) -> None:
    paragraph = next(
        (
            item
            for item in document.paragraphs
            if _heading_key(item.text) == _heading_key(heading)
        ),
        None,
    )
    if paragraph is None:
        raise ValueError(f"E-3707_H template heading is missing: {heading}")
    break_paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run.append(page_break)
    break_paragraph.append(run)
    paragraph._p.addprevious(break_paragraph)


def _keep_heading_with_following_content(document, heading: str) -> None:
    paragraph = next(
        (
            item
            for item in document.paragraphs
            if _heading_key(item.text) == _heading_key(heading)
        ),
        None,
    )
    if paragraph is None:
        raise ValueError(f"E-3707_H template heading is missing: {heading}")
    paragraph.paragraph_format.keep_with_next = True


def _set_document_table_font(document, font_name: str) -> None:
    tables = list(document.tables)
    for section in document.sections:
        for header in (section.header, section.first_page_header):
            tables.extend(header.tables)
        for footer in (section.footer, section.first_page_footer):
            tables.extend(footer.tables)
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                            fonts.set(qn(f"w:{attribute}"), font_name)


def _set_cell_fill(cell: _Cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def _test_sequence_group_label(value: str) -> str:
    return re.sub(r"^Group\s+", "", value.strip(), flags=re.IGNORECASE)


def _default_result(requirement: str) -> str:
    normalized = _normalized(requirement)
    if normalized.casefold() == "no detrimental condition":
        return "No detriment"
    return re.sub(
        r"([≤≥])\s*[+-]?(?:\d+(?:\.\d*)?|\.\d+)",
        r"\1_",
        requirement.strip(),
    )


def _step_description(steps, step_index: int, llcr_indexes: tuple[int, ...]) -> str:
    step = steps[step_index]
    if not is_llcr_test_item(step.test_item):
        description = step.test_item
        if step.suffix_note:
            description = f"{description} ({step.suffix_note})"
        return description
    if len(llcr_indexes) <= 1 or step_index == llcr_indexes[0]:
        return "LLCR"
    if step_index == llcr_indexes[-1]:
        return "Final ΔR"
    previous = next(
        (
            candidate
            for candidate in reversed(steps[:step_index])
            if not is_llcr_test_item(candidate.test_item)
        ),
        None,
    )
    return f"After {previous.test_item}" if previous is not None else "LLCR"


def _stage_result_requirement(
    requirement: str,
    *,
    step_index: int,
    llcr_indexes: tuple[int, ...],
) -> str:
    if step_index not in llcr_indexes:
        return requirement
    clauses = [
        clause.strip()
        for clause in re.split(r"[;；\n]+", requirement)
        if clause.strip()
    ]
    if not clauses:
        return requirement
    if step_index == llcr_indexes[0]:
        return next(
            (clause for clause in clauses if "initial" in clause.casefold()),
            clauses[0],
        )
    return next(
        (clause for clause in clauses if "ΔR" in clause or "∆R" in clause),
        clauses[-1],
    )


def _display_requirement(
    requirement: str,
    *,
    step_index: int,
    llcr_indexes: tuple[int, ...],
) -> str:
    normalized = requirement.strip()
    if (
        llcr_indexes
        and step_index == llcr_indexes[0]
        and "initial" not in normalized.casefold()
        and re.match(r"^[≤≥]", normalized)
    ):
        return f"Initial {normalized}"
    return normalized


def _set_cell_text(cell: _Cell, text: str) -> None:
    first = cell.paragraphs[0]
    for paragraph in list(cell.paragraphs[1:]):
        cell._tc.remove(paragraph._p)
    _set_paragraph_text(first, text)


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    run_properties = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            run_properties = deepcopy(run._r.rPr)
            break
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def _display_date(value: str) -> str:
    normalized = value.strip()
    if not normalized:
        return "[To be confirmed]"
    try:
        parsed = datetime.fromisoformat(normalized.replace("Z", "+00:00"))
    except ValueError:
        return normalized
    return parsed.strftime("%b %d, %Y")


def _display_test_date_range(start_date: str, finish_date: str) -> str:
    start = start_date.strip()
    finish = finish_date.strip()
    if not start or not finish:
        return "TBD"
    return f"{_display_header_date(start)} to {_display_header_date(finish)}"


def _display_header_date(value: str) -> str:
    try:
        parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return value
    return parsed.strftime("%d/%b/%Y")


def _audit_generated_document(path: Path, report: TestReportDraftData) -> None:
    document = Document(path)
    _validate_template_contract(document, populated=True)
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if "[PRODUCT NAME]" in body_text or "[TEST DESCRIPTION]" in body_text:
        raise ValueError("Generated report still contains required identity placeholders.")
    header_text = "".join(
        text
        for section in document.sections
        for header in (section.header, section.first_page_header)
        for text in _header_text_nodes(header)
    )
    if report.report_number not in header_text:
        raise ValueError("Generated report header does not contain the report number.")
    unresolved = [
        placeholder
        for placeholder in _REQUIRED_FIRST_PAGE_HEADER_PLACEHOLDERS
        if placeholder in header_text
    ]
    if unresolved:
        raise ValueError(
            "Generated report header still contains placeholders: "
            + ", ".join(unresolved)
        )


def _header_text_nodes(header) -> tuple[str, ...]:
    return tuple(
        text_node.text or "" for text_node in header._element.xpath(".//w:t")
    )


def _normalized(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def _heading_key(value: str) -> str:
    return re.sub(r"^(\d+)\.\s*", r"\1. ", _normalized(value))
