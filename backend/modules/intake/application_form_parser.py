"""DOCX application form parser for the ConnLab intake flow."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from backend.modules.intake.application_form_parser_patterns import (
    LABEL_ALIASES,
    SAMPLE_ALIASES,
)


@dataclass(frozen=True, slots=True)
class ParsedSampleInfo:
    """Sample row extracted from an application form."""

    product_name: str | None = None
    part_number: str | None = None
    revision: str | None = None
    lot_or_traceability: str | None = None
    material: str | None = None
    plating: str | None = None
    lubricant: str | None = None
    housing_material: str | None = None
    quantity: str | None = None


@dataclass(frozen=True, slots=True)
class ParsedRequestedTestingRow:
    """One row from the application-form requested-testing table."""

    test_to_be_performed: str
    applicable_specification: str


@dataclass(frozen=True, slots=True)
class ParsedLabSection:
    """Section 2 lab fields extracted from an application form."""

    lab: str | None = None
    assigned_personnel: str | None = None
    received_date: str | None = None
    estimated_completion_date: str | None = None
    sample_condition: str | None = None


@dataclass(frozen=True, slots=True)
class ParsedApplicationForm:
    """Structured parser output for a laboratory test request form."""

    form_no: str | None = None
    form_rev: str | None = None
    reference_doc: str | None = None
    lab_test_request_number: str | None = None
    requested_by: str | None = None
    phone: str | None = None
    request_date: str | None = None
    email: str | None = None
    business_unit: str | None = None
    manufacturing_site: str | None = None
    project_number: str | None = None
    requested_completion_date: str | None = None
    results_format: str | None = None
    test_type: str | None = None
    sample_status: str | None = None
    project_type: str | None = None
    post_testing_disposition: str | None = None
    requested_testing_description: str | None = None
    confidential: str | None = None
    subcontract: str | None = None
    additional_information: str | None = None
    send_copies_recipients: str | None = None
    lab_section: ParsedLabSection = field(default_factory=ParsedLabSection)
    samples: tuple[ParsedSampleInfo, ...] = field(default_factory=tuple)
    requested_testing_rows: tuple[ParsedRequestedTestingRow, ...] = field(default_factory=tuple)


class ApplicationFormParser:
    """Parse DOCX application forms into structured intake DTOs."""

    def parse(self, path: Path) -> ParsedApplicationForm:
        """Parse one DOCX file into a structured application form DTO."""
        document = Document(path)
        label_values = _extract_label_values(document)
        _merge_footer_form_metadata(label_values, document)
        samples = _extract_sample_rows(document)
        requested_testing_rows = _extract_requested_testing_rows(document)
        if requested_testing_rows and not label_values.get("requested_testing_description"):
            label_values["requested_testing_description"] = "\n".join(
                r.test_to_be_performed for r in requested_testing_rows if r.test_to_be_performed
            )
        if not label_values.get("requested_testing_description"):
            _merge_legacy_requested_testing(label_values, document)
        _merge_additional_information_block(label_values, document)
        return ParsedApplicationForm(
            form_no=_get(label_values, "form_no"),
            form_rev=_get(label_values, "form_rev"),
            reference_doc=_get(label_values, "reference_doc"),
            lab_test_request_number=_get(label_values, "lab_test_request_number"),
            requested_by=_get(label_values, "requested_by"),
            phone=_get(label_values, "phone"),
            request_date=_get(label_values, "request_date"),
            email=_get(label_values, "email"),
            business_unit=_get(label_values, "business_unit"),
            manufacturing_site=_get(label_values, "manufacturing_site"),
            project_number=_get(label_values, "project_number"),
            requested_completion_date=_get(label_values, "requested_completion_date"),
            results_format=_get(label_values, "results_format"),
            test_type=_get(label_values, "test_type"),
            sample_status=_get(label_values, "sample_status"),
            project_type=_get(label_values, "project_type"),
            post_testing_disposition=_get(label_values, "post_testing_disposition"),
            requested_testing_description=_get(
                label_values,
                "requested_testing_description",
            ),
            confidential=_get(label_values, "confidential"),
            subcontract=_get(label_values, "subcontract"),
            additional_information=_get(label_values, "additional_information"),
            send_copies_recipients=_get(label_values, "send_copies_recipients"),
            lab_section=ParsedLabSection(
                lab=_get(label_values, "lab"),
                assigned_personnel=_get(label_values, "assigned_personnel"),
                received_date=_get(label_values, "received_date"),
                estimated_completion_date=_get(label_values, "estimated_completion_date"),
                sample_condition=_get(label_values, "sample_condition"),
            ),
            samples=tuple(samples),
            requested_testing_rows=requested_testing_rows,
        )

    def table_outline(self, path: Path, limit: int = 8) -> tuple[tuple[str, str], ...]:
        """Return a compact outline of non-empty Word tables."""
        document = Document(path)
        rows: list[tuple[str, str]] = []
        for index, table in enumerate(document.tables[:limit], start=1):
            first_text = _first_non_empty_cell(table)
            if first_text:
                rows.append((f"Table {index}", first_text[:140]))
        return tuple(rows)


def _extract_label_values(document) -> dict[str, str]:
    """Extract normalized label-value pairs from paragraphs and tables."""
    values: dict[str, str] = {}
    for paragraph in _iter_document_paragraphs(document):
        _merge_pair(values, paragraph.text)
    for table in _iter_document_tables(document):
        for row in table.rows:
            cells = [_clean(cell.text) for cell in row.cells]
            _merge_table_row(values, cells)
    _merge_content_control_values(values, document)
    return values


def _merge_table_row(values: dict[str, str], cells: list[str]) -> None:
    """Merge label-value pairs from one table row."""
    cells = _dedupe_cells(cells)
    if len(cells) == 1:
        _merge_pair(values, cells[0])
        return
    if len(cells) == 2:
        _set_value(values, cells[0], cells[1])
        return
    index = 0
    while index < len(cells) - 1:
        if (
            index + 2 < len(cells)
            and _is_known_label(cells[index + 1])
            and not _is_known_label(cells[index + 2])
        ):
            _set_value(values, cells[index + 1], cells[index + 2])
            index += 3
            continue
        _set_value(values, cells[index], cells[index + 1])
        index += 2


def _merge_pair(values: dict[str, str], text: str) -> None:
    """Merge a label-value pair from free text when possible."""
    cleaned = _clean(text)
    if not cleaned:
        return
    for separator in (":", "："):
        if separator in cleaned:
            label, value = cleaned.split(separator, 1)
            _set_value(values, label, value)
            return
    for key, aliases in LABEL_ALIASES.items():
        for alias in sorted(aliases, key=len, reverse=True):
            value = _value_after_alias(cleaned, alias)
            if value:
                values.setdefault(key, value)
                return


def _set_value(values: dict[str, str], label: str, value: str) -> None:
    """Set a normalized label value if both label and value are meaningful."""
    key = _canonical_label(label, LABEL_ALIASES)
    cleaned = _clean(value)
    if key and cleaned and not _is_known_label(cleaned):
        values.setdefault(key, cleaned)


def _extract_sample_rows(document) -> list[ParsedSampleInfo]:
    """Extract sample rows from tables with recognizable sample headers."""
    samples: list[ParsedSampleInfo] = []
    for table in _iter_document_tables(document):
        rows = [[_clean(cell.text) for cell in row.cells] for row in table.rows]
        for index, row in enumerate(rows):
            header = [_canonical_label(cell, SAMPLE_ALIASES) for cell in row]
            if not _is_sample_header(header) or index + 1 >= len(rows):
                continue
            for sample_row in rows[index + 1 :]:
                sample = _sample_from_row(header, sample_row)
                if sample:
                    samples.append(sample)
            return samples
    return samples


def _sample_from_row(
    header: list[str | None],
    row: list[str],
) -> ParsedSampleInfo | None:
    """Convert one table row to a sample DTO when it contains sample data."""
    data = {
        key: row[index]
        for index, key in enumerate(header)
        if key and index < len(row) and row[index]
    }
    if not data:
        return None
    return ParsedSampleInfo(**data)


def _extract_requested_testing_rows(document) -> tuple[ParsedRequestedTestingRow, ...]:
    """Extract requested testing rows from a two-column application-form table."""
    for table in _iter_document_tables(document):
        rows = [
            _dedupe_cells([_clean(cell.text) for cell in row.cells])
            for row in table.rows
        ]
        if not rows:
            continue
        header = [_normalize_label(cell) for cell in rows[0]]
        if "tests to be performed" not in header or "applicable specifications" not in header:
            continue
        test_col = header.index("tests to be performed")
        spec_col = header.index("applicable specifications")
        result: list[ParsedRequestedTestingRow] = []
        for row in rows[1:]:
            if test_col >= len(row) or spec_col >= len(row):
                continue
            test = row[test_col]
            spec = row[spec_col]
            if test or spec:
                result.append(ParsedRequestedTestingRow(test, spec))
        return tuple(result)
    return ()


def _merge_legacy_requested_testing(values: dict[str, str], document) -> None:
    """Fallback: extract requested testing text from header/value testing tables."""
    if values.get("requested_testing_description"):
        return
    for table in _iter_document_tables(document):
        rows = [
            _dedupe_cells([_clean(cell.text) for cell in row.cells])
            for row in table.rows
        ]
        for index, row in enumerate(rows[:-1]):
            labels = [_normalize_label(cell) for cell in row]
            if "tests to be performed" not in labels:
                continue
            column = labels.index("tests to be performed")
            if column >= len(rows[index + 1]):
                continue
            value = rows[index + 1][column]
            if value:
                values["requested_testing_description"] = value
                return


def _merge_additional_information_block(values: dict[str, str], document) -> None:
    """Extract Additional Information from a dedicated block table."""
    if values.get("additional_information"):
        return
    for table in _iter_document_tables(document):
        rows = [
            _dedupe_cells([_clean(cell.text) for cell in row.cells])
            for row in table.rows
        ]
        found = False
        for row in rows:
            for cell in row:
                if "additional information" in _normalize_label(cell):
                    found = True
                    break
            if found:
                break
        if not found:
            continue
        texts: list[str] = []
        for row in rows:
            for cell in row:
                cleaned = _clean(cell)
                if cleaned and "additional information" not in _normalize_label(cleaned) and not _is_known_label(cleaned):
                    texts.append(cleaned)
        if texts:
            values["additional_information"] = "\n".join(texts)
            return
    found_heading = False
    for child in document.element.body:
        text = _body_child_text(child)
        if not text:
            continue
        normalized = _normalize_label(text)
        if not found_heading:
            if "additional information" in normalized:
                found_heading = True
            continue
        if _is_additional_information_skip_block(normalized):
            continue
        if _is_additional_information_stop_block(normalized):
            return
        values["additional_information"] = text
        return


def _merge_content_control_values(values: dict[str, str], document) -> None:
    """Merge relevant Word content-control values that are hidden from cell.text."""
    content_control_fields = {
        "confidential": "confidential",
        "confidiential": "confidential",
        "disposition": "post_testing_disposition",
        "subcontract": "subcontract",
        "subcontracted": "subcontract",
        "test type": "test_type",
    }
    ordered_values: list[str] = []
    for control in document.element.xpath('.//*[local-name()="sdt"]'):
        names = [
            *_control_attribute_values(control, "alias"),
            *_control_attribute_values(control, "tag"),
        ]
        key = next(
            (
                content_control_fields[name]
                for raw_name in names
                if (name := _normalize_label(raw_name)) in content_control_fields
            ),
            None,
        )
        value = _clean("".join(control.xpath('./*[local-name()="sdtContent"]//*[local-name()="t"]/text()')))
        is_placeholder = not value or _is_placeholder_value(value)
        ordered_values.append("" if is_placeholder else value)
        if is_placeholder:
            continue
        if inferred_key := _infer_dropdown_field_from_value(value):
            values.setdefault(inferred_key, value)
        if key:
            values[key] = value
    _merge_ordered_section1_content_controls(values, ordered_values)


def _infer_dropdown_field_from_value(value: str) -> str | None:
    """Infer unlabeled Section 1 dropdown controls from their option text."""
    normalized = _normalize_label(value)
    if normalized in {
        "formal report customer",
        "formal report internal",
        "data and observation",
    }:
        return "results_format"
    if normalized in {
        "product process development",
        "product process qualification",
        "lab failure analysis",
        "customer specific testing",
    }:
        return "test_type"
    if normalized in {
        "new product development",
        "product extension",
        "innovation",
        "lab activities lab use only",
        "operational support",
        "cost reduction",
    }:
        return "project_type"
    if normalized in {
        "prototype",
        "production",
        "pre production",
        "engineering sample",
        "new product",
    }:
        return "sample_status"
    if _looks_like_post_testing_disposition(value):
        return "post_testing_disposition"
    return None


def _merge_ordered_section1_content_controls(
    values: dict[str, str],
    ordered_values: list[str],
) -> None:
    """Merge E-3718 Rev H content controls that have no stable alias/tag."""
    ordered_keys = _section1_content_control_order(ordered_values)
    if len(ordered_values) < 8:
        return
    for key, value in zip(ordered_keys, ordered_values, strict=False):
        if value and not _is_known_label(value):
            if key == "project_type" and _looks_like_post_testing_disposition(value):
                values.setdefault("post_testing_disposition", value)
                continue
            values.setdefault(key, value)


def _section1_content_control_order(ordered_values: list[str]) -> tuple[str, ...]:
    if (
        len(ordered_values) > 2
        and _infer_dropdown_field_from_value(ordered_values[2]) == "results_format"
    ):
        return (
            "request_date",
            "business_unit",
            "results_format",
            "requested_completion_date",
            "test_type",
            "sample_status",
            "project_type",
            "post_testing_disposition",
            "confidential",
            "subcontract",
            "lab",
            "received_date",
            "estimated_completion_date",
            "sample_condition",
        )
    return (
        "request_date",
        "business_unit",
        "manufacturing_site",
        "results_format",
        "requested_completion_date",
        "test_type",
        "sample_status",
        "project_type",
        "post_testing_disposition",
        "confidential",
        "subcontract",
        "lab",
        "received_date",
        "estimated_completion_date",
        "sample_condition",
    )


def _looks_like_post_testing_disposition(value: str) -> bool:
    normalized = _normalize_label(value)
    return normalized in {
        "send back to requestor",
        "scrap",
        "keep in the lab",
        "return samples",
        "return to requestor",
    }


def _control_attribute_values(control, name: str) -> list[str]:
    """Return w:val values for one content-control property name."""
    return [
        _clean(value)
        for value in control.xpath(
            f'./*[local-name()="sdtPr"]/*[local-name()="{name}"]/@*[local-name()="val"]'
        )
        if _clean(value)
    ]


def _first_non_empty_cell(table) -> str:
    """Return the first non-empty cell text in a Word table."""
    for row in table.rows:
        for cell in row.cells:
            value = _clean(cell.text)
            if value:
                return " ".join(value.split())
    return ""


def _iter_document_paragraphs(document):
    """Yield paragraphs from the body, headers, and footers."""
    yield from document.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _iter_document_tables(document):
    """Yield tables from the body, headers, and footers."""
    yield from document.tables
    for section in document.sections:
        yield from section.header.tables
        yield from section.footer.tables


def _body_child_text(child) -> str:
    """Return visible text from one body paragraph/table element."""
    return _clean("".join(child.xpath('.//*[local-name()="t"]/text()')))


def _is_additional_information_skip_block(normalized: str) -> bool:
    """Return whether a block between heading and content is an unrelated yes/no control."""
    return (
        "confidential tests or samples" in normalized
        or "can testing be subcontracted" in normalized
    )


def _is_additional_information_stop_block(normalized: str) -> bool:
    """Return whether Additional Information scanning has reached the next section."""
    return (
        "send copies of test results/reports to" in normalized
        or "section 2 to be completed by the testing laboratory" in normalized
        or normalized.startswith("section 2")
    )


def _merge_footer_form_metadata(values: dict[str, str], document) -> None:
    """Extract form number and revision from compact footer text."""
    footer_text = " ".join(_clean(paragraph.text) for section in document.sections for paragraph in section.footer.paragraphs)
    for section in document.sections:
        for table in section.footer.tables:
            for row in table.rows:
                footer_text = f"{footer_text} {' '.join(_clean(cell.text) for cell in row.cells)}"
    form_match = re.search(r"\b(E-\d{3,})\b", footer_text, flags=re.IGNORECASE)
    if form_match:
        values["form_no"] = form_match.group(1).upper()
    rev_match = re.search(
        r"\bRev(?:ision)?\.?\s*[:：-]?\s*([A-Z0-9]+)\b",
        footer_text,
        flags=re.IGNORECASE,
    )
    if rev_match:
        values["form_rev"] = rev_match.group(1).upper()


def _is_sample_header(header: list[str | None]) -> bool:
    """Return whether a normalized row looks like a sample table header."""
    recognized = {key for key in header if key}
    return "part_number" in recognized and len(recognized) >= 3


def _dedupe_cells(cells: list[str]) -> list[str]:
    """Remove repeated merged-cell text while preserving real columns."""
    deduped: list[str] = []
    for cell in cells:
        if deduped and cell and cell == deduped[-1]:
            continue
        deduped.append(cell)
    return deduped


def _value_after_alias(text: str, alias: str) -> str | None:
    """Return the value that follows a label alias in one cell of text."""
    pattern = rf"^\s*{re.escape(alias)}\s*(?:[:：#-]|\b)\s*(.+)$"
    match = re.match(pattern, text, flags=re.IGNORECASE)
    if not match:
        return None
    return _clean(match.group(1))


def _canonical_label(
    raw_label: str,
    aliases: dict[str, set[str]],
) -> str | None:
    """Return the canonical field name for a raw label."""
    normalized = _normalize_label(raw_label)
    for key, candidates in aliases.items():
        if normalized in candidates:
            return key
    return None


def _is_known_label(value: str) -> bool:
    """Return whether a value is actually another known field label."""
    return _canonical_label(value, LABEL_ALIASES) is not None


def _normalize_label(value: str) -> str:
    """Normalize label text for keyword matching."""
    cleaned = _clean(value).lower().replace("#", " number ")
    cleaned = re.sub(r"[^a-z0-9/]+", " ", cleaned)
    return re.sub(r"\s+", " ", cleaned).strip()


def _get(values: dict[str, str], key: str) -> str | None:
    """Return a parsed value or None when missing."""
    return values.get(key)


def _is_placeholder_value(value: str) -> bool:
    """Return whether a content-control value is an unselected placeholder."""
    return _clean(value).lower() in {
        "choose an item.",
        "click here to enter a date.",
    }


def _clean(value: str) -> str:
    """Collapse Word cell text into a single trimmed string."""
    return re.sub(r"\s+", " ", value).strip()
